"""
pic2table v2 — 图片表格识别 → Excel / Word
===========================================
策略: PaddleOCR 文本检测 + OpenCV 网格分析，而非依赖 PP-Structure。
适用于有框线表格、无框线表格、彩色表头等场景。

用法:
    python pipeline.py <图片路径> [--format excel|word|both]
    python pipeline.py <图片文件夹>
"""

import os
import sys
import argparse
import re
from pathlib import Path

os.environ["FLAGS_use_mkldnn"] = "0"

import torch  # noqa: F401 提前加载，防 shm.dll 问题
import cv2
import numpy as np
from PIL import Image

from paddleocr import PaddleOCR

from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── 引擎 ──────────────────────────────────────────────────
_ocr = None


def get_ocr():
    global _ocr
    if _ocr is None:
        print("[信息] 加载 PaddleOCR 引擎...")
        _ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
        print("[信息] OCR 引擎就绪")
    return _ocr


# ── 图片格式 ──────────────────────────────────────────────
IMG_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"}


def find_images(path):
    p = Path(path)
    if p.is_file() and p.suffix.lower() in IMG_EXTS:
        return [p]
    if p.is_dir():
        files = sorted([f for f in p.iterdir() if f.suffix.lower() in IMG_EXTS])
        if not files:
            print(f"[错误] 目录 '{path}' 中没有图片文件")
            sys.exit(1)
        return files
    print(f"[错误] 路径不存在: {path}")
    sys.exit(1)


# ── 表格线检测 ────────────────────────────────────────────
def detect_table_lines(img_cv):
    """用 OpenCV 检测表格横线和竖线，返回行列分割位置。"""
    gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
    h, w = gray.shape

    # 自适应二值化
    binary = cv2.adaptiveThreshold(
        ~gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 15, -2
    )

    # 水平线核（更长的核对噪声更鲁棒）
    hor_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(w // 15, 1), 1))
    hor_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, hor_kernel)

    # 垂直线核
    ver_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(h // 15, 1)))
    ver_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, ver_kernel)

    # 提取水平线坐标
    hor_pts = np.column_stack(np.where(hor_lines > 0))
    hor_y_vals = sorted(set(hor_pts[:, 0])) if len(hor_pts) > 0 else []

    # 提取垂直线坐标
    ver_pts = np.column_stack(np.where(ver_lines > 0))
    ver_x_vals = sorted(set(ver_pts[:, 1])) if len(ver_pts) > 0 else []

    return hor_y_vals, ver_x_vals, binary


def cluster_lines(lines, threshold=15):
    """聚类相近的线条坐标。"""
    if not lines:
        return []
    clusters = []
    current = [lines[0]]
    for l in lines[1:]:
        if abs(l - current[-1]) <= threshold:
            current.append(l)
        else:
            clusters.append(int(np.median(current)))
            current = [l]
    if current:
        clusters.append(int(np.median(current)))
    return clusters


def get_grid_cells(h_lines, v_lines):
    """根据横竖线位置生成单元格列表。"""
    if len(h_lines) < 2 or len(v_lines) < 2:
        return []
    cells = []
    for r in range(len(h_lines) - 1):
        for c in range(len(v_lines) - 1):
            cells.append(
                {
                    "row": r,
                    "col": c,
                    "x1": v_lines[c],
                    "y1": h_lines[r],
                    "x2": v_lines[c + 1],
                    "y2": h_lines[r + 1],
                }
            )
    return cells


# ── 表格识别核心 ──────────────────────────────────────────
def recognize_table(image_path):
    """识别图片中的表格内容，返回二维数组。"""
    ocr = get_ocr()

    # 读取图片
    img_pil = Image.open(str(image_path)).convert("RGB")
    img_cv = cv2.cvtColor(np.array(img_pil), cv2.COLOR_RGB2BGR)
    h_img, w_img = img_cv.shape[:2]

    # OCR 识别
    print(f"[OCR] {image_path.name}")
    ocr_result = ocr.ocr(str(image_path))

    if not ocr_result or not ocr_result[0]:
        print("[警告] 未识别到文字")
        return []

    # 提取文本框
    boxes = []
    for line in ocr_result[0]:
        pts = line[0]  # 四点坐标 [[x1,y1],[x2,y2],[x3,y3],[x4,y4]]
        text = line[1][0]
        conf = line[1][1]

        x_coords = [p[0] for p in pts]
        y_coords = [p[1] for p in pts]
        x_center = sum(x_coords) / 4
        y_center = sum(y_coords) / 4
        x_min, x_max = min(x_coords), max(x_coords)
        y_min, y_max = min(y_coords), max(y_coords)

        boxes.append(
            {
                "text": text,
                "conf": conf,
                "x_center": x_center,
                "y_center": y_center,
                "x_min": x_min,
                "x_max": x_max,
                "y_min": y_min,
                "y_max": y_max,
                "width": x_max - x_min,
                "height": y_max - y_min,
            }
        )

    print(f"  OCR: {len(boxes)} 个文本块")

    # 基于坐标聚类，不依赖表格线检测
    grid = cluster_by_coords(boxes, h_img, w_img)
    grid = trim_grid(grid)
    print(f"  坐标聚类: {len(grid)} 行 x {len(grid[0]) if grid else 0} 列")

    return grid


# ── 有框线表格：文本分配到单元格 ──────────────────────────
def assign_text_to_cells(boxes, cells, h_lines, v_lines):
    """将 OCR 文本块分配到对应的网格单元格。"""
    n_rows = len(h_lines) - 1
    n_cols = len(v_lines) - 1

    grid = [["" for _ in range(n_cols)] for _ in range(n_rows)]

    for box in boxes:
        cx, cy = box["x_center"], box["y_center"]
        # 找最近的单元格
        min_dist = float("inf")
        best_r, best_c = -1, -1
        for cell in cells:
            if cell["x1"] <= cx <= cell["x2"] and cell["y1"] <= cy <= cell["y2"]:
                # 直接在单元格内
                r, c = cell["row"], cell["col"]
                if grid[r][c]:
                    grid[r][c] += " "
                grid[r][c] += box["text"]
                min_dist = -1
                best_r, best_c = r, c
                break
            else:
                # 找最近的
                dist = min(
                    abs(cx - cell["x1"]),
                    abs(cx - cell["x2"]),
                    abs(cy - cell["y1"]),
                    abs(cy - cell["y2"]),
                )
                if dist < min_dist:
                    min_dist = dist
                    best_r, best_c = cell["row"], cell["col"]

        if min_dist > 0 and best_r >= 0:
            if grid[best_r][best_c]:
                grid[best_r][best_c] += " "
            grid[best_r][best_c] += box["text"]

    return trim_grid(grid)


def trim_grid(grid):
    """删除全空的行和列。"""
    if not grid:
        return grid
    # 删除全空行
    grid = [row for row in grid if any(cell.strip() for cell in row)]
    if not grid:
        return grid
    # 删除全空列
    n_cols = len(grid[0])
    keep_cols = []
    for c in range(n_cols):
        if any(row[c].strip() if c < len(row) else False for row in grid):
            keep_cols.append(c)
    grid = [[row[c] for c in keep_cols] for row in grid]
    return grid


# ── 无框线表格：Y 轴聚类排序 + X 轴列对齐 ─────────────────
def cluster_by_coords(boxes, h_img, w_img):
    """基于 Y 坐标聚类行，X 坐标检测列边界后对齐。"""
    if not boxes:
        return []

    # Y 中心点聚类 → 分出行
    y_centers = sorted([b["y_center"] for b in boxes])
    gaps = []
    for i in range(1, len(y_centers)):
        gaps.append(y_centers[i] - y_centers[i - 1])

    if not gaps:
        return [[b["text"] for b in sorted(boxes, key=lambda x: x["x_center"])]]

    # 用 1.5 倍中位数间隔作为行分割阈值
    median_gap = np.median(gaps) if gaps else 20
    row_threshold = max(median_gap * 1.5, 15)

    # 按 Y 分组
    sorted_boxes = sorted(boxes, key=lambda b: b["y_center"])
    rows = [[sorted_boxes[0]]]
    for b in sorted_boxes[1:]:
        if b["y_center"] - rows[-1][-1]["y_center"] > row_threshold:
            rows.append([b])
        else:
            rows[-1].append(b)

    # === 列边界检测 ===
    # 收集所有 X 中心点，聚类出列边界
    all_x = sorted([b["x_center"] for b in boxes])
    x_gaps = []
    for i in range(1, len(all_x)):
        x_gaps.append(all_x[i] - all_x[i - 1])
    x_median_gap = np.median(x_gaps) if x_gaps else 20
    x_threshold = max(x_median_gap * 2.0, 30)

    col_boundaries = []
    for b in sorted(boxes, key=lambda x: x["x_center"]):
        cx = b["x_center"]
        if not col_boundaries or cx - col_boundaries[-1][1] > x_threshold:
            col_boundaries.append([cx, cx])
        else:
            col_boundaries[-1][1] = cx
    # 列中心点
    col_centers = [(lo + hi) / 2 for lo, hi in col_boundaries]

    # === 分配到网格 ===
    grid = []
    for row in rows:
        row.sort(key=lambda b: b["x_center"])
        row_cells = [""] * len(col_centers)
        for b in row:
            # 找最近的列
            dists = [abs(b["x_center"] - cc) for cc in col_centers]
            best_c = dists.index(min(dists))
            if row_cells[best_c]:
                row_cells[best_c] += " "
            row_cells[best_c] = b["text"]
        grid.append(row_cells)

    return grid


# ── 导出 Excel ────────────────────────────────────────────
def to_excel(grid, output_path, image_name):
    wb = Workbook()
    ws = wb.active
    ws.title = "表格识别结果"

    header_font = Font(bold=True, size=11, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="4472C4")
    cell_font = Font(size=10)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for r_idx, row_data in enumerate(grid):
        for c_idx, val in enumerate(row_data):
            cell = ws.cell(row=r_idx + 1, column=c_idx + 1, value=val)
            cell.font = header_font if r_idx == 0 else cell_font
            cell.alignment = center
            cell.border = thin_border
            if r_idx == 0:
                cell.fill = header_fill

    # 自动列宽
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.value:
                length = sum(2 if ord(c) > 127 else 1 for c in str(cell.value))
                if length > max_len:
                    max_len = length
        ws.column_dimensions[col_letter].width = min(max_len + 4, 60)

    wb.save(output_path)
    print(f"  [Excel] {output_path.name}")


# ── 导出 Word ─────────────────────────────────────────────
def to_word(grid, output_path, image_name):
    doc = Document()
    doc.add_heading(f"表格识别结果 — {image_name}", level=1)

    if grid:
        n_rows, n_cols = len(grid), max(len(r) for r in grid)
        word_table = doc.add_table(rows=n_rows, cols=max(n_cols, 1))
        word_table.style = "Table Grid"
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, row_data in enumerate(grid):
            for c_idx, val in enumerate(row_data):
                cell = word_table.rows[r_idx].cells[c_idx]
                cell.text = val
                if r_idx == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True

    doc.save(output_path)
    print(f"  [Word] {output_path.name}")


# ── 主流程 ────────────────────────────────────────────────
def process(images, output_dir, fmt):
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    for img_path in images:
        print(f"\n{'=' * 60}")
        grid = recognize_table(img_path)

        if not grid or all(not cell for row in grid for cell in row):
            print(f"[跳过] 未提取到数据: {img_path.name}")
            continue

        print(f"  表格: {len(grid)} 行 x {max(len(r) for r in grid)} 列")

        stem = img_path.stem
        if fmt in ("excel", "both"):
            to_excel(grid, output_dir / f"{stem}.xlsx", stem)
        if fmt in ("word", "both"):
            to_word(grid, output_dir / f"{stem}.docx", stem)


def main():
    parser = argparse.ArgumentParser(
        description="pic2table v2 — 图片表格识别 (PaddleOCR + OpenCV)"
    )
    parser.add_argument("input", help="图片路径或文件夹")
    parser.add_argument(
        "--output", "-o", default=None, help="输出目录（默认: input 同级的 output）"
    )
    parser.add_argument(
        "--format",
        "-f",
        choices=["excel", "word", "both"],
        default="excel",
        help="输出格式",
    )
    args = parser.parse_args()

    images = find_images(args.input)
    output_dir = args.output or Path(args.input).parent / "output"

    print(f"发现 {len(images)} 张图片")
    process(images, output_dir, args.format)
    print(f"\n{'=' * 60}")
    print(f"完成。输出目录: {Path(output_dir).resolve()}")


if __name__ == "__main__":
    main()
