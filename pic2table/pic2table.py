"""
pic2table — 图片表格识别 → Excel / Word
=========================================
基于 PaddleOCR PP-Structure，识别图片中的表格并输出结构化文档。

用法:
    python pic2table.py <图片路径> [--output OUTPUT] [--format excel|word|both]
    python pic2table.py <图片文件夹>              # 批量处理目录下所有图片
    python pic2table.py <图片路径> --format word  # 输出 Word 文档

依赖: paddlepaddle, paddleocr, openpyxl, python-docx
"""

import os
import sys
import argparse
import shutil
from pathlib import Path

# ── 修 torch shm.dll 加载问题 ─────────────────────────────
import torch  # noqa: F401 提前加载，防止 albumentations 触发异常

# ── 导入 PPStructure ──────────────────────────────────────
from paddleocr import PPStructure
from PIL import Image

# ── Excel ──────────────────────────────────────────────────
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

# ── Word ───────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ───────────────────────── 引擎（全局单例）─────────────────
_engine = None


def get_engine():
    global _engine
    if _engine is None:
        print("[信息] 加载 PP-Structure 表格识别引擎...")
        _engine = PPStructure(show_log=False, lang="ch")
        print("[信息] 引擎就绪")
    return _engine


# ───────────────────────── 支持的图片格式 ──────────────────
IMG_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"}


def find_images(path):
    p = Path(path)
    if p.is_file() and p.suffix.lower() in IMG_EXTS:
        return [p]
    if p.is_dir():
        files = sorted([f for f in p.iterdir() if f.suffix.lower() in IMG_EXTS])
        if not files:
            print(f"[错误] 目录 '{path}' 中没有找到图片文件")
            sys.exit(1)
        return files
    print(f"[错误] 路径不存在: {path}")
    sys.exit(1)


# ───────────────────────── 识别核心 ────────────────────────
def recognize(image_path):
    engine = get_engine()
    print(f"[识别] {image_path.name}")

    result = engine(str(image_path))

    tables = []
    for block in result:
        if block.get("type") == "table":
            html = block.get("res", {}).get("html", "")
            if html:
                tables.append(html)

    if not tables:
        print(f"[警告] 未检测到表格: {image_path.name}")

    return tables


# ───────────────────────── HTML 表格 → 二维数组 ──────────
def html_to_grid(html):
    """极简 HTML table 解析，提取单元格文本为二维列表。"""
    import re

    # 去掉换行和多余空格
    html = re.sub(r">\s+<", "><", html)
    html = html.strip()

    rows_raw = re.findall(r"<tr>(.*?)</tr>", html, re.DOTALL)
    grid = []
    for row_html in rows_raw:
        cells = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", row_html, re.DOTALL)
        row = []
        for cell in cells:
            # 去掉内部标签
            text = re.sub(r"<[^>]+>", "", cell).strip()
            row.append(text)
        grid.append(row)

    return grid


# ───────────────────────── 导出 Excel ─────────────────────
def to_excel(tables, output_path, image_name):
    wb = Workbook()
    ws = wb.active
    ws.title = "表格识别结果"

    # 样式
    header_font = Font(bold=True, size=11, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="4472C4")
    cell_font = Font(size=10)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)

    current_row = 1

    for idx, html in enumerate(tables, 1):
        grid = html_to_grid(html)
        if not grid:
            continue

        # 表格标题
        if len(tables) > 1:
            ws.merge_cells(
                start_row=current_row,
                start_column=1,
                end_row=current_row,
                end_column=len(grid[0]) if grid else 1,
            )
            title_cell = ws.cell(row=current_row, column=1, value=f"表格 {idx}")
            title_cell.font = Font(bold=True, size=12)
            current_row += 1

        # 写入数据
        for r_idx, row_data in enumerate(grid):
            for c_idx, val in enumerate(row_data):
                cell = ws.cell(row=current_row, column=c_idx + 1, value=val)
                cell.font = header_font if r_idx == 0 else cell_font
                cell.alignment = center
                cell.border = thin_border
                if r_idx == 0:
                    cell.fill = header_fill
            current_row += 1

        current_row += 1  # 表格间空行

    # 自动列宽
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.value:
                # 中文字符算 2 个宽度
                val = str(cell.value)
                length = sum(2 if ord(c) > 127 else 1 for c in val)
                if length > max_len:
                    max_len = length
        ws.column_dimensions[col_letter].width = min(max_len + 4, 60)

    wb.save(output_path)
    print(f"[Excel] 已保存: {output_path}")


# ───────────────────────── 导出 Word ──────────────────────
def to_word(tables, output_path, image_name):
    doc = Document()

    # 标题
    doc.add_heading(f"图片表格识别结果 — {image_name}", level=1)

    for idx, html in enumerate(tables, 1):
        grid = html_to_grid(html)
        if not grid:
            continue

        if len(tables) > 1:
            doc.add_heading(f"表格 {idx}", level=2)

        # 写入 Word 表格
        word_table = doc.add_table(rows=len(grid), cols=len(grid[0]) if grid else 1)
        word_table.style = "Table Grid"
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, row_data in enumerate(grid):
            for c_idx, val in enumerate(row_data):
                cell = word_table.rows[r_idx].cells[c_idx]
                cell.text = val
                # 表头加粗
                if r_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        doc.add_paragraph()  # 表格间间隔

    doc.save(output_path)
    print(f"[Word] 已保存: {output_path}")


# ───────────────────────── 主流程 ─────────────────────────
def process(images, output_dir, fmt):
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    for img_path in images:
        print(f"\n{'=' * 60}")
        tables = recognize(img_path)

        if not tables:
            continue

        stem = img_path.stem

        if fmt in ("excel", "both"):
            to_excel(tables, output_dir / f"{stem}.xlsx", stem)

        if fmt in ("word", "both"):
            to_word(tables, output_dir / f"{stem}.docx", stem)


def main():
    parser = argparse.ArgumentParser(
        description="pic2table — 图片表格识别 → Excel / Word",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例:\n"
            "  python pic2table.py input/报表.jpg\n"
            "  python pic2table.py input/ --format word\n"
            "  python pic2table.py input/发票.jpg --format both\n"
        ),
    )
    parser.add_argument("input", help="图片文件或图片文件夹路径")
    parser.add_argument(
        "--output",
        "-o",
        default=None,
        help="输出目录（默认: input 同级的 output 目录）",
    )
    parser.add_argument(
        "--format",
        "-f",
        choices=["excel", "word", "both"],
        default="excel",
        help="输出格式（默认: excel）",
    )

    args = parser.parse_args()

    images = find_images(args.input)

    output_dir = args.output or Path(args.input).parent / "output"

    if len(images) == 1:
        print(f"发现 1 张图片")
    else:
        print(f"发现 {len(images)} 张图片（批量处理）")

    process(images, output_dir, args.format)

    print(f"\n{'=' * 60}")
    print(f"处理完成。输出目录: {Path(output_dir).resolve()}")


if __name__ == "__main__":
    main()
