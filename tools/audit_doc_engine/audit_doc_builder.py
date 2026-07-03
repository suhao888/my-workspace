"""
审计文档专业排版构建器 — AuditDocBuilder
=========================================
基于 python-docx 封装，提供面向中国审计报告的专业排版能力。

设计参考:
  - gertalot/docx-skill (GitHub) — 品牌化文档构建模式
  - python-docx-template (elapouya) — Jinja2 模板引擎
  - CAS 企业会计准则 — 财务报表附注披露要求

功能:
  - 封面页 (审计报告标准封面)
  - 页眉/页脚 (页码、公司名称)
  - 目录
  - 专业样式体系 (标题/正文/表格)
  - 审计报告专用要素 (标题段、签名段)
  - 表格美化 (交替行色、表头样式)
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional, Dict, List, Tuple

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# ============================================================
# 颜色常量 — 审计报告标准配色
# ============================================================
COLOR_PRIMARY = "#1B3A5C"  # 深蓝 — 一级标题
COLOR_ACCENT = "#C41E3A"  # 暗红 — 强调/警示
COLOR_BODY = "#333333"  # 深灰 — 正文
COLOR_LIGHT_BG = "#F5F7FA"  # 浅灰 — 表头背景
COLOR_BORDER = "#D0D5DD"  # 边框
COLOR_TABLE_ALT = "#F8F9FB"  # 表格交替行
COLOR_WHITE = "#FFFFFF"

# ============================================================
# 字体常量
# ============================================================
FONT_CN_BODY = "宋体"  # 正文
FONT_CN_HEADING = "黑体"  # 标题
FONT_EN = "Times New Roman"  # 英文/数字


# ============================================================
# 底层 XML 工具函数
# ============================================================


def _set_cell_border(cell, **kwargs):
    """设置单元格边框。kwargs: top/left/bottom/right = (sz, color, val)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            sz, color, val = kwargs[edge]
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), val)
            element.set(qn("w:sz"), str(sz))
            element.set(qn("w:color"), color)
            tcBorders.append(element)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, color_hex: str):
    """设置单元格背景色"""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_paragraph_border(para, edge: str, sz: int = 4, color: str = "#D0D5DD"):
    """设置段落边框"""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    element = OxmlElement(f"w:{edge}")
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), str(sz))
    element.set(qn("w:space"), "1")
    element.set(qn("w:color"), color)
    pBdr.append(element)
    pPr.append(pBdr)


def _add_field(para, field_code: str):
    """插入 Word 域代码 (PAGE / NUMPAGES / TOC 等)"""
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    run2._r.append(instr)

    run3 = para.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    # placeholder
    run4 = para.add_run("1")
    run4.font.size = Pt(9)
    run4.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    run5 = para.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def _remove_table_borders(table):
    """移除表格所有边框"""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f"<w:tblBorders {nsdecls('w')}>"
        '<w:top w:val="none" w:sz="0" w:space="0"/>'
        '<w:left w:val="none" w:sz="0" w:space="0"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0"/>'
        '<w:right w:val="none" w:sz="0" w:space="0"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0"/>'
        "</w:tblBorders>"
    )
    tblPr.append(borders)


def _set_run_font(
    run,
    cn_font: str,
    en_font: str = FONT_EN,
    size: int = 12,
    bold: bool = False,
    color: Optional[str] = None,
):
    """设置 run 的中英文字体"""
    run.font.name = en_font
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)
    if color:
        run.font.color.rgb = RGBColor(
            int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
        )


# ============================================================
# AuditDocBuilder — 主构建器
# ============================================================


class AuditDocBuilder:
    """审计文档专业排版构建器。

    用法:
        builder = AuditDocBuilder()
        builder.setup_page()
        builder.setup_styles()
        builder.add_cover("专项审查审计报告", "中国石油华北油田公司")
        builder.add_toc()
        builder.add_heading_1("一、项目背景")
        builder.add_body("正文内容...")
        builder.add_table_from_data(headers, rows)
        builder.save("output.docx")
    """

    def __init__(self):
        self.doc = Document()
        self._body_style_name = "audit_body"
        self._heading1_name = "audit_heading1"
        self._heading2_name = "audit_heading2"
        self._heading3_name = "audit_heading3"

    # ── 页面设置 ────────────────────────────────────────────────

    def setup_page(self, margins: Optional[Dict[str, float]] = None):
        """设置页面大小和边距。"""
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)  # A4
        section.page_height = Cm(29.7)
        m = margins or {"left": 3.0, "right": 2.5, "top": 2.54, "bottom": 2.54}
        section.left_margin = Cm(m["left"])
        section.right_margin = Cm(m["right"])
        section.top_margin = Cm(m["top"])
        section.bottom_margin = Cm(m["bottom"])
        return self

    # ── 样式体系 ────────────────────────────────────────────────

    def setup_styles(self):
        """建立文档样式体系。"""
        styles = self.doc.styles

        # Normal — 正文基准
        normal = styles["Normal"]
        normal.font.name = FONT_EN
        normal.font.size = Pt(12)
        normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        rPr = normal.element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), FONT_CN_BODY)
        rPr.insert(0, rFonts)
        pf = normal.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)

        # 创建自定义样式
        self._create_heading_style(
            self._heading1_name, Pt(16), True, COLOR_PRIMARY, Pt(24), Pt(12)
        )
        self._create_heading_style(
            self._heading2_name, Pt(14), True, COLOR_PRIMARY, Pt(18), Pt(8)
        )
        self._create_heading_style(
            self._heading3_name, Pt(12), True, COLOR_PRIMARY, Pt(12), Pt(6)
        )
        self._create_body_style()

        # 列表样式
        for name in ("List Bullet", "List Number"):
            try:
                ls = styles[name]
                _set_run_font(ls.font, FONT_CN_BODY, FONT_EN, 12)
                ls.paragraph_format.line_spacing = 1.5
            except KeyError:
                pass

        return self

    def _create_heading_style(
        self, name: str, size: Pt, bold: bool, color_hex: str, before: Pt, after: Pt
    ):
        """创建标题样式"""
        style = self.doc.styles.add_style(name, 1)  # paragraph style
        style.font.name = FONT_EN
        style.font.size = size
        style.font.bold = bold
        style.font.color.rgb = RGBColor(
            int(color_hex[1:3], 16), int(color_hex[3:5], 16), int(color_hex[5:7], 16)
        )
        rPr = style.element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), FONT_CN_HEADING)
        rPr.insert(0, rFonts)
        pf = style.paragraph_format
        pf.space_before = before
        pf.space_after = after
        pf.keep_with_next = True
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5

    def _create_body_style(self):
        """创建正文样式（首行缩进2字符）"""
        style = self.doc.styles.add_style(self._body_style_name, 1)
        _set_run_font(style.font, FONT_CN_BODY, FONT_EN, 12)
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pf = style.paragraph_format
        pf.first_line_indent = Pt(24)  # 约2个中文字符
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.space_after = Pt(2)

    # ── 封面页 ──────────────────────────────────────────────────

    def add_cover(
        self,
        title: str,
        subtitle: str = "",
        date_text: str = "",
        confidential: str = "",
    ):
        """生成审计报告标准封面页。

        Args:
            title: 报告主标题
            subtitle: 副标题/公司名称
            date_text: 日期文本
            confidential: 密级标识
        """
        # 顶部密级
        if confidential:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(confidential)
            _set_run_font(r, FONT_CN_BODY, FONT_EN, 11, bold=False, color=COLOR_ACCENT)
            p.paragraph_format.space_after = Pt(80)

        else:
            # 上边距占位
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(100)

        # 公司名称
        if subtitle:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(subtitle)
            _set_run_font(
                r, FONT_CN_HEADING, FONT_EN, 18, bold=False, color=COLOR_PRIMARY
            )
            p.paragraph_format.space_after = Pt(6)

        # 装饰线
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_border(p, "bottom", sz=8, color=COLOR_PRIMARY)
        p.paragraph_format.space_after = Pt(20)

        # 报告标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        _set_run_font(r, FONT_CN_HEADING, FONT_EN, 26, bold=True, color=COLOR_PRIMARY)
        p.paragraph_format.space_after = Pt(12)

        # 副装饰线
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_border(p, "bottom", sz=4, color="#CCCCCC")
        p.paragraph_format.space_after = Pt(60)

        # 日期
        if date_text:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(date_text)
            _set_run_font(r, FONT_CN_BODY, FONT_EN, 14, color=COLOR_BODY)
            p.paragraph_format.space_after = Pt(20)

        self.doc.add_page_break()
        return self

    # ── 目录 ────────────────────────────────────────────────────

    def add_toc(self, title: str = "目  录", levels: int = 3):
        """插入目录域。

        注意: 目录在 Word 中按 Ctrl+A → F9 更新后才显示实际内容。
        """
        p = self.doc.add_paragraph()
        r = p.add_run(title)
        _set_run_font(r, FONT_CN_HEADING, FONT_EN, 16, bold=True, color=COLOR_PRIMARY)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(20)

        _set_paragraph_border(p, "bottom", sz=4, color=COLOR_PRIMARY)

        p_toc = self.doc.add_paragraph()
        _add_field(p_toc, f'TOC \\o "1-{levels}" \\h \\z \\u')

        p_hint = self.doc.add_paragraph()
        r = p_hint.add_run("（在 Word 中按 Ctrl+A → F9 更新目录）")
        _set_run_font(r, FONT_CN_BODY, FONT_EN, 9, color="#999999")
        p_hint.paragraph_format.space_before = Pt(6)

        self.doc.add_page_break()
        return self

    # ── 页眉页脚 ────────────────────────────────────────────────

    def setup_header(self, text: str = ""):
        """设置页眉（首页不同，首页无页眉）"""
        section = self.doc.sections[0]
        section.different_first_page_header_footer = True

        header = section.header
        header.is_linked_to_previous = False

        # 清除默认段落
        for p in header.paragraphs:
            p._p.getparent().remove(p._p)

        p = header.add_paragraph()
        if text:
            r = p.add_run(text)
            _set_run_font(r, FONT_CN_BODY, FONT_EN, 9, color="#666666")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_paragraph_border(p, "bottom", sz=4, color="#CCCCCC")
        return self

    def setup_footer(self, left: str = "", right: str = "page_number"):
        """设置页脚。

        Args:
            left: 页脚左侧文本（如 "第1页"）
            right: "page_number"=页码 或 自定义文本
        """
        section = self.doc.sections[0]
        section.different_first_page_header_footer = True

        footer = section.footer
        footer.is_linked_to_previous = False

        for p in footer.paragraphs:
            p._p.getparent().remove(p._p)

        text_width = section.page_width - section.left_margin - section.right_margin

        p = footer.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)

        # 左侧
        if left:
            r = p.add_run(left)
            _set_run_font(r, FONT_CN_BODY, FONT_EN, 9, color="#666666")

        # 右侧 (制表位分隔)
        r = p.add_run()
        r.add_tab()

        if right == "page_number":
            _add_field(p, "PAGE")
        else:
            r = p.add_run(right)
            _set_run_font(r, FONT_CN_BODY, FONT_EN, 9, color="#666666")

        _set_paragraph_border(p, "top", sz=4, color="#CCCCCC")
        return self

    # ── 内容元素 ────────────────────────────────────────────────

    def add_heading_1(self, text: str):
        """一级标题"""
        p = self.doc.add_paragraph(style=self._heading1_name)
        r = p.add_run(text)
        _set_run_font(r, FONT_CN_HEADING, FONT_EN, 16, bold=True, color=COLOR_PRIMARY)
        return self

    def add_heading_2(self, text: str):
        """二级标题"""
        p = self.doc.add_paragraph(style=self._heading2_name)
        r = p.add_run(text)
        _set_run_font(r, FONT_CN_HEADING, FONT_EN, 14, bold=True, color=COLOR_PRIMARY)
        return self

    def add_heading_3(self, text: str):
        """三级标题"""
        p = self.doc.add_paragraph(style=self._heading3_name)
        r = p.add_run(text)
        _set_run_font(r, FONT_CN_HEADING, FONT_EN, 12, bold=True, color=COLOR_PRIMARY)
        return self

    def add_body(self, text: str):
        """正文段落（首行缩进2字符）"""
        if not text:
            return self
        p = self.doc.add_paragraph(style=self._body_style_name)
        r = p.add_run(text)
        _set_run_font(r, FONT_CN_BODY, FONT_EN, 12, color=COLOR_BODY)
        return self

    def add_body_no_indent(self, text: str):
        """正文段落（无缩进）"""
        if not text:
            return self
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        _set_run_font(r, FONT_CN_BODY, FONT_EN, 12, color=COLOR_BODY)
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        return self

    def add_signature_block(self, items: List[Tuple[str, str]]):
        """签名块。items = [(label, value), ...]"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.first_line_indent = Pt(0)

        for i, (label, value) in enumerate(items):
            r = p.add_run(f"{label}：{value}")
            _set_run_font(r, FONT_CN_BODY, FONT_EN, 12, color=COLOR_BODY)
            if i < len(items) - 1:
                r2 = p.add_run("\n")
                _set_run_font(r2, FONT_CN_BODY, FONT_EN, 12)
        p.paragraph_format.line_spacing = 2.0
        return self

    def add_space(self, pt: float = 12):
        """添加空行"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(pt)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run("")
        r.font.size = Pt(pt)
        return self

    def add_page_break(self):
        """分页"""
        self.doc.add_page_break()
        return self

    def add_note_box(self, text: str):
        """添加带边框的提示框"""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        _set_paragraph_border(p, "top", sz=6, color=COLOR_ACCENT)
        _set_paragraph_border(p, "bottom", sz=6, color=COLOR_ACCENT)
        r = p.add_run(text)
        _set_run_font(r, FONT_CN_BODY, FONT_EN, 11, color=COLOR_ACCENT)
        r.italic = True
        return self

    # ── 表格 ────────────────────────────────────────────────────

    def add_table_from_data(
        self,
        headers: List[str],
        rows: List[List],
        col_widths: Optional[List[float]] = None,
        caption: str = "",
    ):
        """从数据创建专业格式化表格。

        Args:
            headers: 表头列表
            rows: 数据行列表
            col_widths: 列宽比例 (如 [2, 3, 2])，自动匹配
            caption: 表题
        """
        if caption:
            p = self.doc.add_paragraph()
            r = p.add_run(caption)
            _set_run_font(r, FONT_CN_HEADING, FONT_EN, 11, bold=True, color=COLOR_BODY)
            p.paragraph_format.space_after = Pt(4)

        ncols = len(headers)
        nrows = len(rows) + 1
        table = self.doc.add_table(rows=nrows, cols=ncols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # 表头
        for j, header_text in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(header_text)
            _set_run_font(
                r, FONT_CN_HEADING, FONT_EN, 10.5, bold=True, color=COLOR_WHITE
            )
            _set_cell_shading(cell, COLOR_PRIMARY)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 数据行
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                cell = table.rows[i + 1].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if isinstance(cell_text, (int, float))
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                r = p.add_run(str(cell_text))
                _set_run_font(r, FONT_CN_BODY, FONT_EN, 10.5, color=COLOR_BODY)

                # 交替行背景
                if i % 2 == 1:
                    _set_cell_shading(cell, COLOR_TABLE_ALT)

        # 列宽
        if col_widths and len(col_widths) == ncols:
            total = sum(col_widths)
            available = 15.0  # cm
            for j, w in enumerate(col_widths):
                for row in table.rows:
                    row.cells[j].width = Cm(available * w / total)

        self.add_space(6)
        return self

    def add_key_value_table(
        self, items: List[Tuple[str, str]], label_width: float = 3.5
    ):
        """键值对表格 (如 项目：XXX 类型：XXX)"""
        table = self.doc.add_table(rows=len(items), cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for i, (key, value) in enumerate(items):
            # 标签列
            cell_k = table.rows[i].cells[0]
            cell_k.text = ""
            _set_cell_shading(cell_k, COLOR_LIGHT_BG)
            p = cell_k.paragraphs[0]
            r = p.add_run(key)
            _set_run_font(
                r, FONT_CN_HEADING, FONT_EN, 10.5, bold=True, color=COLOR_PRIMARY
            )
            cell_k.width = Cm(label_width)

            # 值列
            cell_v = table.rows[i].cells[1]
            cell_v.text = ""
            p = cell_v.paragraphs[0]
            r = p.add_run(str(value))
            _set_run_font(r, FONT_CN_BODY, FONT_EN, 10.5, color=COLOR_BODY)

        self.add_space(6)
        return self

    # ── 审计报告专用 ────────────────────────────────────────────

    def add_audit_opinion(self, opinion_text: str, opinion_type: str = "无保留意见"):
        """添加审计意见段（带样式标记）"""
        # 意见类型标签
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

        type_colors = {
            "无保留意见": "#006600",
            "保留意见": "#CC6600",
            "否定意见": COLOR_ACCENT,
            "无法表示意见": "#990000",
        }
        tag_color = type_colors.get(opinion_type, COLOR_BODY)

        r = p.add_run(f"【{opinion_type}】")
        _set_run_font(r, FONT_CN_HEADING, FONT_EN, 12, bold=True, color=tag_color)

        # 意见正文
        self.add_body(opinion_text)
        return self

    def add_finding(self, title: str, detail: str, level: str = "一般"):
        """添加审计发现项（现象→依据→影响→建议）"""
        levels = {"高": COLOR_ACCENT, "中": "#CC6600", "一般": COLOR_BODY}

        # 标题带级别标签
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
        r = p.add_run(f"▎{title}")
        _set_run_font(r, FONT_CN_HEADING, FONT_EN, 12, bold=True, color=COLOR_PRIMARY)

        level_color = levels.get(level, COLOR_BODY)
        r2 = p.add_run(f"  [{level}]")
        _set_run_font(r2, FONT_CN_HEADING, FONT_EN, 10, bold=True, color=level_color)

        # 详情
        self.add_body(detail)
        return self

    # ── 保存 ────────────────────────────────────────────────────

    def save(self, path: str):
        """保存文档。"""
        self.doc.save(str(path))
        return self


# ============================================================
# 快捷函数 — 一行生成专业文档
# ============================================================


def quick_report(
    title: str,
    subtitle: str,
    body_sections: List[Tuple[str, str]],
    output_path: str,
    date_text: str = "",
):
    """快速生成结构化审计报告。

    Args:
        title: 报告标题
        subtitle: 副标题
        body_sections: [(标题, 正文内容), ...]
        output_path: 输出路径
        date_text: 日期
    """
    builder = AuditDocBuilder()
    (
        builder.setup_page()
        .setup_styles()
        .add_cover(title, subtitle, date_text)
        .setup_header(subtitle)
        .setup_footer()
    )

    for i, (sec_title, sec_body) in enumerate(body_sections):
        builder.add_heading_1(sec_title)
        for para_text in sec_body.split("\n"):
            para_text = para_text.strip()
            if para_text:
                builder.add_body(para_text)

    builder.save(output_path)
    return output_path
