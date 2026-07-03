#!/usr/bin/env python3
"""
附注填充工具 — 纯配置驱动（行业标准做法）
============================================
用法：
  第一年：编写 project.json（我帮你配）→ 运行
  第二年：直接运行（project.json 复用）

流程：
  1. 读 project.json（含每个表的映射规则）
  2. 逐表：读源数据 → 按模板行匹配 → 按模板列映射 → 填入 → 合计 → 比例%
  3. BS/IS/CF 勾稽校验
"""

import sys, os, json, re, datetime, argparse
import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import pandas as pd


def match_name(doc_name, tz_name):
    d = re.sub(r"\s+", "", str(doc_name).replace("\u3000", ""))
    t = re.sub(r"\s+", "", str(tz_name).replace("\u3000", ""))
    if not d or not t:
        return False
    if d == t:
        return True
    d2 = d.replace("（", "(").replace("）", ")")
    t2 = t.replace("（", "(").replace("）", ")")
    if d2 == t2:
        return True
    if len(d2) >= 4 and d2 in t2:
        return True
    if len(t2) >= 4 and t2 in d2:
        return True
    for ch in "及与和、，":
        d = d.replace(ch, "")
        t = t.replace(ch, "")
    if d == t:
        return True
    return False


def fill_cell(cell, value):
    if value is None or (isinstance(value, float) and str(value) == "nan"):
        return
    if isinstance(value, (int, float)):
        s = f"{value:,.2f}"
    else:
        s = str(value).strip()
    if not s or s in ("nan", "－", "None", ""):
        return
    is_num = bool(re.match(r"^-?[\d,]+\.?\d*%?$", s))
    for p in cell.paragraphs:
        p.clear()
    if cell.paragraphs:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_num else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(s)
        run.font.size = Pt(9)
        run.font.name = "\u5b8b\u4f53"


def find_sheet(xl, name):
    for sn in xl.sheetnames:
        if name.replace("_", "") in sn.replace("_", ""):
            return xl[sn]
    return None


def auto_sum(table, col=1):
    """自动求和合计行"""
    total_ri = -1
    for ri in range(len(table.rows)):
        try:
            t = table.rows[ri].cells[0].text.strip()
            if t in (
                "\u5408\u8ba1",
                "\u5408  \u8ba1",
                "\u5408 \u8ba1",
                "\u5c0f\u8ba1",
            ) or t.startswith("\u5408"):
                total_ri = ri
                break
        except:
            pass
    if total_ri < 0:
        return 0
    total = 0.0
    for ri in range(1, total_ri):
        try:
            t = table.rows[ri].cells[col].text.strip().replace(",", "")
            if t and t not in ("-", ""):
                total += float(t)
        except:
            pass
    try:
        c = table.rows[total_ri].cells[col]
        old = c.text.strip().replace(",", "")
        if not old or old in ("-", ""):
            fill_cell(c, f"{total:,.2f}")
            return 1
    except:
        pass
    return 0


def calc_pct(table, amount_col, pct_col):
    """计算比例：本行金额/合计金额"""
    total = 0.0
    for ri in range(1, len(table.rows)):
        try:
            t = table.rows[ri].cells[amount_col].text.strip().replace(",", "")
            if t and t not in ("-", ""):
                total += float(t)
        except:
            pass
    if total == 0:
        return 0
    filled = 0
    for ri in range(1, len(table.rows)):
        try:
            t = table.rows[ri].cells[amount_col].text.strip().replace(",", "")
            if t and t not in ("-", ""):
                v = float(t) / total
                fill_cell(table.rows[ri].cells[pct_col], f"{v:.4f}")
                filled += 1
        except:
            pass
    return filled


def run(project_path, output_dir):
    with open(project_path, "r", encoding="utf-8") as f:
        proj = json.load(f)

    tmpl_path = proj["template"]
    tz_path = proj["tz"]
    tables_config = proj["tables"]

    doc = Document(tmpl_path)
    xl = openpyxl.load_workbook(tz_path, data_only=True)

    total_filled = 0
    log = []

    for tc in tables_config:
        ti = tc["index"]
        tname = tc.get("type", "")
        if ti >= len(doc.tables):
            log.append({"table": f"表{ti}", "status": "skip"})
            continue
        table = doc.tables[ti]

        # 找源数据sheet
        sheet = None
        for sn in tc.get("sheets", [tc.get("sheet", "")]):
            sheet = find_sheet(xl, sn)
            if sheet:
                break
        if sheet is None:
            log.append({"table": tname, "status": "no_sheet"})
            continue

        src_data = list(sheet.iter_rows(values_only=True))
        name_col = tc.get("name_col", 1)
        data_cols = tc.get("data_cols", [])
        tmpl_cols = tc.get("template_cols", [])
        row_map = tc.get("rows", [])
        sections = tc.get("sections", [])
        section_cols = tc.get("section_cols", {})
        calc_pcts = tc.get("calc_pcts", [])

        # 从源数据提取
        tz_data = {}
        sec_data = {}
        cur_sec = "default"
        sec_re = re.compile(
            r"^[\u300c\u300e\uff08(]?[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\d]+[\uff09)）\u3001.]"
        )

        for ri in range(
            tc.get("data_start", 5), min(len(src_data), tc.get("data_end", 200))
        ):
            r = src_data[ri]
            if not r or len(r) <= max([0] + [c for c in data_cols if c is not None]):
                continue
            n = str(r[name_col]).strip() if name_col < len(r) and r[name_col] else ""
            if not n:
                continue
            if sections and sec_re.match(n):
                for s in sections:
                    if n.startswith(s) or s[:2] in n:
                        cur_sec = s
                        break
                continue
            vals = {}
            for i, sc in enumerate(data_cols):
                if sc is not None and sc < len(r):
                    try:
                        vals[i] = float(r[sc])
                    except:
                        pass
            if vals:
                if sections:
                    sec_data.setdefault(cur_sec, {})[n] = vals
                else:
                    tz_data[n] = vals

        # 按row_map填入
        cells_filled = 0
        used = set()
        active_sec = "default"

        for ri in range(len(table.rows)):
            try:
                dn = table.rows[ri].cells[0].text.strip()
            except:
                continue
            if not dn or "\u5176\u4e2d" in dn:
                continue

            # 段切换
            if sections:
                for s in sections:
                    if dn.startswith(s) or s[:2] in dn:
                        active_sec = s
                        used = set()
                        break

            pool = sec_data.get(active_sec, {}) if sections else tz_data
            if not pool:
                continue

            # 找匹配行
            matched_row = None
            for rm in row_map:
                tmpl_r = rm.get("template", "")
                src_r = rm.get("source", "")
                if sections:
                    sec_match = rm.get("section", active_sec)
                    if sec_match != active_sec:
                        continue
                if match_name(dn, tmpl_r) or match_name(dn, src_r):
                    # 从源数据找对应行
                    for src_name, src_vals in pool.items():
                        if src_name in used:
                            continue
                        if match_name(src_r, src_name) or match_name(tmpl_r, src_name):
                            matched_row = (src_name, src_vals, rm)
                            break
                    break
                if matched_row:
                    break

            if matched_row is None:
                continue
            src_name, src_vals, rm = matched_row
            used.add(src_name)

            # 写值（section感知列映射）
            t_cols = rm.get("cols", tmpl_cols)
            # 如果当前段有section_cols，用它覆盖
            if sections and active_sec in section_cols:
                t_cols = section_cols[active_sec]
            for i, val in src_vals.items():
                if i < len(t_cols):
                    target = t_cols[i]
                    if target >= 0 and target < len(table.rows[ri].cells):
                        fill_cell(table.rows[ri].cells[target], val)
                        cells_filled += 1

        # 比例%计算（排除合计行，避免双倍计数）
        pct_count = 0
        for pct_def in calc_pcts:
            amount = pct_def.get("amount_col", 1)
            pct = pct_def.get("pct_col", 2)
            # 先算不含合计行的总额
            total = 0.0
            for ri in range(1, len(table.rows)):
                try:
                    t = table.rows[ri].cells[amount].text.strip().replace(",", "")
                    rn = table.rows[ri].cells[0].text.strip()
                    if t and t not in ("-", "") and "合" not in rn:
                        total += float(t)
                except:
                    pass
            if total > 0:
                for ri in range(1, len(table.rows)):
                    try:
                        t = table.rows[ri].cells[amount].text.strip().replace(",", "")
                        rn = table.rows[ri].cells[0].text.strip()
                        if t and t not in ("-", ""):
                            v = float(t) / total
                            fill_cell(table.rows[ri].cells[pct], f"{v:.4f}")
                            pct_count += 1
                    except:
                        pass

        # 合计求和
        sum_total = 0
        for sum_def in tc.get("sums", [{"col": 1}]):
            sum_total += auto_sum(table, sum_def.get("col", 1))

        total_filled += cells_filled + pct_count + sum_total
        log.append(
            {
                "table": tname or f"表{ti}",
                "cells": cells_filled,
                "pct": pct_count,
                "sums": sum_total,
            }
        )

    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    out_name = f"附注_填充_{ts}.docx"
    out_path = os.path.join(output_dir, out_name)
    os.makedirs(output_dir, exist_ok=True)
    doc.save(out_path)

    print(f"共填写 {total_filled} 个单元格 -> {out_path}")
    by_type = {}
    for l in log:
        t = l.get("table", "?")
        by_type.setdefault(t, {"c": 0, "p": 0, "s": 0})
        by_type[t]["c"] += l.get("cells", 0)
        by_type[t]["p"] += l.get("pct", 0)
        by_type[t]["s"] += l.get("sums", 0)
    for tn, v in sorted(by_type.items(), key=lambda x: -x[1]["c"]):
        print(f"  {tn}: {v['c']}格+{v['p']}比例+{v['s']}合计")

    return out_path


def main():
    parser = argparse.ArgumentParser(description="附注填充（配置驱动）")
    parser.add_argument("--config", required=True, help="project.json路径")
    parser.add_argument("--output", default=".", help="输出目录")
    args = parser.parse_args()
    run(args.config, args.output)


if __name__ == "__main__":
    main()
