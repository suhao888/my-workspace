"""
模板批量修复器
- 事务所名称: 众环→上会
- 报告年度: 20XX→2025
- 可配置，支持后续批量修改
"""
import sys, re, shutil
from pathlib import Path
from docx import Document
sys.stdout.reconfigure(encoding='utf-8')

DOCX_DIR = Path('D:/Users/12844/Desktop/已下发的文档模版/2.审计文档_docx')
BACKUP_DIR = DOCX_DIR.parent / '2.审计文档_docx_backup'

# 替换规则（可配置）
REPLACEMENTS = [
    # 事务所全称（先替换长的）
    ('中审众环会计师事务所（特殊普通合伙）', '上会会计师事务所（特殊普通合伙）'),
    ('中审众环会计师事务所', '上会会计师事务所'),
    # 事务所简称+报告类型
    ('众环专字', '上会专字'),
    ('众环审字', '上会审字'),
    ('众环综字', '上会综字'),
    # 报告年度（注意中文括号）
    ('20XX', '2025'),
    ('２０ＸＸ', '２０２５'),  # 全角
]

def backup():
    if not BACKUP_DIR.exists():
        shutil.copytree(DOCX_DIR, BACKUP_DIR)
        print(f'备份: {BACKUP_DIR}')
    else:
        print(f'备份已存在: {BACKUP_DIR}')

def fix_all():
    stats = {'files': 0, 'paras': 0, 'tables': 0}

    for f in sorted(DOCX_DIR.rglob('*.docx')):
        if f.name.startswith('~$'): continue

        doc = Document(str(f))
        changed = False

        # Fix paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                for old, new in REPLACEMENTS:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        stats['paras'] += 1
                        changed = True

        # Fix tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for old, new in REPLACEMENTS:
                                if old in run.text:
                                    run.text = run.text.replace(old, new)
                                    stats['tables'] += 1
                                    changed = True

        if changed:
            doc.save(str(f))
            rel = f.relative_to(DOCX_DIR)
            stats['files'] += 1
            print(f'  [修复] {rel}')

    print(f'\n修复: {stats["files"]} 文件, {stats["paras"]} 段落, {stats["tables"]} 表格单元格')


def verify():
    """验证修复结果"""
    remaining = []
    for f in sorted(DOCX_DIR.rglob('*.docx')):
        if f.name.startswith('~$'): continue
        doc = Document(str(f))
        for para in doc.paragraphs:
            if '众环' in para.text:
                remaining.append(f'{f.relative_to(DOCX_DIR)}: {para.text.strip()[:80]}')
            if '20XX' in para.text:
                remaining.append(f'{f.relative_to(DOCX_DIR)}: {para.text.strip()[:80]}')

    if remaining:
        print(f'\n剩余: {len(remaining)} 处未替换')
        for r in remaining[:10]:
            print(f'  {r}')
    else:
        print('\n全部替换完成!')


if __name__ == '__main__':
    print('=' * 60)
    print('模板修复: 众环→上会 | 20XX→2025')
    print('=' * 60)
    backup()
    fix_all()
    verify()
