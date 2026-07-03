"""
模板结构分析 — 从第7项开始逐科目解读
"""

import sys

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document

doc = Document(
    "D:/Users/12844/Desktop/自动化填模板/已下发的文档模版/2.审计文档_docx/1.审计报告/1-2单体审计报告模板/4.2025年财务报表附注模板-单户.docx"
)

print(f"模板总表格数: {len(doc.tables)}")
print()

# 找"七、财务报表重要项目的说明"位置
section7_start = None
for ti in range(len(doc.tables)):
    t = doc.tables[ti]
    heading = ""
    for ri in range(min(3, len(t.rows))):
        for ci in range(min(3, len(t.rows[ri].cells))):
            txt = t.rows[ri].cells[ci].text.strip()
            if len(txt) > 3:
                heading = txt[:60]
                break
        if heading:
            break
    if "七" in heading and ("重要" in heading or "报表" in heading):
        section7_start = ti
        print(f'>>> 找到"七、财务报表重要项目的说明": Table #{ti}')
        print(f"    内容: {heading}")
        break

if section7_start is None:
    # 可能前面的表格是封面/目录，直接从正文开始找
    # 检查有没有 "一、公司基本情况" 之类的
    for ti in range(len(doc.tables)):
        t = doc.tables[ti]
        heading = ""
        for ri in range(min(3, len(t.rows))):
            for ci in range(min(3, len(t.rows[ri].cells))):
                txt = t.rows[ri].cells[ci].text.strip()
                if len(txt) > 3:
                    heading = txt[:60]
                    break
            if heading:
                break
        if "公司基本情况" in heading and "一" in heading:
            section7_start = ti
            print(f'>>> 找到"一、公司基本情况": Table #{ti}')
            break
        if "重要" in heading and ("会计" in heading or "政策" in heading):
            section7_start = ti
            print(f'>>> 找到"重要会计政策": Table #{ti}')
            break

if section7_start is None:
    # 从头开始列出表格
    print("未找到章节标题，列出所有表格:")
    section7_start = 0

# 列出从section7_start开始的表格
print(f"\n=== 从Table #{section7_start}开始的表格结构 ===")
for ti in range(section7_start, min(section7_start + 100, len(doc.tables))):
    t = doc.tables[ti]
    # 提取前两行文本
    texts = []
    for ri in range(min(3, len(t.rows))):
        row_texts = []
        for ci in range(min(3, len(t.rows[ri].cells))):
            txt = t.rows[ri].cells[ci].text.strip()[:40]
            if txt:
                row_texts.append(txt)
        if row_texts:
            texts.append(" | ".join(row_texts))
    heading = " | ".join(texts)
    print(f"T{ti:3d} ({len(t.rows):2d}行×{len(t.columns):2d}列): {heading[:100]}")
