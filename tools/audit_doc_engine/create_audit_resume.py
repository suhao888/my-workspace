from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(10.5)

    # 标题
    title = doc.add_heading("审计师简历模板", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 个人信息表格
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 填充个人信息
    personal_info = [
        ["姓名：", "张三", "电话：", "138-0013-8000"],
        ["邮箱：", "zhangsan@email.com", "地址：", "上海市浦东新区"],
        ["出生年月：", "1990年1月", "政治面貌：", "中共党员"],
        ["求职意向：", "审计师", "期望薪资：", "面议"],
    ]

    for i, row_data in enumerate(personal_info):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            # 设置第一列和第三列为粗体
            if j % 2 == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()  # 空行

    # 教育背景
    add_heading(doc, "教育背景", level=1)
    edu_info = [
        ["2015.09 - 2018.06", "上海财经大学", "会计学硕士", "GPA: 3.8/4.0"],
        ["2011.09 - 2015.06", "上海财经大学", "会计学学士", "GPA: 3.6/4.0"],
    ]

    edu_table = doc.add_table(rows=len(edu_info), cols=4)
    edu_table.style = "Table Grid"
    for i, row_data in enumerate(edu_info):
        row = edu_table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text

    doc.add_paragraph()

    # 工作经历
    add_heading(doc, "工作经历", level=1)

    # 工作经历1
    add_paragraph(doc, "2020.07 - 2024.06  普华永道（中国）  高级审计师", bold=True)
    duties1 = [
        "主导并执行3个大型IPO上市审计项目，包括前期尽职调查、财务报表审阅及上市申报材料复核",
        "负责5家上市公司及2家大型民营企业的年度财务报表审计工作，严格遵循四大审计底稿编制规范",
        "深入分析客户内控体系，识别关键风险点并提出改进建议，帮助客户优化内控流程",
        "指导并培训初级审计师，负责项目团队管理与协调，有效提升团队工作效率15%",
    ]
    for duty in duties1:
        doc.add_paragraph(duty, style="List Bullet")

    doc.add_paragraph()

    # 工作经历2
    add_paragraph(doc, "2018.07 - 2020.06  德勤华永会计师事务所  审计师", bold=True)
    duties2 = [
        "参与8家跨国企业和国内知名企业的年度审计项目，负责底稿编制、往来款项函证、存货监盘等基础审计工作",
        "协助高级审计师进行风险评估程序和内控测试，撰写相关工作底稿和内部控制缺陷报告",
        "参与完成2个特殊目的审计项目，包括并购尽职调查及专项审计，撰写相关审计报告",
    ]
    for duty in duties2:
        doc.add_paragraph(duty, style="List Bullet")

    doc.add_paragraph()

    # 专业技能
    add_heading(doc, "专业技能", level=1)
    skills = [
        "审计专业技能：IPO上市审计、四大审计规范、财务报表审计、内部控制审计、风险评估",
        "会计准则与政策：中国会计准则（CAS）、国际财务报告准则（IFRS）、美国通用会计准则（US GAAP）",
        "审计工具：Auditor Assistant、ACL、Excel (高级)、SAP、用友/金蝶",
        "语言能力：英语（商务流利）、普通话（母语）",
    ]
    for skill in skills:
        doc.add_paragraph(skill, style="List Bullet")

    doc.add_paragraph()

    # 证书资质
    add_heading(doc, "证书资质", level=1)
    certs = [
        ["中国注册会计师（CPA）", "中国注册会计师协会", "2019.08"],
        ["ACCA（特许公认会计师）", "特许公认会计师公会", "2017.06"],
    ]

    cert_table = doc.add_table(rows=len(certs), cols=3)
    cert_table.style = "Table Grid"
    for i, row_data in enumerate(certs):
        row = cert_table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text

    doc.add_paragraph()

    # 获奖经历
    add_heading(doc, "获奖经历", level=1)
    awards = [
        ["优秀员工奖", "普华永道（中国）", "2023.12"],
        ["专业技能竞赛一等奖", "上海财经大学", "2017.05"],
    ]

    award_table = doc.add_table(rows=len(awards), cols=3)
    award_table.style = "Table Grid"
    for i, row_data in enumerate(awards):
        row = award_table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text

    # 保存文档
    output_path = r"D:\Users\12844\Desktop\审计师简历模板.docx"
    doc.save(output_path)
    print(f"模板已保存到: {output_path}")


if __name__ == "__main__":
    main()
