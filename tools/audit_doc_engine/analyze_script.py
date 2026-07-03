# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding="utf-8")
import docx
from docx.shared import RGBColor
from docx.oxml.ns import qn

fp = "D:/Users/12844/Desktop/已下发的文档模版/2.审计文档_docx/1.审计报告/1-2单体审计报告模板/4.2025年财务报表附注模板-单户.docx"
op = "D:/Users/12844/Desktop/analysis_output.txt"
doc = docx.Document(fp)
tables = doc.tables
lines = []
def o(s=""): lines.append(s)

def gct(cell):
    ts = []
    for p in cell.paragraphs:
        t = p.text.strip()
        if t: ts.append(t)
    return " | ".join(ts) if ts else ""

def grt(row):
    return [gct(c) for c in row.cells]

def pt(t, i, n=5):
    rr = len(t.rows)
    cc = len(t.columns) if t.columns else (len(t.rows[0].cells) if t.rows else 0)
    o(f"--- Table {i} (rows={rr}, cols={cc}) ---")
    for ri in range(min(n, rr)):
        rd = grt(t.rows[ri])
        o("  R" + str(ri) + ": " + " | ".join(rd))
    if rr > n:
        o(f"  ... ({rr - n} more rows)")

def classify(table):
    tc = 0; ec = 0; hp = False
    for row in table.rows:
        for cell in row.cells:
            text = gct(cell)
            tc += 1
            if not text.strip(): ec += 1
            if ("(" in text and ")" in text and len(text) < 25):
                hp = True
            if (chr(65288) in text and chr(65289) in text and len(text) < 25):
                hp = True
    er = ec / max(tc, 1) if tc > 0 else 0
    if hp: return "DATA(ph)"
    if er > 0.3 and len(table.rows) > 2: return "DATA(emp)"
    if len(table.rows) > 5: return "DATA(row)"
    return "FIXED"

# PART 0: All 223 tables
o("=" * 120)
o("PART 0: ALL 223 TABLES OVERVIEW")
o("=" * 120)
for i, table in enumerate(tables):
    rr = len(table.rows)
    cc = len(table.columns) if table.columns else (len(table.rows[0].cells) if table.rows else 0)
    fr = grt(table.rows[0]) if table.rows else []
    o("T" + str(i).rjust(3) + " (r=" + str(rr).rjust(2) + ", c=" + str(cc).rjust(2) + "): " + " | ".join(fr))

# PART 1: Every 10th table
o("")
o("=" * 120)
o("PART 1: EVERY 10TH TABLE")
o("=" * 120)
for i, table in enumerate(tables):
    if i % 10 == 0:
        pt(table, i, 4)

# PART 2: Classification
o("")
o("=" * 120)
o("PART 2: TABLE CLASSIFICATION")
o("=" * 120)
di = []; fi = []
for i, table in enumerate(tables):
    cls = classify(table)
    if cls.startswith("DATA"): di.append(i)
    else: fi.append(i)
    fr = grt(table.rows[0]) if table.rows else []
    o("T" + str(i).rjust(3) + " [" + cls.ljust(10) + "] " + " | ".join(fr)[:100])
o("Data tables: " + str(len(di)) + ", Fixed tables: " + str(len(fi)))
o("Fixed indices: " + str(fi))
o("")
o("--- Fixed Tables Full Content ---")
for idx in fi:
    pt(tables[idx], idx, 100)

# PART 3: Keyword search
o("")
o("=" * 120)
o("PART 3: KEYWORD SEARCH")
o("=" * 120)

kw_pairs = [
    ("huobi_zijin", ["库存现金", "货币资金"]),
    ("yingshou_zhangkuan", ["应收账款"]),
    ("yufu_kuanxiang", ["预付款项"]),
    ("qita_yingshou", ["其他应收款"]),
    ("guding_zichan", ["固定资产", "账面原值"]),
    ("wuxing_zichan", ["无形资产"]),
    ("changqi_daitan", ["长期待摊费用"]),
    ("diyan_suodeshui", ["递延所得税"]),
    ("yingfu_zhangkuan", ["应付账款"]),
    ("yingfu_xinchou", ["应付职工薪酬"]),
    ("yingjiao_shuifei", ["应交税费"]),
    ("qita_yingfu", ["其他应付款"]),
    ("shishou_ziben", ["实收资本"]),
    ("ziben_gongji", ["资本公积"]),
    ("yingyu_gongji", ["盈余公积"]),
    ("weifenpei_lirun", ["未分配利润"]),
    ("yingye_shouru", ["营业收入"]),
    ("shuijin_fujia", ["税金及附加"]),
    ("guanli_feiyong", ["管理费用"]),
    ("caiwu_feiyong", ["财务费用"]),
    ("xinyong_jianzhi", ["信用减值损失"]),
    ("zichan_jianzhi", ["资产减值损失"]),
    ("yingwai_shouru", ["营业外收入"]),
    ("yingwai_zhichu", ["营业外支出"]),
    ("suodeshui_feiyong", ["所得税费用"]),
    ("xianjin_liuliang", ["现金流量", "间接法", "净利润"]),
    ("shouxian_zichan", ["所有权受限", "受限资产"]),
    ("zhengfu_buzhu", ["政府补助"]),
]

matched = {}
for i, table in enumerate(tables):
    all_text = ""
    for row in table.rows:
        for cell in row.cells:
            all_text += gct(cell) + " "
    for cat, kwlist in kw_pairs:
        if cat not in matched: matched[cat] = []
        for kw in kwlist:
            if kw in all_text:
                matched[cat].append((i, kw))
                break

for cat, matches in matched.items():
    if matches:
        o("### " + cat + " ###")
        seen = set()
        for idx, kw in matches:
            if idx not in seen:
                seen.add(idx)
                o("  Table " + str(idx) + " (kw: " + kw + ")")
                pt(tables[idx], idx, 6)
    else:
        o("### " + cat + " ### (NO MATCH)")

# PART 4: Tables with percentage-style headers
o("")
o("=" * 120)
o("PART 4: TABLES WITH bili/pct/yuqixinyong")
o("=" * 120)
for i, table in enumerate(tables):
    fh = set()
    for row in table.rows:
        for cell in row.cells:
            text = gct(cell)
            if "比例" in text: fh.add("比例")
            if "百分比" in text: fh.add("百分比")
            if "预期信用损失率" in text: fh.add("预期信用损失率")
    if fh:
        o("Table " + str(i) + " (" + ", ".join(fh) + ")")
        pt(table, i, 5)

# PART 5: Red font and italic
o("")
o("=" * 120)
o("PART 5: RED FONT & ITALIC")
o("=" * 120)
red_info = []
ital_info = []
for i, table in enumerate(tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                for run in p.runs:
                    txt = run.text.strip()
                    if not txt: continue
                    is_red = False
                    if run.font.color and run.font.color.rgb:
                        if run.font.color.rgb == RGBColor(0xFF, 0x00, 0x00): is_red = True
                    if not is_red:
                        rPr = run._element.find(qn("w:rPr"))
                        if rPr is not None:
                            clr = rPr.find(qn("w:color"))
                            if clr is not None:
                                v = clr.get(qn("w:val"))
                                if v and v.upper() in ("FF0000", "RED"): is_red = True
                    if is_red:
                        red_info.append((i, ri, ci, txt[:60]))
                    is_ital = False
                    if run.font.italic: is_ital = True
                    if not is_ital:
                        rPr = run._element.find(qn("w:rPr"))
                        if rPr is not None:
                            it = rPr.find(qn("w:i"))
                            if it is not None:
                                v = it.get(qn("w:val"))
                                if v is None or v in ("1", "true"): is_ital = True
                    if is_ital:
                        ital_info.append((i, ri, ci, txt[:60]))

o("Red cells: " + str(len(red_info)))
for idx, ri, ci, txt in red_info:
    o("  T" + str(idx) + " R" + str(ri) + " C" + str(ci) + ": [" + txt + "]")
o("Italic cells: " + str(len(ital_info)))
for idx, ri, ci, txt in ital_info:
    o("  T" + str(idx) + " R" + str(ri) + " C" + str(ci) + ": [" + txt + "]")

# Write final output
with open(op, "w", encoding="utf-8") as f:
    f.write(chr(10).join(lines))
print("DONE: " + op)
