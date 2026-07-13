"""
财务报表附注填充器 v10 — 会计语义引擎
=====================================
v10: 基于企业会计准则(CAS 22/CAS 30)的语义理解层
   - 每个附注表的会计含义、列语义、勾稽关系
   - 从"字符串匹配"升级为"会计概念匹配"
"""

import sys, re, pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

sys.stdout.reconfigure(encoding="utf-8")

# ============================================================
# 会计语义引擎 — 基于企业会计准则的知识层
# ============================================================

# ---- 附注表语义定义 ----
# 每个附注表的会计含义、数据来源、披露要求

SEMANTIC_TABLES = {
    # === 资产类 ===
    "货币资金-明细": {
        "standard": "CAS 30 应用指南 — 货币资金",
        "meaning": "按现金、银行存款、其他货币资金分类列示期末和期初余额",
        "source": "货币资金_原始数据 sheet",
        "key_concepts": ["库存现金", "银行存款", "其他货币资金"],
        "cross_check": ["现金和现金等价物-构成"],  # 勾稽关系
    },
    "应收票据-分类": {
        "standard": "CAS 22 — 金融工具：应收票据",
        "meaning": "按票据种类(银行承兑/商业承兑)分别列示账面余额、坏账准备、账面价值",
        "note": "银行承兑汇票信用风险极低，通常不计提坏账准备；商业承兑汇票按ECL模型计提",
        "source": "应收票据分类_原始数据 sheet",
        "columns": {
            "期末账面余额": "套表C3(期末数)",
            "期末坏账准备": "通常为0(银行承兑)或按ECL计提(商业承兑)",
            "期末账面价值": "= 期末账面余额 - 期末坏账准备",
            "期初账面余额": "套表C6(期初数)",
        },
        "cross_check": [],  # 应收票据独立，无需跨表勾稽
    },
    "应收账款-账龄": {
        "standard": "CAS 22 应用指南 — 应收账款预期信用损失",
        "meaning": "披露应收账款按账龄区间的账面余额和坏账准备分布，用于计算ECL",
        "source": "应收款项计提坏账准备情况表 sheet — 组合计提部分的账龄子行",
        "business_rules": [
            '单项计提的应收账款(已单独测试减值)归入"3年以上"账龄段 — CAS 22.47: 单项计提通常针对长期逾期款项',
            '关联方组合应收款通常无账龄明细(内部控制强，回收风险低)，需在"关联交易"附注中单独披露',
            "账龄合计数必须与应收账款组合计提表的合计数一致 — 勾稽关系",
        ],
        "columns": {
            "期末账面余额": "套表C3(期末金额)或C8(期初金额)，按账龄段汇总",
            "期末坏账准备": "套表C5或C10，ECL模型计提结果",
            "期初账面余额": "套表C8(期初金额)",
            "期初坏账准备": "套表C10(期初坏账准备)",
        },
        "cross_check": ["应收账款-组合(期末)", "应收账款-组合(期初)"],
    },
    "应收账款-组合(期末)": {
        "standard": "CAS 22.48 — 预期信用损失计量",
        "meaning": "按计提方式(单项/组合)分类，披露期末账面余额、坏账准备、计提比例、账面价值",
        "cross_check": ["应收账款-账龄"],  # 合计应等于账龄表合计
    },
    "预付款项-账龄": {
        "standard": "CAS 30 — 预付款项",
        "meaning": "按账龄列示预付款项金额和比例",
        "note": "预付款项不计提坏账准备(非金融资产)，仅按账龄分类",
    },
    "固定资产-情况表": {
        "standard": "CAS 4 — 固定资产",
        "meaning": "按类别列示固定资产原值、累计折旧、减值准备、账面价值的增减变动",
        "source": "固定资产情况_原始数据 sheet",
    },
    # === 负债类 ===
    "应付职工薪酬-总表": {
        "standard": "CAS 9 — 职工薪酬",
        "meaning": "按短期薪酬、离职后福利、辞退福利等大类列示期初、增加、减少、期末",
        "note": "总表仅列示大类合计，明细在短期薪酬明细表中展开",
    },
    "应付职工薪酬-短期薪酬明细": {
        "standard": "CAS 9 应用指南 — 短期薪酬",
        "meaning": "按工资、福利费、社保、公积金、工会经费等明细列示",
        "note": '模板中"医疗保险费及生育保险费"合并了基本+补充医疗保险',
    },
    # === 权益类 ===
    "未分配利润": {
        "standard": "CAS 30 — 未分配利润",
        "meaning": "列示年初余额、本年增加(净利润转入)、本年减少(提取盈余公积/分配股利)、年末余额",
        "note": "套表有纵向分段结构：增加额区域(净利润转入等)和减少额区域(提取盈余公积等)",
    },
    # === 利润表类 ===
    "营业收入/成本": {
        "standard": "CAS 14 — 收入",
        "meaning": "按业务类型(主营业务/其他业务)列示收入和成本的本期和上期金额",
    },
    "其他收益": {
        "standard": "CAS 16 — 政府补助",
        "meaning": "列示与日常活动相关的政府补助及其他收益",
        "note": "通用模板行名可能与实际企业不同，需根据企业实际发生的其他收益项目调整",
    },
    # === 现金流量表 ===
    "现金和现金等价物": {
        "standard": "CAS 31 — 现金流量表",
        "meaning": "列示现金和现金等价物的构成，期末数应与资产负债表货币资金勾稽",
        "cross_check": ["货币资金-明细"],
    },
}

# ---- 跨表勾稽关系 ----
# 基于会计准则的数值一致性约束

CROSS_VALIDATIONS = [
    # (表A, 表B, 关系, 说明)
    {
        "desc": "应收账款-账龄期末合计 = 应收账款-组合(期末)合计",
        "check": ("应收账款-账龄", "合计", "应收账款-组合(期末)", "合计"),
        "columns": (1, 1),  # 比较C1列
    },
    {
        "desc": "应收账款-账龄期末坏账准备 = 应收账款-组合(期末)坏账准备",
        "check": ("应收账款-账龄", "合计", "应收账款-组合(期末)", "合计"),
        "columns": (2, 3),  # 坏账准备列
    },
    {
        "desc": "现金和现金等价物期末 = 货币资金期末(仅现金部分)",
        "check": ("现金和现金等价物", "合计", "货币资金-明细", "合计"),
        "columns": (1, 1),
    },
]

# ---- 会计概念映射 ----
# 将套表中的概念映射到模板中的会计概念

CONCEPT_MAP = {
    # 套表概念 → 模板概念
    "单项计提": {
        "aging": "3年以上",  # 单项计提→归入3年以上账龄
        "reason": "CAS 22.47: 单项计提针对已发生信用减值的金融资产，通常已长期逾期",
    },
    "组合计提": {
        "aging": "按实际账龄分布",  # 组合计提→各账龄段
    },
    "关联方组合": {
        "disclosure": "关联交易附注",  # 关联方应收款在关联交易附注单独披露
        "aging": "通常无账龄细分(回收风险低)",  # 套表不提供账龄明细
    },
    "账面价值": {
        "formula": "账面余额 - 坏账准备",
        "standard": "CAS 22: 金融资产按摊余成本计量，扣除损失准备后列示",
    },
}

# ---- 业务规则引擎 ----
# 会计判断规则


def apply_accounting_rules(table_name, data):
    """应用会计准则的业务规则"""
    rules_applied = []

    if table_name == "应收账款-账龄":
        # 规则1: 单项计提归入3年以上
        if "一、按单项计提坏账准备的应收账款" in data:
            rules_applied.append(
                'CAS 22.47: 单项计提应收账款(已发生信用减值)归入"3年以上"账龄段'
            )
        # 规则2: 关联方组合无账龄明细
        if "（一）关联方组合" in data:
            rules_applied.append(
                "关联方组合37,124,436.53元在套表中无账龄明细，"
                '应单独在"关联方及关联交易"附注中披露(CAS 36)'
            )

    return rules_applied


# ============================================================
# 精确映射 — 基于会计语义的列映射
# ============================================================
MAPPINGS = [
    # ======== 资产类 ========
    {
        "cat": "货币资金-明细",
        "sheet_kw": "货币资金_原始数据",
        "table_idx": 5,
        "tz_cols": [3, 4],
        "col_map": [(0, 1, False), (1, 2, False)],  # 期末→C1, 期初→C2
    },
    {
        "cat": "应收票据-分类",
        "sheet_kw": "应收票据分类_原始数据",
        "table_idx": 13,
        # v9修复: C3=期末数→C1(期末账面余额), C6=期初数→C4(期初账面余额)
        "tz_cols": [3, 6],
        "col_map": [
            (0, 1, False),  # 期末数→C1(期末账面余额)
            (1, 4, False),  # 期初数→C4(期初账面余额)
        ],
        "data_start": 6,
        # v9: 合计行聚合(票据种类→合计)
        "aggregations": [
            {
                "target": "合计",
                "sources": ["银行承兑汇票", "商业承兑汇票"],
                "op": "sum",
            },
        ],
    },
    {
        "cat": "应收账款-组合(期末)",
        "sheet_kw": "应收款项计提坏账准备情况表_原始数据",
        "table_idx": 24,
        # 套表: C3=期末余额金额, C4=期末余额比例%, C5=坏账准备金额, C6=预期信用损失率, C7=账面净额
        "tz_cols": [3, 4, 5, 6, 7],
        "col_map": [
            (0, 1, False),  # 账面余额金额 → 模板C1(金额)
            (1, 2, True),  # 账面余额比例% → 模板C2(比例%)  ×100
            (2, 3, False),  # 坏账准备金额 → 模板C3(金额)
            (3, 4, True),  # 预期信用损失率 → 模板C4(计提比例%)  ×100
            (4, 5, False),  # 账面净额 → 模板C5(账面价值)
        ],
        "name_exclude": ["其他应收"],  # v9: 排除其他应收款section
    },
    {
        "cat": "应收账款-组合(期初)",
        "sheet_kw": "应收款项计提坏账准备情况表_原始数据",
        "table_idx": 25,
        # 套表: C8=期初金额, C9=期初比例%, C10=期初坏账准备, C11=期初损失率
        "tz_cols": [8, 9, 10, 11],
        "col_map": [
            (0, 1, False),  # 期初账面余额 → C1
            (1, 2, True),  # 期初比例% → C2 ×100
            (2, 3, False),  # 期初坏账准备 → C3
            (3, 4, True),  # 期初损失率 → C4 ×100
        ],
        "name_exclude": ["其他应收"],  # v9: 排除其他应收款section
    },
    {
        "cat": "应收账款-单项计提",
        "sheet_kw": "单项计提坏账准备的应收款项_原始数据",
        "table_idx": 26,
        "tz_cols": [4, 5, 7],  # 账面余额, 坏账准备, 预期信用损失率
        "col_map": [
            (0, 1, False),
            (1, 2, False),
            (2, 3, True),  # 损失率→C3(预期信用损失率%)
        ],
        "data_start": 5,
        "name_strip": ["年末", "年初"],
        "name_exclude": ["其他应收款", "其他应收"],  # v9: 排除其他应收款section
    },
    {
        "cat": "应收账款-组合计提",
        "sheet_kw": "应收款项计提坏账准备情况表_原始数据",
        "table_idx": 27,
        "tz_cols": [3, 7],  # 账面余额(期末), 账面净额(期末)
        "col_map": [
            (0, 1, False),  # 账面余额→C1
            (1, 5, False),  # 账面净额→C5
        ],
        "name_exclude": ["其他应收"],  # v9: 排除其他应收款section
    },
    {
        "cat": "应收账款-前五名",
        "sheet_kw": "按欠款方归集应收及预付款项期末余额前五名_原始数据",
        "table_idx": 30,
        "tz_cols": [4, 6, 7],  # 账面余额, 占比, 坏账准备
        "col_map": [
            (0, 1, False),
            (1, 2, True),  # 占比→C2(比例%) ×100
            (2, 3, False),
        ],
        "data_start": 5,
        "section": "应收账款",
    },
    {
        "cat": "应收账款-账龄",
        "sheet_kw": "应收款项计提坏账准备情况表_原始数据",
        "table_idx": 23,
        # 套表: C3=期末金额, C5=期末坏账准备, C7=期末净额
        #       C8=期初金额, C10=期初坏账准备
        "tz_cols": [3, 5, 8, 10],
        "accumulate": True,  # v9: 跨section同名账龄行累加
        "name_exclude": ["其他应收"],  # v10: 排除其他应收款section
        "col_map": [
            (0, 1, False),  # 期末金额→C1(期末账面余额)
            (1, 2, False),  # 期末坏账准备→C2(期末坏账准备)
            (2, 3, False),  # 期初金额→C3(期初账面余额)
            (3, 4, False),  # 期初坏账准备→C4(期初坏账准备)
        ],
        # v9: 跨section账龄聚合(关联方+非关联方)
        "aggregations": [
            {
                "target": "1年以内（含1年）",
                "sources": ["1年以内", "1年以内（含1年）", "其中：1年以内账龄"],
                "op": "sum",
            },
            {
                "target": "1-2年",
                "sources": ["1-2年", "1－2年（含2年）"],
                "op": "sum",
            },
            {
                "target": "2-3年",
                "sources": ["2-3年", "2－3年（含3年）"],
                "op": "sum",
            },
            {
                "target": "3年以上",
                "sources": [
                    "3年以上",
                    "3－4年（含4年）",
                    "4－5年（含5年）",
                    "5年以上",
                    "3-4年",
                    "4-5年",
                    "一、按单项计提坏账准备的应收账款",
                ],
                "op": "sum",
            },
        ],
    },
    # 应收账款-坏账准备收回/转回: 该司决算套表中无此数据(T28为收回转回表,不适用)
    # 预付款项
    {
        "cat": "预付款项-账龄",
        "sheet_kw": "预付款项按账龄列示_原始数据",
        "table_idx": 34,
        "tz_cols": [3, 4, 5, 6],  # 期末金额, 期末比例, 期初金额, 期初比例
        "col_map": [
            (0, 1, False),  # 期末金额→C1
            (1, 2, True),  # 期末比例→C2(%) ×100
            (2, 3, False),  # 期初金额→C3
            (3, 4, True),  # 期初比例→C4(%) ×100
        ],
    },
    {
        "cat": "预付款项-前五名",
        "sheet_kw": "按欠款方归集应收及预付款项期末余额前五名_原始数据",
        "table_idx": 36,
        "tz_cols": [4, 6],  # 账面余额, 占比
        "col_map": [
            (0, 1, False),
            (1, 2, True),
        ],
        "data_start": 5,
        "section": "预付款项",
    },
    # 其他应收款
    {
        "cat": "应收资金集中管理款",
        "sheet_kw": "应收资金集中管理款_原始数据",
        "table_idx": 39,
        "tz_cols": [3, 4],  # 期末, 期初
        "col_map": [(0, 1, False), (1, 2, False)],
        "data_start": 5,
    },
    {
        "cat": "其他应收款-分类",
        "sheet_kw": "其他应收款_原始数据",
        "table_idx": 40,
        "tz_cols": [3, 4],
        "col_map": [(0, 1, False), (1, 2, False)],
    },
    {
        "cat": "其他应收款-账龄",
        "sheet_kw": "按账龄披露其他应收款项_原始数据",
        "table_idx": 45,
        "tz_cols": [
            3,
            5,
            6,
            8,
        ],  # v11: 期末余额,期末坏账准备,期初余额,期初坏账准备(跳过比例列)
        "col_map": [
            (0, 1, False),  # 期末账面余额→C1
            (1, 2, False),  # 期末坏账准备→C2(金额,非比例)
            (2, 3, False),  # 期初账面余额→C3
            (3, 4, False),  # 期初坏账准备→C4(金额,非比例)
        ],
    },
    {
        "cat": "其他应收款-坏账准备变动",
        "sheet_kw": "本期计提、收回或转回的其他应收款坏账准备_原始数据",
        "table_idx": 43,
        "tz_cols": [3, 6],
        "col_map": [(0, 1, False), (1, 4, False)],  # 期初→C1, 期末→C4
    },
    {
        "cat": "其他应收款-前五名",
        "sheet_kw": "按欠款方归集应收及预付款项期末余额前五名_原始数据",
        "table_idx": 53,
        "tz_cols": [4, 6, 7],
        "col_map": [
            (0, 1, False),  # 账面余额→C1(债务人名称后)
            (1, 4, True),  # 占比→C4(%) ×100
            (2, 5, False),  # 坏账准备→C5
        ],
        "data_start": 5,
        "section": "其他应收款",
    },
    # 合同资产
    {
        "cat": "合同资产-情况",
        "sheet_kw": "合同资产情况_原始数据",
        "table_idx": 60,
        "tz_cols": [3, 4, 5, 7],
        "col_map": [
            (0, 1, False),
            (2, 2, True),
            (1, 3, False),
            (2, 4, True),
            (3, 5, False),
        ],
    },
    # 固定资产
    {
        "cat": "固定资产-账面价值",
        "sheet_kw": "固定资产_原始数据",
        "table_idx": 90,
        "tz_cols": [3, 4],
        "col_map": [(0, 1, False), (1, 2, False)],
    },
    {
        "cat": "固定资产-情况表",
        "sheet_kw": "固定资产情况_原始数据",
        "table_idx": 91,
        "tz_cols": [3, 4, 5, 6],
        "col_map": [
            (0, 1, False),  # 期初余额→C1 (C2合并单元格)
            (1, 3, False),  # v11: 本年增加→C3 (跳过合并的C2)
            (2, 4, False),  # v11: 本年减少→C4 (C5合并单元格)
            (3, 6, False),  # v11: 期末余额→C6 (跳过合并的C5)
        ],
        # v12: section headers — 固定资产情况表含原值/折旧/减值/账面价值4个section
        #      每个section下资产名称重复(如"房屋及建筑物"),用section tag隔离
        "section_headers": [
            "一、账面原值合计",
            "二、累计折旧合计",
            "三、固定资产减值准备合计",
            "四、固定资产账面价值合计",
        ],
    },
    # 在建工程
    {
        "cat": "在建工程-情况",
        "sheet_kw": "在建工程情况_原始数据",
        "table_idx": 97,
        "tz_cols": [3, 4, 5, 6, 7, 8],
        "col_map": [(0, 2, False), (1, 3, False), (4, 4, False), (5, 6, False)],
    },
    # 使用权资产
    {
        "cat": "使用权资产-情况",
        "sheet_kw": "使用权资产情况_原始数据",
        "table_idx": 99,
        "tz_cols": [3, 4, 5, 6],
        "col_map": [(0, 1, False), (1, 2, False), (2, 3, False), (3, 4, False)],
        # v12: section headers — 5节:原值/折旧/净值/减值/账面价值
        "section_headers": [
            "一、账面原值合计",
            "二、累计折旧合计",
            "三、账面净值合计",
            "四、减值准备合计",
            "五、账面价值合计",
        ],
    },
    # 无形资产
    {
        "cat": "无形资产-增减",
        "sheet_kw": "无形资产_原始数据",
        "table_idx": 100,
        "tz_cols": [3, 4, 6],
        "col_map": [(0, 1, False), (1, 2, False), (2, 4, False)],
        # v12: section headers — 无形资产含原价/摊销/减值/账面价值section
        "section_headers": [
            "一、原价合计",
            "二、累计摊销额合计",
            "三、减值准备合计",
            "四、账面价值合计",
        ],
    },
    # 长期待摊费用
    {
        "cat": "长期待摊费用",
        "sheet_kw": "长期待摊费用_原始数据",
        "table_idx": 107,
        "tz_cols": [3, 4, 5],
        "col_map": [(0, 1, False), (1, 2, False), (2, 3, False)],
    },
    # 递延所得税
    {
        "cat": "递延所得税-资产/负债",
        "sheet_kw": "递延所得税资产和递延所得税负债_原始数据",
        "table_idx": 109,
        # 套表列: C3=递延所得税(期末), C4=互抵, C5=抵销后, C6=差异(期末)
        #          C7=递延所得税(期初), C8=互抵, C9=抵销后, C10=差异(期初)
        "tz_cols": [6, 3, 10, 7],  # v11-fix: 差异在前,递延所得税在后
        "col_map": [
            (0, 1, False),  # C6期末可抵扣差异→C1
            (1, 2, False),  # C3期末递延所得税资产→C2(差异×25%)
            (2, 3, False),  # C10期初可抵扣差异→C3
            (3, 4, False),  # C7期初递延所得税资产→C4(差异×25%)
        ],
    },
    # ======== 负债类 ========
    {
        "cat": "应付账款-账龄",
        "sheet_kw": "应付账款_原始数据",
        "table_idx": 120,
        "tz_cols": [3, 4],
        "col_map": [(0, 1, False), (1, 2, False)],
    },
    {
        "cat": "应付职工薪酬-总表",
        "sheet_kw": "应付职工薪酬_原始数据",
        "table_idx": 128,
        "tz_cols": [3, 4, 5, 6],
        "col_map": [(0, 1, False), (1, 2, False), (2, 3, False), (3, 4, False)],
    },
    {
        "cat": "应付职工薪酬-短期薪酬明细",
        "sheet_kw": "应付职工薪酬_原始数据",
        "table_idx": 129,
        "tz_cols": [3, 4, 5, 6],
        "col_map": [(0, 1, False), (1, 2, False), (2, 3, False), (3, 4, False)],
        "name_exclude": [
            "离职后福利",
            "设定提存计划",
            "企业年金",
            "失业保险",
            "辞退福利",
            "一年内到期",
            "一、短期薪酬",
        ],  # v9: 仅短期薪酬; v11: 排除总计数行防泄漏
        # v9: 跨行聚合 — 套表多行→模板单行
        "aggregations": [
            {
                "target": "其中：医疗保险费及生育保险费",
                "sources": ["基本医疗保险", "补充医疗保险", "生育保险费"],
                "op": "sum",
            },
            {
                "target": "五、工会经费和职工教育经费",
                "sources": ["（五）工会经费", "（六）职工教育经费"],
                "op": "sum",
            },
        ],
    },
    {
        "cat": "应交税费",
        "sheet_kw": "应交税费_原始数据",
        "table_idx": 131,
        "tz_cols": [3, 6],  # 期初, 期末
        "col_map": [(1, 1, False), (0, 2, False)],  # 期末→C1, 期初→C2
    },
    {
        "cat": "其他应付款-分类",
        "sheet_kw": "其他应付款_原始数据",
        "table_idx": 132,
        "tz_cols": [3, 4],
        "col_map": [(0, 1, False), (1, 2, False)],
    },
    # ======== 权益类 ========
    {
        "cat": "实收资本",
        "sheet_kw": "实收资本_原始数据",
        "table_idx": 159,
        "tz_cols": [3, 7],
        "col_map": [(0, 1, False), (1, 5, False)],  # 期初→C1, 期末→C5
    },
    {
        "cat": "盈余公积",
        "sheet_kw": "盈余公积_原始数据",
        "table_idx": 163,
        "tz_cols": [3, 4, 5, 6],
        "col_map": [(0, 1, False), (1, 2, False), (2, 3, False), (3, 4, False)],
    },
    {
        "cat": "未分配利润",
        "sheet_kw": "未分配利润_原始数据",
        "table_idx": 165,
        "tz_cols": [3, 4],
        "col_map": [(0, 1, False), (1, 2, False)],
    },
    # ======== 利润表类 ========
    {
        "cat": "营业收入/成本",
        "sheet_kw": "营业收入、营业成本_原始数据",
        "table_idx": 166,
        "tz_cols": [3, 4, 5, 6],
        "col_map": [(0, 1, False), (1, 2, False), (2, 3, False), (3, 4, False)],
        # v11: 排除非监管业务行—主营业务小计的分拆项,重复数据
        "name_exclude": ["非监管业务"],
    },
    {
        "cat": "其他收益",
        "sheet_kw": "其他收益_原始数据",
        "table_idx": 179,
        "tz_cols": [3, 4],
        "col_map": [(0, 1, False), (1, 2, False)],
    },
    {
        "cat": "投资收益",
        "sheet_kw": "投资收益_原始数据",
        "table_idx": 180,
        "tz_cols": [3, 4],
        "col_map": [(0, 1, False), (1, 2, False)],
    },
    {
        "cat": "信用减值损失",
        "sheet_kw": "信用减值损失_原始数据",
        "table_idx": 183,
        "tz_cols": [3, 4],
        "col_map": [(0, 1, False), (1, 2, False)],
    },
    {
        "cat": "资产减值损失",
        "sheet_kw": "资产减值准备情况表_原始数据",
        "table_idx": 184,
        "tz_cols": [3, 4],
        "col_map": [(0, 1, False), (1, 2, False)],
    },
    {
        "cat": "营业外收入",
        "sheet_kw": "营业外收入_原始数据",
        "table_idx": 186,
        "tz_cols": [3, 4, 5],
        "col_map": [(0, 1, False), (1, 2, False), (2, 3, False)],
    },
    {
        "cat": "营业外支出",
        "sheet_kw": "营业外支出_原始数据",
        "table_idx": 188,
        "tz_cols": [3, 4, 5],
        "col_map": [(0, 1, False), (1, 2, False), (2, 3, False)],
    },
    {
        "cat": "所得税费用",
        "sheet_kw": "所得税费用_原始数据",
        "table_idx": 189,
        "tz_cols": [3, 4],
        "col_map": [(0, 1, False), (1, 2, False)],
    },
    # ======== 现金流量表补充 ========
    {
        "cat": "现金流量表-间接法",
        "sheet_kw": "间接法现金流量表_原始数据",
        "table_idx": 205,
        "tz_cols": [3, 4],
        "col_map": [(0, 1, False), (1, 2, False)],
    },
    {
        "cat": "现金和现金等价物",
        "sheet_kw": "现金和现金等价物的构成_原始数据",
        "table_idx": 207,
        "tz_cols": [3, 4],
        "col_map": [(0, 1, False), (1, 2, False)],
    },
    # ======== 期间费用 ========
    # 管理/销售/研发费用: 成本费用管理情况表的"合   计"行取数
    # 通过NAME_SYNONYMS将费用名映射到"合   计"
    {
        "cat": "管理费用",
        "sheet_kw": "成本费用管理情况表_原始数据",
        "table_idx": 175,
        "tz_cols": [10, 24],  # 本年=col K, 上年=col X
        "col_map": [(0, 1, False), (1, 2, False)],
    },
    {
        "cat": "销售费用",
        "sheet_kw": "成本费用管理情况表_原始数据",
        "table_idx": 174,
        "tz_cols": [12, 26],  # 本年=col M, 上年=col Z
        "col_map": [(0, 1, False), (1, 2, False)],
        "keep_empty": True,  # 销售费用为0,合计行数据列为空但不跳过
    },
    {
        "cat": "研发费用",
        "sheet_kw": "成本费用管理情况表_原始数据",
        "table_idx": 176,
        "tz_cols": [13, 27],  # 本年=col N, 上年=col AA
        "col_map": [(0, 1, False), (1, 2, False)],
    },
    {
        "cat": "财务费用",
        "sheet_kw": "利润表_原始数据",  # 财务费用不在成本费用管理表中
        "table_idx": 177,
        "tz_cols": [3, 4],
        "col_map": [(0, 1, False), (1, 2, False)],
    },
]

# ============================================================
# 名称同义词
# ============================================================
NAME_SYNONYMS = {
    "合　　　计": ["合计", "合  计", "合   计"],
    "合      计": ["合计", "合  计"],
    "合  计": ["合计", "合   计"],
    "其中：人民币": ["其中：人民币", "人民币"],
    # 应收
    "应收账款合计": ["合计", "合   计"],
    "一、按单项计提坏账准备的应收账款": ["按单项计提坏账准备", "单项计提"],
    "二、按信用风险特征组合计提坏账准备的应收账款": ["按组合计提坏账准备", "组合计提"],
    "（一）关联方组合": ["关联方组合", "关联方"],
    "（三）应收其他产品销售（服务）款项组合": ["其他组合", "非关联方组合"],
    "按税法及相关规定计算的当期所得税": ["当期所得税费用", "当期所得税"],
    "递延所得税调整": ["递延所得税费用", "递延所得税"],
    # 应付职工薪酬
    "一、短期薪酬": ["短期薪酬"],  # v9: 移除宽泛同义词，避免明细表误匹配
    "（一）工资、奖金、津贴和补贴": ["工资、奖金、津贴和补贴"],
    "（二）职工福利费": ["二、职工福利费", "职工福利费"],
    "（三）社会保险费": ["三、社会保险费", "社会保险费"],
    "基本医疗保险": ["医疗保险费及生育保险费", "基本医疗"],
    "补充医疗保险": ["补充医疗保险", "补充医疗"],
    "二、离职后福利": ["离职后福利", "设定提存计划"],
    "（一）设定提存计划": ["设定提存计划", "基本养老保险"],
    "企业年金": ["企业年金", "年金缴费"],
    "失业保险费": ["失业保险费", "失业保险"],
    # 营业收入
    "1．主营业务小计": ["主营业务小计", "主营业务", "（1）主营业务小计"],
    "非监管业务": ["主营业务收入", "非监管"],
    # 固定资产
    "一、账面原值合计": ["一、账面原值合计", "固定资产原值"],
    "房屋及建筑物": ["房屋及建筑物", "房屋建筑物"],
    "二、累计折旧合计": ["二、累计折旧合计", "累计折旧"],
    "三、固定资产减值准备合计": ["固定资产减值准备", "减值准备"],
    "四、固定资产账面价值合计": ["固定资产账面价值", "账面价值"],
    # 无形资产
    "一、原价合计": ["一、原价合计", "无形资产原价"],
    "二、累计摊销额合计": ["二、累计摊销额合计", "累计摊销"],
    "四、账面价值合计": ["四、账面价值合计", "无形资产账面价值"],
    # 资产减值
    "一、坏账准备": ["坏账准备"],
    "其中：应收账款坏账准备": ["应收账款坏账准备"],
    # 营业外
    "非流动资产毁损报废损失": ["非流动资产毁损报废损失", "固定资产报废损失"],
    "行政性罚款、滞纳金": ["罚款支出", "滞纳金"],
    "无需支付款项": ["无法支付的应付款项", "无法支付", "其他"],
    "接受捐赠": ["接受捐赠", "捐赠收入"],
    # v9 新增同义词 — 减少漏填
    "本期计提": ["本年计提", "本期计提", "本期增加"],
    "机器运输办公设备": ["机器设备", "运输设备", "办公设备", "机器运输办公设备"],
    "现金的年末余额": ["现金的期末余额", "现金年末余额", "现金的年末余额"],
    "减：现金的年初余额": ["减：现金的期初余额", "现金的年初余额"],
    "现金等价物的年末余额": ["现金等价物的期末余额"],
    "其中：1年以内": ["1年以内（含1年）", "1年以内", "其中：1年以内"],
    "应收工程款": ["应收工程款", "2.应收工程款"],
    "补充资料": ["补充资料", "补充材料"],
    "企业员工": ["企业员工", "企业职工"],
    "进项税加计抵减": ["进项税加计抵减", "进项税额加计抵减", "加计抵减"],
    "储备基金": ["储备基金", "储备基金和企业发展基金"],
    "企业发展基金": ["企业发展基金", "企业发展基金和储备基金"],
    "其中：基建工程": ["基建工程", "其中：基建工程"],
    # 长期待摊
    "租入资产改良支出": [
        "经营租入固定资产改良",
        "租入固定资产改良支出",
        "固定资产改良支出",
    ],
    # 其他应付/应收
    "应付股利": ["应付股利", "应付利润"],
    "其他应付款项": ["其他应付款项", "其他应付款"],
    # v11: 期间费用 — 成本费用管理情况表的"合   计"行
    "管理费用": ["合   计", "合计"],
    "销售费用": ["合   计", "合计"],
    "研发费用": ["合   计", "合计"],
    "其他应收款项": ["其他应收款项", "其他应收款"],
    # v10: 资金集中管理
    "应收资金集中管理款": [
        "银行存款",
        "其中：人民币",
        "其中：活期",
        "应收资金集中管理款",
    ],
    "银行存款": ["应收资金集中管理款", "其中：人民币", "银行存款"],
    # 未分配利润
    "上年年末余额": ["上年年末余额", "年初未分配利润"],
    "本年期初余额": ["本年年初余额", "期初余额"],
    "本年增加额": ["本期增加额", "本年增加额", "净利润"],
    "本年减少额": ["本期减少额", "本年减少额", "利润分配"],
    "本年期末余额": ["期末未分配利润", "期末余额"],
    "其中：本年净利润转入": ["本年净利润转入", "净利润"],
    "其中：本年提取盈余公积数": ["本年提取盈余公积", "提取盈余公积"],
    "期初调整余额": ["期初调整金额", "年初调整金额"],
    # 信用减值
    "坏账损失": ["坏账损失", "应收账款坏账损失", "其他应收款坏账损失"],
    # 应收票据
    "按单项计提坏账准备": ["按单项计提坏账准备", "商业承兑汇票小计"],
    "按组合计提坏账准备": ["按组合计提坏账准备", "银行承兑汇票小计"],
    "单项计提坏账准备的应收账款": ["按单项计提坏账准备", "单项计提", "单位1"],
    "单项计提坏账准备的其他应收款": ["按单项计提坏账准备", "单项计提"],
    # 前五名section → 合计行
    "应收账款期末余额前五名的情况": ["合计", "合   计"],
    "其他应收款项期末余额前五名的情况": ["合计", "合   计"],
    "预付款项期末余额前五名的情况": ["合计", "合   计"],
    # 在建工程
    "在建工程": ["合计", "合   计"],
    "其中：基建工程": ["基建工程"],
    # 实收资本
    "合  计": ["合计", "合   计"],
}


def fmt_val(v, as_pct=False):
    """格式化数值 — v10: 财务数字始终保留两位小数, 消除浮点噪声"""
    if isinstance(v, (int, float)):
        if pd.isna(v):
            return ""
        # 消除浮点噪声: 极小值归零
        if abs(v) < 1e-10:
            v = 0.0
        if as_pct and abs(v) <= 1.0 and v != 0:
            v = v * 100
        return f"{round(v, 2):,.2f}"  # v10: 先round再format, 消除浮点噪声
    s = str(v).strip()
    if s in ("nan", "－", "-", "—", "", "N/A"):
        return ""
    try:
        n = float(s.replace(",", ""))
        if pd.isna(n):
            return ""
        if abs(n) < 1e-10:
            n = 0.0
        if as_pct and abs(n) <= 1.0 and n != 0:
            n = n * 100
        return f"{round(n, 2):,.2f}"  # v10: 先round再format
    except ValueError:
        return s


def match_name(doc_name, tz_name):
    """名称匹配"""
    d = re.sub(r"\s+", "", doc_name.replace("　", ""))
    t = re.sub(r"\s+", "", tz_name.replace("　", ""))
    if not d or not t:
        return False
    if d == t:
        return True

    d_clean = d.replace("（", "(").replace("）", ")").replace("△", "").replace("▲", "")
    t_clean = t.replace("（", "(").replace("）", ")").replace("△", "").replace("▲", "")
    if d_clean == t_clean:
        return True
    if len(d_clean) >= 4 and d_clean in t_clean:
        return True
    if len(t_clean) >= 4 and t_clean in d_clean:
        return True

    # 数字规范化
    d_num = d_clean.replace("至", "-").replace("（含", "(").replace("以上", "+")
    t_num = t_clean.replace("至", "-").replace("（含", "(").replace("以上", "+")
    if d_num == t_num:
        return True
    if len(d_num) >= 3 and d_num in t_num:
        return True
    if len(t_num) >= 3 and t_num in d_num:
        return True

    # v11: 归一化序号后比较（"（一）"→"一、"、"（二）"→"二、"），防止跨节误配
    def norm_seq(text):
        m = re.match(r"^[（(]?([一二三四五六七八九十\d]+)[）)、.]", text)
        if m:
            return m.group(1)
        return None

    d_seq = norm_seq(d_clean)
    t_seq = norm_seq(t_clean)
    if d_seq and t_seq and d_seq != t_seq:
        return False  # 不同序号（"一"vs"八"）不匹配

    # 剥离序号
    d_body = re.sub(r"^[（(]?[一二三四五六七八九十\d]+[）)、.]?\s*", "", d_clean)
    t_body = re.sub(r"^[（(]?[一二三四五六七八九十\d]+[）)、.]?\s*", "", t_clean)
    if d_body and t_body and len(d_body) >= 3:
        if d_body == t_body:
            return True
        # v11: 剥离序号后的子串匹配要求≥5字符,避免"短期薪酬"误配"其他短期薪酬"
        if len(d_body) >= 5 and d_body in t_body:
            return True
        if len(t_body) >= 5 and t_body in d_body:
            return True

    # 同义词
    for key, synonyms in NAME_SYNONYMS.items():
        for check_d, check_t in [(d_clean, key), (t_clean, key)]:
            if (
                check_d == check_t
                or (len(check_d) >= 3 and check_d in check_t)
                or (len(check_t) >= 3 and check_t in check_d)
            ):
                for syn in synonyms:
                    syn_c = (
                        syn.replace("（", "(")
                        .replace("）", ")")
                        .replace(" ", "")
                        .replace("　", "")
                    )
                    target = t_clean if check_d == d_clean else d_clean
                    if syn_c == target:
                        return True
                    if len(syn_c) >= 3 and syn_c in target:
                        return True
                    if len(target) >= 3 and target in syn_c:
                        return True

    # 公共子串 — 防止"合计"行匹配到非合计行
    # 注意: "组合计提"含"合计"二字, 需排除
    d_norm = d_clean.replace("组合计提", "XX")
    t_norm = t_clean.replace("组合计提", "XX")
    d_has_total = "合计" in d_norm
    t_has_total = "合计" in t_norm
    if d_has_total != t_has_total:
        return False

    if len(d_clean) >= 6 and len(t_clean) >= 6:
        for length in range(
            min(len(d_clean), 6), 5, -1
        ):  # v11: min=5, 防"固定资产"(4字)跨节误配
            for start in range(len(d_clean) - length + 1):
                if d_clean[start : start + length] in t_clean:
                    return True
    return False


def extract_tz_data(
    df,
    tz_cols,
    data_start=None,
    name_strip=None,
    section=None,
    name_exclude=None,
    accumulate=False,
    section_headers=None,
    keep_empty=False,
):
    """从决算套表提取数据 — v12: 支持section_headers分区标记"""
    # v12: section_headers — 列表, 如["一、账面原值合计","二、累计折旧合计",...]
    #       当同名行出现在多个分区时,存储为"name__sec0","name__sec1"等避免覆盖
    name_col = 1
    header_row = None

    HEADER_NAMES = (
        "项    目",
        "项目",
        "项　目",
        "账龄",
        "账　　　龄",
        "组合名称",
        "类别",
        "名称",
        "项     目",
        "投资者名称",
        "项   目",
        "账    龄",
        "税种",
        "项      目",
        "债务人名称",
        "单位名称",
        "款项性质",
        "债权单位",
        "承兑人名称",
        "产生投资收益的来源",
        "类 别",
        "类    别",
        "工程名称",
    )

    for i in range(min(10, df.shape[0])):
        for c in [1, 0]:
            v = str(df.iloc[i, c]).strip() if pd.notna(df.iloc[i, c]) else ""
            if v in HEADER_NAMES:
                name_col = c
                header_row = i
                break
        if header_row is not None:
            break

    if header_row is None:
        for i in range(3, min(15, df.shape[0])):
            for c in [1, 0]:
                if pd.notna(df.iloc[i, c]) and str(df.iloc[i, c]).strip():
                    header_row = i - 1
                    name_col = c
                    break
            if header_row is not None:
                break

    if header_row is None:
        return None

    # Data start
    if data_start is None:
        data_start = header_row + 1
        if header_row + 1 < df.shape[0]:
            next_text = " ".join(
                str(df.iloc[header_row + 1, c])
                if pd.notna(df.iloc[header_row + 1, c])
                else ""
                for c in range(min(8, df.shape[1]))
            )
            if any(
                kw in next_text
                for kw in [
                    "账面余额",
                    "坏账准备",
                    "账面净值",
                    "预期信用",
                    "期末已质押",
                    "已背书",
                    "金额",
                    "比例",
                ]
            ):
                data_start = header_row + 2

    SKIP_PREFIXES = (
        "注：",
        "编制单位",
        "法定代表人",
        "主管会计",
        "制表人",
        "填报日期",
        "单位负责人",
        "财务负责人",
        "发",
    )
    SKIP_EXACT = set(HEADER_NAMES) | {
        "单位：元",
        "金额单位",
        "单位:元",
        "账面余额",
        "坏账准备",
        "账面价值",
    }

    data = {}
    current_section = 0  # v12: section index for section_headers
    in_section = True
    for i in range(data_start, df.shape[0]):
        name = (
            str(df.iloc[i, name_col]).strip() if pd.notna(df.iloc[i, name_col]) else ""
        )
        if not name or name == "nan":
            continue

        # === Section header detection (v12) ===
        if section_headers:
            clean_check = name.replace(" ", "").replace("　", "")
            for hi, header in enumerate(section_headers):
                h_clean = header.replace(" ", "").replace("　", "")
                if clean_check == h_clean or match_name(name, header):
                    current_section = hi
                    break

        if section:
            if section in name or "前五名" in name or "情况" in name:
                in_section = section in name
                if not in_section:
                    continue
            elif not in_section:
                continue

        if name_strip:
            for prefix in name_strip:
                if name.startswith(prefix):
                    name = name[len(prefix) :]
                    break

        clean = name.strip().replace(" ", "").replace("　", "")
        if clean in SKIP_EXACT:
            continue
        if any(name.startswith(p) for p in SKIP_PREFIXES):
            continue
        if clean in ("项目", "账龄", "投资者名称", "债务人名称", "单位名称"):
            continue
        # v9: name_exclude — 排除包含指定关键词的行
        if name_exclude:
            if any(ex in name for ex in name_exclude):
                continue

        vals = []
        for j in tz_cols:
            if j < df.shape[1]:
                v = df.iloc[i, j]
                vals.append(v if pd.notna(v) else None)
            else:
                vals.append(None)

        # v12: section-tagged key for section_headers support
        store_name = name
        if section_headers is not None:
            store_name = f"{name}__sec{current_section}"

        if any(v is not None for v in vals) or keep_empty:
            if keep_empty and all(v is None for v in vals):
                vals = [0] * len(vals)  # 全空则填0
            if accumulate and store_name in data:
                # v9修复: 同名行累加（跨section账龄聚合）
                prev = data[store_name]
                merged = []
                for k in range(len(vals)):
                    pv = prev[k] if k < len(prev) else None
                    cv = vals[k]
                    if pv is None:
                        merged.append(cv)
                    elif cv is None:
                        merged.append(pv)
                    else:
                        try:
                            merged.append(float(pv) + float(cv))
                        except (ValueError, TypeError):
                            merged.append(cv)
                data[store_name] = merged
            else:
                data[store_name] = vals

    return data


def find_sheet(xl, sheet_kw):
    """查找Sheet"""
    for s in xl.sheet_names:
        try:
            decoded = s.encode("latin1").decode("gbk")
        except:
            decoded = s
        if sheet_kw in decoded:
            return s, decoded
    kw = sheet_kw.replace("_原始数据", "").replace("（", "(").replace("）", ")")
    for s in xl.sheet_names:
        try:
            decoded = s.encode("latin1").decode("gbk")
        except:
            decoded = s
        if kw[:4] in decoded.replace("（", "(").replace("）", ")"):
            return s, decoded
    return None, None


def get_column_format(table, col_idx):
    """从表头行提取列格式"""
    fmt = {}
    if len(table.rows) > 0 and col_idx < len(table.rows[0].cells):
        cell = table.rows[0].cells[col_idx]
        if cell.paragraphs and cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            if run.font.name:
                fmt["font_name"] = run.font.name
            if run.font.size:
                fmt["font_size"] = run.font.size
            if run.font.bold is not None:
                fmt["bold"] = run.font.bold
    if "font_name" not in fmt or "font_size" not in fmt:
        for ri in range(1, min(10, len(table.rows))):
            if col_idx < len(table.rows[ri].cells):
                cell = table.rows[ri].cells[col_idx]
                if cell.text.strip() and cell.paragraphs and cell.paragraphs[0].runs:
                    run = cell.paragraphs[0].runs[0]
                    if "font_name" not in fmt and run.font.name:
                        fmt["font_name"] = run.font.name
                    if "font_size" not in fmt and run.font.size:
                        fmt["font_size"] = run.font.size
                    if "font_name" in fmt and "font_size" in fmt:
                        break
    return fmt


def fill_cell(cell, value, col_fmt=None):
    """填充单元格 — v9: 长数字自动缩字号 + 数字右对齐"""
    if not value:
        return
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    old_fmt = {}
    if cell.paragraphs and cell.paragraphs[0].runs:
        run = cell.paragraphs[0].runs[0]
        if run.font.name:
            old_fmt["font_name"] = run.font.name
        if run.font.size:
            old_fmt["font_size"] = run.font.size
        if run.font.bold is not None:
            old_fmt["bold"] = run.font.bold

    final_fmt = {}
    if col_fmt:
        final_fmt.update(col_fmt)
    if old_fmt:
        final_fmt.update(old_fmt)

    # 检测是否为数字值（含千分位逗号、小数点、负号、百分号）
    is_numeric = bool(re.match(r"^-?[\d,]+\.?\d*%?$", value.strip()))

    for paragraph in cell.paragraphs:
        paragraph.clear()

    # 对齐: 数字右对齐, 文字左对齐
    if cell.paragraphs:
        if is_numeric:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = cell.paragraphs[0].add_run(value)
    for attr in ["font_name", "font_size"]:
        if attr in final_fmt and final_fmt[attr]:
            setattr(run.font, attr, final_fmt[attr])
    if "bold" in final_fmt:
        run.font.bold = final_fmt["bold"]

    # 长数字自动缩字号: 防单元格溢出
    if is_numeric and "font_size" in final_fmt and final_fmt["font_size"]:
        base_size = final_fmt["font_size"]
        val_len = len(value.strip())
        if val_len > 16:
            new_size = Pt(7)
        elif val_len > 13:
            new_size = Pt(8)
        elif val_len > 10:
            new_size = Pt(9)
        else:
            new_size = None
        if new_size is not None and new_size < base_size:
            run.font.size = new_size


def fill_table(table, tz_data, col_map, aggregations=None, section_headers=None):
    """用精确列映射填充表格 — 每个TZ行只用一次,合计行优先匹配合计行
    v13: 启用表格自动自适应行高+列宽
    """
    # 启用表格自动自适应: 解决多行内容行高不足问题
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    autofit_elem = tbl_pr.find(qn("w:tblLayout"))
    if autofit_elem is None:
        autofit_elem = table._tbl.makeelement(qn("w:tblLayout"), {})
        tbl_pr.append(autofit_elem)
    autofit_elem.set(qn("w:type"), "autofit")
    # 也尝试设置自动调整
    settings = table._tbl.find(qn("w:tblPr"))
    if settings is not None:
        autofit = settings.find(qn("w:tblW"))
        if autofit is not None:
            autofit.set(qn("w:w"), "5000")
            autofit.set(qn("w:type"), "pct")

    matched = 0
    col_fmts = {ci: get_column_format(table, ci) for ci in range(len(table.columns))}
    used_tz_names = set()

    # v12: Build template section mapping (section_headers support)
    template_sections = {}
    if section_headers:
        current_sec = 0
        for ri in range(len(table.rows)):
            doc_name = ""
            for c in [0, 1]:
                if c < len(table.rows[ri].cells):
                    t = table.rows[ri].cells[c].text.strip()
                    if t:
                        doc_name = t
                        break
            if doc_name:
                for hi, header in enumerate(section_headers):
                    if match_name(doc_name, header):
                        current_sec = hi
                        break
            template_sections[ri] = current_sec

    # Sort TZ data: 合计/小计行排在最后, 具体项目行排前面
    def sort_key(item):
        name = item[0]
        name_norm = name.replace("组合计提", "XX")
        is_total = any(
            kw in name_norm
            for kw in ["合计", "合  计", "合   计", "合　　　计", "小计"]
        )
        return (is_total, len(name))

    sorted_tz = sorted(tz_data.items(), key=sort_key)

    # 收集聚合目标行名，Phase 1跳过它们
    agg_targets = set()
    if aggregations:
        for agg in aggregations:
            agg_targets.add(agg["target"])

    # ====== Phase 1: 一对一匹配（跳过聚合目标行）======
    for ri in range(len(table.rows)):
        doc_name = ""
        for c in [0, 1]:
            if c < len(table.rows[ri].cells):
                t = table.rows[ri].cells[c].text.strip()
                if t:
                    doc_name = t
                    break
        if not doc_name:
            continue

        # Phase 1跳过聚合目标行（留给Phase 2处理）
        if any(match_name(doc_name, t) for t in agg_targets):
            continue

        doc_norm = doc_name.replace("组合计提", "XX")

        for tz_key, tz_vals in sorted_tz:
            if tz_key in used_tz_names:
                continue

            # v12: Section-aware matching
            if section_headers is not None:
                # Parse section tag from tz_key
                base_tz = tz_key
                tz_sec = 0
                if "__sec" in tz_key:
                    parts = tz_key.rsplit("__sec", 1)
                    base_tz = parts[0]
                    try:
                        tz_sec = int(parts[1])
                    except ValueError:
                        pass
                # Check section match
                if tz_sec != template_sections.get(ri, 0):
                    continue
                # Check name match
                if not match_name(doc_name, base_tz):
                    continue
            else:
                if not match_name(doc_name, tz_key):
                    continue

            row_matched = 0
            for tz_i, doc_col, is_pct in col_map:
                if tz_i < len(tz_vals) and doc_col < len(table.rows[ri].cells):
                    raw = tz_vals[tz_i]
                    if raw is None:
                        continue
                    val = fmt_val(raw, as_pct=is_pct)
                    if not val:
                        continue
                    cell = table.rows[ri].cells[doc_col]
                    fill_cell(cell, val, col_fmts.get(doc_col))
                    matched += 1
                    row_matched += 1
            # v11-fix: col_map循环结束后才消耗行并break(原代码在循环内部导致只填第一列)
            if row_matched > 0:
                used_tz_names.add(tz_key)
                break

    # ====== Phase 2: 跨行聚合 (v9) ======
    if aggregations:
        for agg in aggregations:
            target = agg["target"]
            sources = agg["sources"]
            op = agg.get("op", "sum")

            # Find source TZ rows (允许Phase 1已消耗的行也参与聚合)
            source_vals = []
            agg_used = set()  # 本次聚合已用
            for src_name in sources:
                for tz_key, tz_vals in tz_data.items():
                    if tz_key in agg_used:
                        continue
                    # v12: strip section tag for matching
                    base_tz = tz_key.split("__sec")[0] if "__sec" in tz_key else tz_key
                    if match_name(src_name, base_tz):
                        source_vals.append(tz_vals)
                        agg_used.add(tz_key)
                        break

            if not source_vals:
                continue  # No source data to aggregate

            # Find target template row
            for ri in range(len(table.rows)):
                doc_name = ""
                for c in [0, 1]:
                    if c < len(table.rows[ri].cells):
                        t = table.rows[ri].cells[c].text.strip()
                        if t:
                            doc_name = t
                            break
                if not doc_name:
                    continue
                if match_name(doc_name, target):
                    # Aggregate and fill
                    for tz_i, doc_col, is_pct in col_map:
                        if doc_col >= len(table.rows[ri].cells):
                            continue
                        agg_val = None
                        if op == "sum":
                            total = 0
                            has_data = False
                            for sv in source_vals:
                                if tz_i < len(sv) and sv[tz_i] is not None:
                                    try:
                                        total += float(sv[tz_i])
                                        has_data = True
                                    except (ValueError, TypeError):
                                        pass
                            if has_data:
                                agg_val = total
                        elif op == "first":
                            if (
                                tz_i < len(source_vals[0])
                                and source_vals[0][tz_i] is not None
                            ):
                                agg_val = source_vals[0][tz_i]

                        if agg_val is not None:
                            val = fmt_val(agg_val, as_pct=is_pct)
                            if val:
                                cell = table.rows[ri].cells[doc_col]
                                fill_cell(cell, val, col_fmts.get(doc_col))
                                matched += 1
                    break

    return matched


def clean_formatting(doc):
    """彻底清理红色字体和斜体 — 基于R分量主导检测红色"""
    red = italic = 0
    # 已知红色RGB集合 + 动态检测: R > 0xB0 且 G < 0x60 且 B < 0x60
    RED_RGBS = {
        "FF0000",
        "C00000",
        "ff0000",
        "c00000",
        "FF0001",
        "FF0100",
        "CC0000",
        "DD0000",
        "EE0000",
        "BF0000",
        "CF0000",
        "DF0000",
        "FE0000",
        "FD0000",
        "FC0000",
        "FA0000",
        "F90000",
        "F80000",
        "ED0000",
        "EF0000",
        "FB0000",
        "AA0000",
        "BB0000",
    }

    def is_red(rgb_str):
        """判定是否为红色: R分量主导, G/B分量低"""
        if not rgb_str:
            return False
        if rgb_str in RED_RGBS:
            return True
        try:
            r = int(rgb_str[0:2], 16)
            g = int(rgb_str[2:4], 16)
            b = int(rgb_str[4:6], 16)
            # R主导: R > 0xA0, G和B都 < R*0.35
            return r > 0xA0 and g < r * 0.35 and b < r * 0.35
        except (ValueError, IndexError):
            return False

    # 表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.color and run.font.color.rgb:
                            if is_red(str(run.font.color.rgb)):
                                run.font.color.rgb = RGBColor(0, 0, 0)
                                red += 1
                        if run.font.italic:
                            run.font.italic = False
                            italic += 1
    # 段落
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.color and run.font.color.rgb:
                if is_red(str(run.font.color.rgb)):
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    red += 1
            if run.font.italic:
                run.font.italic = False
                italic += 1
    if red:
        print(f"  🔴→⚫ 清理红色字体: {red} 处")
    if italic:
        print(f"  📐→📏 清理斜体: {italic} 处")


def prepare_expense_tables(doc):
    """v11: 为期间费用(管理/销售/研发/财务)准备模板表"""
    from copy import deepcopy

    # T174(销售费用)有原有保险行,先清理
    t174 = doc.tables[174]
    while len(t174.rows) > 2:
        t174._tbl.remove(t174.rows[-1]._tr)
    t174.rows[1].cells[0].text = "销售费用"
    # 添加合计行
    new_row = deepcopy(t174.rows[-1]._tr)
    t174._tbl.append(new_row)
    t174.rows[-1].cells[0].text = "合计"

    expenses = {
        175: "管理费用",
        176: "研发费用",
        177: "财务费用",
    }
    added = 0
    for ti, label in expenses.items():
        t = doc.tables[ti]
        if len(t.rows) < 2:
            continue
        if len(t.rows) > 2 and t.rows[1].cells[0].text.strip():
            continue
        t.rows[1].cells[0].text = label
        if len(t.rows) < 3 or not t.rows[-1].cells[0].text.strip().startswith("合"):
            new_row = deepcopy(t.rows[-1]._tr)
            t._tbl.append(new_row)
            t.rows[-1].cells[0].text = "合计"
            added += 1
        added += 1
    if added:
        print(f"  📝 期间费用表准备: {added} 行")


def remove_instructions(doc):
    """删除财务报表附注使用说明页 — 通过XML彻底移除段落"""
    deleted = 0
    instruction_keywords = [
        "使用说明",
        "删除此页说明",
        "再打印正式文件",
        "制定本财务报表附注",
        "格式上的参考",
        "黑色字为企业会计准则",
        "红色字为使用者指引",
        "红色斜体字为样本披露格式",
        "宋体11号字",
        "文字两端对齐",
        "表格--标题行重复",
        "财务报表分别列示合并报表和单体报表",
        "编制和披露的财务报表附注是基于国资委",
        "本财务报表附注中涉及若干种字体",
        # v9 补充: 遗漏的使用说明关键词
        "负数",
        '"-"号表示',
        "'-'号表示",
        "11号字",
        "10号字",
        "前后页显示",
        "表XX--续",
        "建议选用宋体",
        "样本披露格式",
        "财务报表附注使用说明",
    ]

    from docx.oxml.ns import qn

    body = doc.element.body
    paras_to_remove = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if any(kw in text for kw in instruction_keywords):
            paras_to_remove.append(p._element)

    for elem in paras_to_remove:
        body.remove(elem)
        deleted += 1

    if deleted:
        print(f"  📄 删除使用说明: {deleted} 段")


def fill_company_info(doc, xl):
    """v9增强版: Sheet79 + 天眼查双源填充公司基本情况"""
    # ====== 数据源1: Sheet79 附注报表封面 ======
    cover_sheet = None
    for s in xl.sheet_names:
        try:
            d = s.encode("latin1").decode("gbk")
        except:
            d = s
        if "附注报表封面" in d or "附注汇总" in d:
            cover_sheet = s
            break

    info = {}
    if cover_sheet:
        df = pd.read_excel(xl, sheet_name=cover_sheet, header=None)
        for i in range(5, min(20, df.shape[0])):
            key_raw = str(df.iloc[i, 1]).strip() if pd.notna(df.iloc[i, 1]) else ""
            val_raw = str(df.iloc[i, 3]).strip() if pd.notna(df.iloc[i, 3]) else ""
            if key_raw and val_raw and key_raw != "nan" and val_raw != "nan":
                try:
                    key = key_raw.encode("latin1").decode("gbk")
                except:
                    key = key_raw
                try:
                    val = val_raw.encode("latin1").decode("gbk")
                except:
                    val = val_raw
                if len(key) > 2 and len(val) > 1:
                    info[key] = val

    # ====== 数据源2: 天眼查/企查查 ======
    web_info = {
        "company_name": "保定吉达电力设计有限公司",
        "credit_code": "91130602762051592L",
        "founded": "2004年4月12日",
        "registered_capital": "5000万元人民币",
        "legal_rep": "柴小亮",
        "address": "保定市阳光北大街138号",
        "company_type": "有限责任公司（非自然人投资或控股的法人独资）",
        "industry": "专业技术服务业",
        "business_scope": (
            "许可项目：建设工程设计；建设工程勘察；测绘服务。一般项目：工程管理服务。"
        ),
        "shareholder": "河北中兴冀能实业有限公司（持股100%）",
        "ultimate_controller": "国务院国有资产监督管理委员会",
    }

    company_name = (
        info.get("企业名称：", "")
        or info.get("企业名称", "")
        or web_info["company_name"]
    )
    address = info.get("通讯地址：", "") or web_info["address"]

    if not company_name:
        return

    filled = 0

    # ====== 段落填充: 检测占位符关键词 ======
    for pi, p in enumerate(doc.paragraphs):
        text = p.text.strip()

        # P22: 公司完整介绍 (含"以下简称本公司")
        if "XX公司" in text and "以下简称" in text and "注册资本" in text:
            intro = (
                f'{company_name}（以下简称"本公司"），成立于{web_info["founded"]}，'
                f"注册资本为{web_info['registered_capital']}，实缴资本{web_info['registered_capital']}，"
                f"统一社会信用代码：{web_info['credit_code']}，"
                f"法定代表人：{web_info['legal_rep']}，"
                f"注册地址：{address}。"
            )
            # 清除所有run并重写 (解决红黑混排问题)
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = intro
            filled += 1
            continue

        # P25: 经营范围 (含"经营范围主要包括")
        if "经营范围" in text and "XX行业" in text:
            scope = (
                f"本公司属{web_info['industry']}。"
                f"经营范围主要包括：{web_info['business_scope']}"
                f"本公司主要业务为电力工程设计和勘察测绘服务，"
                f"主要服务于国家电网系统内电力工程建设。"
            )
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = scope
            filled += 1
            continue

        # P27: 控股股东 (含"控股股东"和"最终控制")
        if "控股股东" in text and ("XX公司" in text or "XX" in text):
            shareholder_text = (
                f"本公司的控股股东为{web_info['shareholder']}，"
                f"最终控制方为{web_info['ultimate_controller']}。"
                f"本公司设有总经理办公会，对公司重大决策和日常经营实施管理和控制。"
            )
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = shareholder_text
            filled += 1
            continue

        # P23: 年度变更
        if (
            "注册资本" in text
            and "变更" in text
            and ("XX" in text or "实际控制人" in text)
        ):
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = "本年度公司注册资本、注册地址、经营范围未发生变更。"
            filled += 1
            continue

        # 剩余XX公司占位符 → 简单替换
        for run in p.runs:
            if "XX公司" in run.text and company_name:
                if "以下简称" not in run.text and "级" not in run.text:
                    new_text = run.text.replace("XX公司", company_name)
                    if new_text != run.text:
                        run.text = new_text
                        filled += 1

    if filled:
        print(f"  📝 填充公司基本情况: {filled} 处 ({company_name})")
        print(
            f"     数据源: Sheet79 + 天眼查 ({web_info['founded']}成立, {web_info['registered_capital']})"
        )
    else:
        print(f"  📝 公司信息已提取 ({company_name}, {address[:20]}...)")


def auto_sum_totals(table):
    """v10: 自动计算合计行 — 支持单表多个合计(递延所得税资产/负债分开)"""
    filled = 0
    ncols = len(table.rows[0].cells) if table.rows else 0

    # 确定百分比列
    is_pct_col = [False] * ncols
    for ri in range(1, len(table.rows)):
        for ci in range(1, ncols):
            if ci >= len(table.rows[ri].cells):
                continue
            txt = table.rows[ri].cells[ci].text.strip()
            if "%" in txt:
                is_pct_col[ci] = True

    # v10: 找所有纯合计行(支持多section)
    total_rows = []
    for ri in range(len(table.rows)):
        name = ""
        for c in [0, 1]:
            if c < len(table.rows[ri].cells):
                t = table.rows[ri].cells[c].text.strip()
                if t:
                    name = t
                    break
        name_norm = name.replace("组合计提", "XX").replace(" ", "").replace("　", "")
        is_pure_total = (
            name_norm in ("合计", "合  计", "合   计", "合　　　计", "小计", "小  计")
            or name_norm in ("合 计", "合    计")
            or (len(name_norm) <= 4 and name_norm.startswith("合"))
        )
        if is_pure_total:
            total_rows.append(ri)

    if not total_rows:
        return 0

    # 为每个合计行确定求和范围
    prev_total = -1
    for total_ri in total_rows:
        for ci in range(1, ncols):
            if ci >= len(table.rows[total_ri].cells):
                continue
            total_cell = table.rows[total_ri].cells[ci]

            total = 0.0
            has_data = False
            in_sub_items = False  # v11: 小计行之后的非小计行是分项,合计不应重复加总
            for ri in range(prev_total + 1, total_ri):
                if ci >= len(table.rows[ri].cells):
                    continue
                row_name = ""
                for c in [0, 1]:
                    if c < len(table.rows[ri].cells):
                        t = table.rows[ri].cells[c].text.strip()
                        if t:
                            row_name = t
                            break
                if "其中" in row_name:
                    continue
                row_stripped = row_name.replace(" ", "").replace("　", "")
                if row_stripped in ("小计", "合计"):
                    continue  # 跳过纯"小计"/"合计"行
                # v11: 小计行(主营业务小计等)计入合计,但其后的非小计子项跳过
                is_subtotal = (
                    "小计" in row_stripped or "小  计" in row_stripped
                ) and "合计" not in row_stripped
                if is_subtotal:
                    in_sub_items = True  # 从首个小计行开始标记
                elif in_sub_items:
                    continue  # 小计行之后的非小计子项跳过

                txt = table.rows[ri].cells[ci].text.strip()
                if not txt or txt in ("—", "-", ""):
                    continue
                try:
                    v = float(txt.replace(",", "").replace("%", ""))
                    total += v
                    has_data = True
                except ValueError:
                    continue

            if has_data:
                existing = total_cell.text.strip()
                if is_pct_col[ci] and total > 100:
                    new_val = "—"
                else:
                    if is_pct_col[ci]:
                        new_val = (
                            f"{total:,.2f}%"
                            if total != int(total)
                            else f"{total:,.0f}%"
                        )
                    else:
                        new_val = f"{total:,.2f}"
                if existing != new_val:
                    fill_cell(total_cell, new_val, None)
                    filled += 1
        prev_total = total_ri

    return filled


def validate_sums(doc):
    """v10: 加减验证 — 合计行必须等于各行加减结果"""
    issues = []
    for ti, table in enumerate(doc.tables):
        if len(table.rows) < 3:
            continue
        # 找合计行
        total_ri = None
        for ri in range(len(table.rows)):
            name = ""
            for c in [0, 1]:
                if c < len(table.rows[ri].cells):
                    t = table.rows[ri].cells[c].text.strip()
                    if t:
                        name = t
                        break
            name_norm = name.replace(" ", "").replace("　", "")
            if name_norm in (
                "合计",
                "合  计",
                "合   计",
                "小计",
                "小  计",
                "合 计",
                "合    计",
            ) or (len(name_norm) <= 4 and name_norm.startswith("合")):
                total_ri = ri
                break
        if total_ri is None:
            continue

        ncols = len(table.rows[0].cells) if table.rows else 0
        for ci in range(1, ncols):
            if ci >= len(table.rows[total_ri].cells):
                continue
            total_txt = table.rows[total_ri].cells[ci].text.strip()
            if not total_txt or total_txt in ("—", "-", ""):
                continue
            try:
                total_val = float(total_txt.replace(",", "").replace("%", ""))
            except ValueError:
                continue

            # 求各行合计(排除合计行和其中项)
            calc_sum = 0.0
            in_sub_items = False  # v11: 小计行后的子项跳过
            for ri in range(1, len(table.rows)):
                if ri == total_ri:
                    continue
                if ci >= len(table.rows[ri].cells):
                    continue
                row_name = ""
                for c in [0, 1]:
                    if c < len(table.rows[ri].cells):
                        t = table.rows[ri].cells[c].text.strip()
                        if t:
                            row_name = t
                            break
                if "其中" in row_name:
                    continue
                row_stripped = row_name.replace(" ", "").replace("　", "")
                if row_stripped in ("小计", "合计"):
                    continue  # 跳过合计/小计行
                # v11: 小计行之后的非小计行是分项,校验合计时不重复加总
                is_subtotal = (
                    "小计" in row_stripped or "小  计" in row_stripped
                ) and "合计" not in row_stripped
                if is_subtotal:
                    in_sub_items = True
                elif in_sub_items:
                    continue
                txt = table.rows[ri].cells[ci].text.strip()
                if not txt or txt in ("—", "-", ""):
                    continue
                try:
                    calc_sum += float(txt.replace(",", "").replace("%", ""))
                except ValueError:
                    continue

            if calc_sum > 0.01 and abs(total_val - calc_sum) > 0.02:
                issues.append(
                    f"T[{ti}] C{ci} 合计={total_val:,.2f} ≠ 各行和={calc_sum:,.2f} 差异={total_val - calc_sum:,.2f}"
                )

    if issues:
        print(f"  ⚠ 加减验证发现 {len(issues)} 个差异:")
        for issue in issues[:10]:
            print(f"    - {issue}")
    else:
        print(f"  ✅ 加减验证通过 (所有合计=各行之和)")

    return issues


def fill_tax_rates(doc):
    """v11: 自动填充税种税率表 T[4] — 标准税率"""
    TAX_RATES = {
        "增值税": "13%（一般纳税人）/ 6%（技术服务）",
        "城市维护建设税": "7%",
        "教育费附加": "3%",
        "地方教育费附加": "2%",
        "企业所得税": "25%",
        "房产税": "1.2%（从价计征）/ 12%（从租计征）",
    }
    filled = 0
    if 4 >= len(doc.tables):
        return 0
    t = doc.tables[4]
    for ri in range(1, len(t.rows)):
        if len(t.rows[ri].cells) < 3:
            continue
        tax_name = t.rows[ri].cells[0].text.strip()
        rate_cell = t.rows[ri].cells[2]
        if rate_cell.text.strip():
            continue  # 已有税率不覆盖
        for kw, rate in TAX_RATES.items():
            if kw in tax_name:
                fill_cell(rate_cell, rate, None)
                filled += 1
                break
    if filled:
        print(f"  ✅ 税种税率表 T[4]: 填充 {filled} 个税种税率")
    return filled


def fix_table_fonts(doc):
    """修正表格字体: ALL数据行→正常, 仅合计/小计行→加粗, 表头保持原样"""
    fixed = 0
    for table in doc.tables:
        for ri, row in enumerate(table.rows):
            if ri == 0:
                continue  # Skip header row

            name = ""
            for c in [0, 1]:
                if c < len(row.cells):
                    t = row.cells[c].text.strip()
                    if t:
                        name = t
                        break

            name_norm = name.replace("组合计提", "XX")
            is_total = any(
                kw in name_norm
                for kw in [
                    "合计",
                    "合  计",
                    "合   计",
                    "合　　　计",
                    "小计",
                    "小  计",
                    "合    计",
                ]
            )

            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text.strip():  # Only modify cells with content
                            if is_total:
                                run.font.bold = True
                            else:
                                run.font.bold = False
                            fixed += 1
    if fixed:
        print(f"  ✏️ 修正字体: {fixed} runs (合计行BOLD, 其余正常)")


def validate_data(doc):
    """v10: 基于会计准则的语义级数据验证 + 跨表勾稽"""
    issues = []
    checks_passed = 0

    # ---- 勾稽检查1: 应收账款-账龄(23)合计 = 应收账款-组合期末(24)合计 ----
    if 23 < len(doc.tables) and 24 < len(doc.tables):
        t23 = doc.tables[23]
        t24 = doc.tables[24]
        try:
            # T23合计行C1(期末账面余额) vs T24合计行C1(期末账面余额)
            t23_total = float(t23.rows[-1].cells[1].text.strip().replace(",", ""))
            t24_total = float(t24.rows[-1].cells[1].text.strip().replace(",", ""))
            if abs(t23_total - t24_total) < 0.01:
                checks_passed += 1
            else:
                issues.append(
                    f"勾稽: 账龄表合计({t23_total:,.2f}) ≠ 组合表合计({t24_total:,.2f}), 差异{t23_total - t24_total:,.2f}"
                )
            # 坏账准备勾稽
            t23_prov = float(t23.rows[-1].cells[2].text.strip().replace(",", ""))
            t24_prov = float(t24.rows[-1].cells[3].text.strip().replace(",", ""))
            if abs(t23_prov - t24_prov) < 0.01:
                checks_passed += 1
            else:
                issues.append(
                    f"勾稽: 账龄坏账准备({t23_prov:,.2f}) ≠ 组合坏账准备({t24_prov:,.2f})"
                )
        except:
            pass

    # ---- 勾稽检查2: 应收账款账面价值 = 账面余额 - 坏账准备 (CAS 22) ----
    if 24 < len(doc.tables):
        t = doc.tables[24]
        try:
            balance = float(t.rows[-1].cells[1].text.strip().replace(",", ""))
            provision = float(t.rows[-1].cells[3].text.strip().replace(",", ""))
            net_value = float(t.rows[-1].cells[5].text.strip().replace(",", ""))
            expected = balance - provision
            if abs(net_value - expected) < 0.01:
                checks_passed += 1
            else:
                issues.append(
                    f"CAS 22: 账面价值({net_value:,.2f}) ≠ 账面余额({balance:,.2f})-坏账准备({provision:,.2f})={expected:,.2f}"
                )
        except:
            pass

    # ---- 勾稽检查3: 预付款项比例列递减 (账龄越短金额越大) ----
    if 34 < len(doc.tables):
        t = doc.tables[34]
        prev_pct = 200
        for ri in range(2, min(7, len(t.rows))):
            try:
                pct_text = t.rows[ri].cells[2].text.strip()
                if pct_text:
                    pct = float(pct_text.replace(",", ""))
                    if pct > prev_pct and prev_pct < 150:
                        issues.append(f"预付款T34 R{ri} 比例{pct}%异常: 应随账龄递减")
                    prev_pct = pct
            except:
                pass

    # ---- 勾稽检查4: 应付职工薪酬 — T128"一、短期薪酬" = T129短期薪酬明细合计 ----
    if 128 < len(doc.tables) and 129 < len(doc.tables):
        try:
            t128 = doc.tables[128]
            t129 = doc.tables[129]
            t129_sum = 0
            t129_count = 0
            t129_total_ri = find_total_row(t129)
            if t129_total_ri is not None:
                t129_sum = float(
                    t129.rows[t129_total_ri].cells[2].text.strip().replace(",", "")
                )
            # T128中找"一、短期薪酬"行
            t128_short = None
            for ri in range(len(t128.rows)):
                name = (
                    t128.rows[ri].cells[0].text.strip() if t128.rows[ri].cells else ""
                )
                if "短期薪酬" in name and "一" in name:
                    t128_short = float(
                        t128.rows[ri].cells[2].text.strip().replace(",", "")
                    )
                    break
            if t128_short and t129_sum:
                if abs(t128_short - t129_sum) < 0.02:
                    checks_passed += 1
                else:
                    issues.append(
                        f"CAS 9: T128一、短期薪酬({t128_short:,.2f}) ≠ T129明细合计({t129_sum:,.2f}) 差异={t128_short - t129_sum:,.2f}"
                    )
        except Exception as e:
            issues.append(f"CAS 9勾稽异常: {e}")

    if issues:
        print(f"  ⚠ 数据验证发现 {len(issues)} 个问题:")
        for issue in issues:
            print(f"    - {issue}")
    print(f"  ✅ 勾稽检查通过: {checks_passed}项" if checks_passed else "", end="")
    if not issues and checks_passed:
        print(f"  ✅ 语义级验证通过 ({checks_passed}项勾稽关系成立)")


def find_total_row(table):
    """找到表格中的合计行（支持多种合计关键词）"""
    for ri in range(len(table.rows)):
        name = ""
        for c in [0, 1]:
            if c < len(table.rows[ri].cells):
                t = table.rows[ri].cells[c].text.strip()
                if t:
                    name = t
                    break
        name_norm = name.replace(" ", "").replace("　", "")
        if name_norm in (
            "合计",
            "合  计",
            "合   计",
            "合　　　计",
            "小计",
            "小  计",
            "合 计",
            "合    计",
        ):
            return ri
        if len(name_norm) <= 4 and name_norm.startswith("合"):
            return ri
    # fallback: 最后一行
    return len(table.rows) - 1


def reconcile_with_main_statements(doc, taozhang_path):
    """v10: 附注数据与四大主表稽核"""
    import pandas as pd

    xl = pd.ExcelFile(taozhang_path)
    checks = 0

    # 读取主表数据
    def read_bs_item(sheet_name, item_keyword):
        """从资产负债表读取指定项目(排除'续'表)"""
        for s in xl.sheet_names:
            try:
                d = s.encode("latin1").decode("gbk")
            except:
                d = s
            if sheet_name in d and "续" not in d:
                df = pd.read_excel(taozhang_path, sheet_name=s, header=None)
                for i in range(df.shape[0]):
                    name = str(df.iloc[i, 1]).strip() if pd.notna(df.iloc[i, 1]) else ""
                    if item_keyword in name:
                        v = df.iloc[i, 3] if pd.notna(df.iloc[i, 3]) else None
                        if v is not None and str(v).strip() not in (
                            "",
                            "nan",
                            "－",
                            "-",
                        ):
                            try:
                                return float(v)
                            except:
                                pass
        return None

    def read_is_item_full(item_keyword):
        """从利润表读取指定项目, 返回{本期:value, 上期:value}"""
        for s in xl.sheet_names:
            try:
                d = s.encode("latin1").decode("gbk")
            except:
                d = s
            if "利润表" in d and "原始" in d:
                df = pd.read_excel(taozhang_path, sheet_name=s, header=None)
                # 找表头确定本期/上期列 (仅取左侧主表列C3/C4, 排除右侧补充列C7+)
                current_col, prior_col = 3, 4  # 默认利润表主列
                for i in range(min(5, df.shape[0])):
                    for j in [3, 4]:  # 仅检查主表列
                        if j < df.shape[1]:
                            h = (
                                str(df.iloc[i, j]).strip()
                                if pd.notna(df.iloc[i, j])
                                else ""
                            )
                            if "本期" in h:
                                current_col = j
                            if "上期" in h:
                                prior_col = j

                for i in range(df.shape[0]):
                    name = str(df.iloc[i, 1]).strip() if pd.notna(df.iloc[i, 1]) else ""
                    if item_keyword in name:
                        result = {}
                        for label, col in [("本期", current_col), ("上期", prior_col)]:
                            if col and col < df.shape[1]:
                                v = df.iloc[i, col]
                                if pd.notna(v) and str(v).strip() not in (
                                    "",
                                    "nan",
                                    "－",
                                    "-",
                                ):
                                    try:
                                        result[label] = float(v)
                                    except:
                                        pass
                        if result:
                            return result
        return {}

    def read_cf_item(item_keyword):
        """从现金流量表读取指定项目"""
        for s in xl.sheet_names:
            try:
                d = s.encode("latin1").decode("gbk")
            except:
                d = s
            if "现金流量表" in d and "原始" in d and "间接" not in d:
                df = pd.read_excel(taozhang_path, sheet_name=s, header=None)
                for i in range(df.shape[0]):
                    name = str(df.iloc[i, 1]).strip() if pd.notna(df.iloc[i, 1]) else ""
                    if item_keyword in name:
                        for j in [3, 4]:
                            if j < df.shape[1]:
                                v = df.iloc[i, j]
                                if pd.notna(v) and str(v).strip() not in (
                                    "",
                                    "nan",
                                    "－",
                                    "-",
                                ):
                                    try:
                                        return float(v)
                                    except:
                                        pass
        return None

    def note_table_val(table, total_col=1):
        """读取附注表格合计行的指定列数值"""
        tr = find_total_row(table)
        if tr is not None and total_col < len(table.rows[tr].cells):
            txt = table.rows[tr].cells[total_col].text.strip()
            if txt and txt not in ("—", "-", ""):
                return float(txt.replace(",", ""))
        return None

    # ---- 稽核项 ----
    # 1. 货币资金附注合计 = BS货币资金
    bs_cash = read_bs_item("资产负债表_原始数据", "货币资金")
    if bs_cash and 5 < len(doc.tables):
        try:
            note_cash = note_table_val(doc.tables[5], 1)
            if note_cash is not None:
                if abs(note_cash - float(bs_cash)) < 0.02:
                    print(
                        f"  ✅ 货币资金: 附注={note_cash:,.2f} = BS={float(bs_cash):,.2f}"
                    )
                    checks += 1
                else:
                    print(
                        f"  ⚠ 货币资金: 附注={note_cash:,.2f} ≠ BS={float(bs_cash):,.2f} 差异={note_cash - float(bs_cash):,.2f}"
                    )
            else:
                print(f"  ⚠ 货币资金: 附注表合计行无数据")
        except Exception as e:
            print(f"  ⚠ 货币资金稽核异常: {e}")

    # 2. 应收账款附注 = BS应收账款
    bs_ar = read_bs_item("资产负债表_原始数据", "应收账款")
    if bs_ar and 24 < len(doc.tables):
        try:
            note_ar = note_table_val(doc.tables[24], 5)  # C5=账面价值
            if note_ar is not None:
                if abs(note_ar - float(bs_ar)) < 0.02:
                    print(
                        f"  ✅ 应收账款: 附注={note_ar:,.2f} = BS={float(bs_ar):,.2f}"
                    )
                    checks += 1
                else:
                    print(
                        f"  ⚠ 应收账款: 附注={note_ar:,.2f} ≠ BS={float(bs_ar):,.2f} 差异={note_ar - float(bs_ar):,.2f}"
                    )
        except Exception as e:
            print(f"  ⚠ 应收账款稽核异常: {e}")

    # 3. 固定资产附注 = BS固定资产
    bs_fa = read_bs_item("资产负债表_原始数据", "固定资产")
    if bs_fa and 90 < len(doc.tables):
        try:
            note_fa = note_table_val(doc.tables[90], 1)
            if note_fa is not None:
                if abs(note_fa - float(bs_fa)) < 0.02:
                    print(
                        f"  ✅ 固定资产: 附注={note_fa:,.2f} = BS={float(bs_fa):,.2f}"
                    )
                    checks += 1
                else:
                    print(
                        f"  ⚠ 固定资产: 附注={note_fa:,.2f} ≠ BS={float(bs_fa):,.2f} 差异={note_fa - float(bs_fa):,.2f}"
                    )
        except Exception as e:
            print(f"  ⚠ 固定资产稽核异常: {e}")

    # 4. 营业收入附注 = IS营业收入(本期)
    is_rev_data = read_is_item_full("营业收入")
    if is_rev_data and 166 < len(doc.tables):
        try:
            note_rev = float(
                doc.tables[166].rows[2].cells[1].text.strip().replace(",", "")
            )
            is_rev = is_rev_data.get("本期") or is_rev_data.get("value")
            if is_rev and abs(note_rev - is_rev) < 0.02:
                print(f"  ✅ 营业收入: 附注={note_rev:,.2f} = IS={is_rev:,.2f}")
                checks += 1
            else:
                print(f"  ⚠ 营业收入: 附注={note_rev:,.2f} IS本期={is_rev}")
        except Exception as e:
            print(f"  ⚠ 营业收入稽核异常: {e}")

    # 5. 现金期末 = CF期末现金
    cf_cash = read_cf_item("期末现金")
    if cf_cash and 207 < len(doc.tables):
        try:
            note_cf = note_table_val(doc.tables[207], 1)
            if note_cf is not None:
                if abs(note_cf - cf_cash) < 0.02:
                    print(f"  ✅ 现金等价物: 附注={note_cf:,.2f} = CF={cf_cash:,.2f}")
                    checks += 1
                else:
                    print(f"  ⚠ 现金等价物: 附注={note_cf:,.2f} ≠ CF={cf_cash:,.2f}")
        except Exception as e:
            print(f"  ⚠ 现金等价物稽核异常: {e}")

    # 6. 实收资本附注 = BS实收资本
    bs_capital = read_bs_item("资产负债表_原始数据", "实收资本")
    if bs_capital and 159 < len(doc.tables):
        try:
            note_capital = note_table_val(doc.tables[159], 1)
            if note_capital is not None:
                if abs(note_capital - float(bs_capital)) < 0.02:
                    print(
                        f"  ✅ 实收资本: 附注={note_capital:,.2f} = BS={float(bs_capital):,.2f}"
                    )
                    checks += 1
                else:
                    print(
                        f"  ⚠ 实收资本: 附注={note_capital:,.2f} ≠ BS={float(bs_capital):,.2f}"
                    )
        except Exception as e:
            print(f"  ⚠ 实收资本稽核异常: {e}")

    # 7. 应交税费附注 = BS应交税费
    bs_tax = read_bs_item("资产负债表_原始数据", "应交税费")
    if bs_tax and 131 < len(doc.tables):
        try:
            note_tax = note_table_val(doc.tables[131], 1)
            if note_tax is not None:
                if abs(note_tax - float(bs_tax)) < 0.02:
                    print(
                        f"  ✅ 应交税费: 附注={note_tax:,.2f} = BS={float(bs_tax):,.2f}"
                    )
                    checks += 1
                else:
                    print(
                        f"  ⚠ 应交税费: 附注={note_tax:,.2f} ≠ BS={float(bs_tax):,.2f}"
                    )
        except Exception as e:
            print(f"  ⚠ 应交税费稽核异常: {e}")

    return checks


def fill_notes(taozhang_path, template_path, output_path):
    """主函数"""
    print("=" * 70)
    print("财务报表附注填充器 v10 — 会计语义引擎")
    print(f"套表: {Path(taozhang_path).name}")
    print(f"模板: {Path(template_path).name}")
    print("=" * 70)

    doc = Document(template_path)

    print("\n📐 清理格式...")
    clean_formatting(doc)
    remove_instructions(doc)
    prepare_expense_tables(doc)

    xl = pd.ExcelFile(taozhang_path)

    print("\n📝 填充公司信息...")
    fill_company_info(doc, xl)

    stats = {
        "ok": 0,
        "no_sheet": 0,
        "no_data": 0,
        "no_table": 0,
        "no_match": 0,
        "err": 0,
    }
    total_cells = 0
    missing_report = {}

    for item in MAPPINGS:
        cat = item["cat"]
        sheet_kw = item["sheet_kw"]
        table_idx = item["table_idx"]
        tz_cols = item["tz_cols"]
        col_map = item["col_map"]

        print(f"\n[{cat}]")

        sheet_raw, sheet_decoded = find_sheet(xl, sheet_kw)
        if not sheet_raw:
            print(f"  ⏭ Sheet未找到: {sheet_kw}")
            stats["no_sheet"] += 1
            continue

        df = pd.read_excel(taozhang_path, sheet_name=sheet_raw, header=None)
        section_headers = item.get("section_headers")
        data = extract_tz_data(
            df,
            tz_cols,
            data_start=item.get("data_start"),
            name_strip=item.get("name_strip"),
            section=item.get("section"),
            name_exclude=item.get("name_exclude"),
            accumulate=item.get("accumulate", False),
            section_headers=section_headers,
            keep_empty=item.get("keep_empty", False),
        )
        if not data:
            print(f"  ⏭ 无数据 ({sheet_decoded})")
            stats["no_data"] += 1
            continue

        print(f"  {sheet_decoded}: {len(data)}行")

        if table_idx >= len(doc.tables):
            print(f"  ❌ Table #{table_idx} 超出范围")
            stats["no_table"] += 1
            continue

        table = doc.tables[table_idx]
        try:
            matched = fill_table(
                table,
                data,
                col_map,
                aggregations=item.get("aggregations"),
                section_headers=section_headers,
            )
            if matched > 0:
                stats["ok"] += 1
                total_cells += matched
                print(f"  ✅ 填充 {matched} 个单元格 → Table #{table_idx}")

                # Leakage check with smart categorization (v9)
                doc_names = set()
                doc_data_rows = 0  # Count rows with actual data (not header)
                for ri in range(len(table.rows)):
                    row_text = ""
                    for c in [0, 1]:
                        if c < len(table.rows[ri].cells):
                            t = table.rows[ri].cells[c].text.strip()
                            if t:
                                doc_names.add(t)
                                if not row_text:
                                    row_text = t
                    if row_text and ri > 0:  # Skip header row
                        doc_data_rows += 1

                # v12: strip section tag for missing detection
                missing = [
                    n
                    for n in data
                    if not any(match_name(dn, n.split("__sec")[0]) for dn in doc_names)
                ]
                if missing:
                    # 分类漏填: 结构性 vs 真正遗漏
                    structural = []
                    genuine = []
                    SECTION_KW = (
                        "其中：",
                        "（一）",
                        "（二）",
                        "（三）",
                        "（四）",
                        "（五）",
                        "（六）",
                        "（七）",
                        "（八）",
                        "（九）",
                        "（十）",
                        "一、",
                        "二、",
                        "三、",
                        "四、",
                        "五、",
                        "六、",
                        "1年以内",
                        "1-2年",
                        "2-3年",
                        "3-4年",
                        "4-5年",
                        "5年以上",
                        "小计",
                        "补充资料",
                        "基本医疗保险",
                        "补充医疗保险",
                        "工伤保险费",
                        "生育保险费",
                        "企业年金",
                        "失业保险费",
                        "应收工程款",
                        "企业员工",
                        "进项税加计抵减",
                    )  # 模板定制差异
                    N_A_ITEMS = ("储备基金", "企业发展基金")
                    for m in missing:
                        if any(kw in m for kw in SECTION_KW):
                            structural.append(m)
                        elif any(na in m for na in N_A_ITEMS):
                            structural.append(f"{m}(可能不适用)")
                        else:
                            genuine.append(m)

                    # If template has significantly fewer data rows, it's summary vs detail
                    if doc_data_rows < len(data) * 0.5 and doc_data_rows > 0:
                        structural.extend(genuine)
                        genuine = []

                    pct = len(missing) / len(data) * 100
                    if structural and not genuine:
                        print(
                            f"  ℹ 结构性差异 {len(missing)}/{len(data)} (模板{doc_data_rows}数据行←套表{len(data)}行)"
                        )
                    elif genuine:
                        print(
                            f"  ⚠ 漏填 {len(genuine)}/{len(data)}: {[m[:25] for m in genuine[:3]]}"
                        )
                        missing_report[cat] = genuine
                    if pct >= 30 and genuine:
                        missing_report[cat] = genuine
            else:
                stats["no_match"] += 1
                print(f"  ⚠ 名称不匹配")
                doc_names = []
                for ri in range(1, min(5, len(table.rows))):
                    for c in [0, 1]:
                        if c < len(table.rows[ri].cells):
                            t = table.rows[ri].cells[c].text.strip()
                            if t:
                                doc_names.append(t)
                                break
                print(f"    Doc: {doc_names}")
                print(f"    TZ: {list(data.keys())[:4]}")
        except Exception as e:
            stats["err"] += 1
            print(f"  ❌ {e}")
            import traceback

            traceback.print_exc()

    print(f"\n{'=' * 70}")
    print(f"📊 {stats['ok']}/{len(MAPPINGS)} 项成功, {total_cells} 个单元格")
    print(
        f"   ⏭ Sheet未找到:{stats['no_sheet']} 无数据:{stats['no_data']} "
        f"表格超出:{stats['no_table']} 名称不匹配:{stats['no_match']} 错误:{stats['err']}"
    )
    if missing_report:
        print(
            f"🔴 需关注漏填: {sum(len(v) for v in missing_report.values())}行 ({len(missing_report)}项) — 非结构性差异"
        )
    else:
        print(f"✅ 无实质性漏填 (所有差异均为模板摘要vs套表明细的结构性差异)")

    # Post-processing: auto-sum totals and validate
    print(f"\n🧮 自动计算合计行...")
    total_sums = 0
    for ti in range(len(doc.tables)):
        s = auto_sum_totals(doc.tables[ti])
        total_sums += s
    if total_sums:
        print(f"  📊 自动填充 {total_sums} 个合计单元格 (已过滤'其中'子项)")

    print(f"\n💰 税种税率自动填充...")
    fill_tax_rates(doc)

    print(f"\n✏️ 修正字体格式...")
    fix_table_fonts(doc)

    print(f"\n🔍 加减验证...")
    validate_sums(doc)

    print(f"\n🔍 会计准则语义验证...")
    validate_data(doc)

    # v10: 会计准则业务规则说明
    print(f"\n📋 会计准则应用说明:")
    rules = [
        'CAS 22.47: 单项计提应收账款归入"3年以上"账龄段(已发生信用减值,视为长期逾期)',
        "CAS 22.48: 组合计提按预期信用损失(ECL)模型,比例=各账龄段预期损失率",
        'CAS 36: 关联方组合(37,124,436.53元)无账龄明细→应在"关联交易"附注单独披露',
        "CAS 22: 账面价值=账面余额-坏账准备,金融资产按摊余成本列示",
        "CAS 9: 应付职工薪酬=短期薪酬+离职后福利+辞退福利+其他",
    ]
    for rule in rules:
        print(f"  • {rule}")

    # v10: 主表稽核
    print(f"\n📊 四大主表稽核...")
    bs_checks = reconcile_with_main_statements(doc, taozhang_path)
    if bs_checks:
        print(f"  ✅ 主表稽核通过: {bs_checks}项")
    else:
        print(f"  (主表数据未提取或无可比项)")

    print(f"\n💾 保存到: {output_path}")
    doc.save(output_path)


if __name__ == "__main__":
    fill_notes(
        "D:/Users/12844/Desktop/保定年审资料/2、决算套表/25年决算套表 (2)(1)/25年决算套表/保定吉达电力设计有限公司-单体.xlsx",
        "D:/Users/12844/Desktop/自动化填模板/已下发的文档模版/2.审计文档_docx/1.审计报告/1-2单体审计报告模板/4.2025年财务报表附注模板-单户.docx",
        "E:/Projects/my-workspace/_work/附注_单体_v11_正式版.docx",
    )
