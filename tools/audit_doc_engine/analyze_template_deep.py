"""
模板深度解读 — 从T5开始，逐科目分析每个表格的含义、披露目的、对应会计准则
"""

import sys

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document

doc = Document(
    "D:/Users/12844/Desktop/自动化填模板/已下发的文档模版/2.审计文档_docx/1.审计报告/1-2单体审计报告模板/4.2025年财务报表附注模板-单户.docx"
)

# 逐表分析
for ti in range(5, min(55, len(doc.tables))):
    t = doc.tables[ti]

    # 提取表头行
    header = []
    if len(t.rows) > 0:
        for ci in range(min(8, len(t.rows[0].cells))):
            txt = t.rows[0].cells[ci].text.strip()[:30]
            if txt:
                header.append(txt)

    # 提取第一列名称（通常是科目项目名）
    first_col = []
    for ri in range(1, min(8, len(t.rows))):
        txt = t.rows[ri].cells[0].text.strip()[:40]
        if txt:
            first_col.append(txt)

    # 提取行列数
    rows = len(t.rows)
    cols = len(t.columns)

    # 判断科目类别
    all_text = " ".join(header + first_col)

    print(f"\n{'=' * 70}")
    print(f"【T{ti}】({rows}行×{cols}列)")
    print(f"  表头: {header}")
    print(f"  首列: {first_col}")

    # 简要解读
    if "库存现金" in all_text or "银行存款" in all_text:
        print(f"  📖 科目: 货币资金 — 按现金/银行存款/其他货币资金分类列示期末期初余额")
        print(f"  📖 准则: CAS 30 应用指南 — 货币资金")
        print(f"  📖 披露目的: 展示企业最基础的流动性资产构成")
    elif "银行承兑汇票" in all_text:
        if "坏账准备" in all_text:
            print(f"  📖 科目: 应收票据 — 按票据种类列账面余额/坏账准备/账面价值")
            print(f"  📖 准则: CAS 22 — 金融工具：应收票据")
        elif "已质押" in all_text:
            print(f"  📖 科目: 应收票据-已质押金额")
        elif "终止确认" in all_text:
            print(f"  📖 科目: 应收票据-已背书未到期")
        elif "转应收账款" in all_text:
            print(f"  📖 科目: 应收票据-出票人未履约转应收")
        else:
            print(f"  📖 科目: 应收票据")
    elif "账龄" in all_text and ("1年以内" in first_col[0] if first_col else False):
        if "坏账准备" in all_text:
            print(f"  📖 科目: 应收账款-账龄分析 — 按账龄段披露账面余额和坏账准备")
            print(f"  📖 准则: CAS 22 应用指南 — 预期信用损失(ECL)模型")
            print(f"  📖 披露目的: 展示应收账款账龄结构和信用风险分布")
        else:
            print(f"  📖 科目: 预付款项-账龄分析")
    elif "类别" in header and ("账面余额" in all_text or "坏账准备" in all_text):
        if "单项计提" in first_col[0] if first_col else False:
            print(f"  📖 科目: 应收账款-按计提方式分类")
        else:
            print(f"  📖 科目: 应收账款/其他应收款-组合分类")
    elif "债务人名称" in header and "比例" in all_text:
        print(f"  📖 科目: 应收账款-前五名")
    elif "预付款项" in all_text and "账龄" in header:
        print(f"  📖 科目: 预付款项-账龄分析")
    elif "应收资金集中管理款" in all_text:
        print(f"  📖 科目: 应收资金集中管理款（集团资金归集）")
    elif "应收利息" in all_text or "应收股利" in all_text:
        print(f"  📖 科目: 其他应收款-应收利息/应收股利")
    elif "坏账准备" in header and "第一阶段" in all_text:
        print(f"  📖 科目: 其他应收款-坏账准备变动（三阶段ECL模型）")
    elif "应收分保" in all_text:
        print(f"  📖 科目: 应收分保账款（保险公司专用）")
    elif "拆放" in all_text:
        print(f"  📖 科目: 拆出资金（银行专用）")
    elif "定期存款" in all_text or "委托贷款" in all_text:
        print(f"  📖 科目: 其他流动资产")
    elif "租赁" in first_col[0] if first_col else False:
        print(f"  📖 科目: 应收融资租赁款")
    elif "合同资产" in all_text and "减值" not in all_text:
        print(f"  📖 科目: 合同资产")
    elif "持有待售" in all_text:
        print(f"  📖 科目: 持有待售资产")
    elif "原材料" in all_text or "库存商品" in all_text or "跌价准备" in all_text:
        print(f"  📖 科目: 存货")
    elif "数据资源" in all_text:
        print(f"  📖 科目: 数据资源存货（2024新增）")
    elif "债权" in all_text and "投资" in all_text and "面值" not in all_text:
        print(f"  📖 科目: 债权投资")
    elif "其他债权投资" in all_text:
        print(f"  📖 科目: 其他债权投资")
    elif (
        "长期股权投资" in all_text
        or "对子公司投资" in all_text
        or "对合营企业" in all_text
    ):
        print(f"  📖 科目: 长期股权投资")
    elif "固定资产" in all_text and "账面原值" in all_text:
        print(f"  📖 科目: 固定资产-增减变动（原值/折旧/减值/账面价值）")
    elif "在建工程" in all_text:
        if "减值" in all_text:
            print(f"  📖 科目: 在建工程-减值情况")
        elif "工程名称" in header[0] if header else False:
            print(f"  📖 科目: 在建工程-明细项目")
        else:
            print(f"  📖 科目: 在建工程-账面情况")
    elif "使用权资产" in all_text:
        print(f"  📖 科目: 使用权资产（CAS 21 租赁）")
    elif "无形资产" in all_text:
        print(f"  📖 科目: 无形资产-增减变动")
    elif "长期待摊" in all_text:
        print(f"  📖 科目: 长期待摊费用")
    elif "递延所得税" in all_text:
        print(f"  📖 科目: 递延所得税资产/负债")
    elif "其他非流动资产" in all_text:
        print(f"  📖 科目: 其他非流动资产")
    elif "短期借款" in all_text or "质押借款" in all_text or "抵押借款" in all_text:
        print(f"  📖 科目: 短期借款")
    elif "应付票据" in all_text:
        print(f"  📖 科目: 应付票据")
    elif "应付账款" in all_text and "账龄" in header:
        print(f"  📖 科目: 应付账款-账龄分析")
    elif "应付职工薪酬" in all_text and "总表" not in all_text:
        if "短期薪酬" in all_text:
            pass  # 不重复
        else:
            print(f"  📖 科目: 应付职工薪酬")
    elif "应交税费" in all_text:
        print(f"  📖 科目: 应交税费")
    elif "应付利息" in all_text or "应付股利" in all_text:
        print(f"  📖 科目: 其他应付款-应付利息/应付股利")
    elif "长期借款" in all_text:
        print(f"  📖 科目: 长期借款")
    elif "应付债券" in all_text:
        print(f"  📖 科目: 应付债券")
    elif "预计负债" in all_text:
        print(f"  📖 科目: 预计负债")
    elif "递延收益" in all_text:
        print(f"  📖 科目: 递延收益")
    elif "实收资本" in all_text:
        print(f"  📖 科目: 实收资本")
    elif "资本公积" in all_text:
        print(f"  📖 科目: 资本公积")
    elif "盈余公积" in all_text:
        print(f"  📖 科目: 盈余公积")
    elif "未分配利润" in all_text:
        print(f"  📖 科目: 未分配利润")
    elif "营业收入" in all_text and "营业成本" in all_text:
        print(f"  📖 科目: 营业收入/营业成本")
    elif "税金及附加" in all_text:
        print(f"  📖 科目: 税金及附加")
    elif "销售费用" in all_text or "管理费用" in all_text:
        print(f"  📖 科目: 销售费用/管理费用")
    elif "研发费用" in all_text:
        print(f"  📖 科目: 研发费用")
    elif "财务费用" in all_text:
        print(f"  📖 科目: 财务费用")
    elif "其他收益" in all_text:
        print(f"  📖 科目: 其他收益（政府补助等）")
    elif "投资收益" in all_text:
        print(f"  📖 科目: 投资收益")
    elif "信用减值损失" in all_text:
        print(f"  📖 科目: 信用减值损失（ECL模型下的坏账损失）")
    elif "资产减值损失" in all_text:
        print(f"  📖 科目: 资产减值损失")
    elif "营业外收入" in all_text:
        print(f"  📖 科目: 营业外收入")
    elif "营业外支出" in all_text:
        print(f"  📖 科目: 营业外支出")
    elif "所得税费用" in all_text:
        print(f"  📖 科目: 所得税费用")
    elif "现金流量表" in all_text and "间接法" not in all_text:
        print(f"  📖 科目: 现金流量表补充资料")
    elif "间接法" in all_text:
        print(f"  📖 科目: 现金流量表-间接法（将净利润调节为经营活动现金流量）")
    elif "现金" in all_text and "等价物" in all_text:
        print(f"  📖 科目: 现金和现金等价物构成")
    else:
        print(f"  📖 待识别")
