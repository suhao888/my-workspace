"""
匹配诊断脚本 — 检查T91(固定资产情况表)和T45(其他应收款账龄)的name matching
"""

import sys

sys.stdout.reconfigure(encoding="utf-8")
sys.path.insert(0, "E:/Projects/my-workspace/tools/audit_doc_engine")

import re, pandas as pd
from docx import Document
from fill_notes import match_name, extract_tz_data, MAPPINGS

taozhang_path = "D:/Users/12844/Desktop/保定年审资料/2、决算套表/25年决算套表 (2)(1)/25年决算套表/保定吉达电力设计有限公司-单体.xlsx"
xl = pd.ExcelFile(taozhang_path)

# =============== T91 固定资产情况表 ===============
# 查找套表 固定资产情况
sheet_raw = None
for s in xl.sheet_names:
    try:
        d = s.encode("latin1").decode("gbk")
    except:
        d = s
    if "固定资产情况_原始数据" in d:
        sheet_raw = s
        break

if sheet_raw:
    df = pd.read_excel(taozhang_path, sheet_name=sheet_raw, header=None)
    # 提取taozhang中的名称
    tz_names = {}
    name_col = 1
    for i in range(df.shape[0]):
        name = (
            str(df.iloc[i, name_col]).strip() if pd.notna(df.iloc[i, name_col]) else ""
        )
        if name and name != "nan":
            rows = [
                str(df.iloc[i, j])[:15] if pd.notna(df.iloc[i, j]) else ""
                for j in range(min(10, df.shape[1]))
            ]
            tz_names[i] = (name, rows)

    # 输出套表行
    print("=== 套表 固定资产情况_原始数据 ===")
    for i, (name, rows) in sorted(tz_names.items()):
        print(f"  R{i}: [{name[:25]}] 数据:{rows}")
    print()

    # 读取Docx模板T91
    doc = Document("D:/Users/12844/Desktop/审计输出_测试/附注_单体_v11_测试.docx")
    t = doc.tables[91]

    print("=== T91 模板行 vs 套表匹配 ===")
    for ri in range(len(t.rows)):
        doc_name = ""
        for c in [0, 1]:
            if c < len(t.rows[ri].cells):
                txt = t.rows[ri].cells[c].text.strip()
                if txt:
                    doc_name = txt
                    break
        if not doc_name:
            continue

        best_match = None
        for tz_i, (tz_name, _) in tz_names.items():
            if match_name(doc_name, tz_name):
                best_match = (tz_i, tz_name)

        if best_match:
            match_name_clean = best_match[1][:30]
            print(f"  R{ri}[{doc_name[:25]}] → 套R{best_match[0]}[{match_name_clean}]")
        else:
            print(f"  R{ri}[{doc_name[:25]}] → 无匹配")

# =============== T45 其他应收款-账龄 ===============
print()
for s in xl.sheet_names:
    try:
        d = s.encode("latin1").decode("gbk")
    except:
        d = s
    if "按账龄披露其他应收" in d:
        df = pd.read_excel(taozhang_path, sheet_name=s, header=None)
        tz_names2 = {}
        for i in range(df.shape[0]):
            name = str(df.iloc[i, 1]).strip() if pd.notna(df.iloc[i, 1]) else ""
            if name and name != "nan" and not name.startswith("注"):
                rows = [
                    str(df.iloc[i, j])[:15] if pd.notna(df.iloc[i, j]) else ""
                    for j in range(min(12, df.shape[1]))
                ]
                tz_names2[i] = (name, rows)

        print("=== 套表 按账龄披露其他应收款 ===")
        for i, (name, rows) in sorted(tz_names2.items()):
            print(f"  R{i}: [{name[:25]}] 数据:{rows}")
        print()

        t2 = doc.tables[45]
        print("=== T45 模板行 vs 套表匹配 ===")
        for ri in range(len(t2.rows)):
            doc_name = ""
            for c in [0, 1]:
                if c < len(t2.rows[ri].cells):
                    txt = t2.rows[ri].cells[c].text.strip()
                    if txt:
                        doc_name = txt
                        break
            if not doc_name:
                continue

            best_match = None
            for tz_i, (tz_name, _) in tz_names2.items():
                if match_name(doc_name, tz_name):
                    best_match = (tz_i, tz_name)

            if best_match:
                print(
                    f"  R{ri}[{doc_name[:25]}] → 套R{best_match[0]}[{best_match[1][:30]}]"
                )
            else:
                print(f"  R{ri}[{doc_name[:25]}] → 无匹配")
        break
