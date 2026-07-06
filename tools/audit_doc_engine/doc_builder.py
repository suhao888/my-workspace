"""
DocBuilder — 简洁文档排版引擎
=============================
基于 python-docx 封装，输出符合日常排版习惯的中文文档。

功能:
  - 封面页（标题/副标题/日期/密级）
  - 页眉/页脚（页码、第X页共Y页）
  - 目录域（Word 中按 Ctrl+A → F9 更新）
  - 样式体系（多级标题/正文/列表）
  - 标准表格（浅色表头、列宽控制）
  - 键值对表、提示框、签名块
  - 中英文字体自动设置（宋体+Times New Roman）
  - 所有文字默认黑色，版面简洁

用法:
    builder = DocBuilder()
    builder.setup_page().setup_styles()
    builder.add_heading_1("第一章")
    builder.add_body("正文内容...")
    builder.add_table(headers=["列1","列2"], rows=[["a","b"]])
    builder.save("output.docx")
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional, Dict, List, Tuple

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# ============================================================
# 配色 — 全部黑色，保持简洁
# ============================================================
COLOR_BLACK = "#000000"
COLOR_LIGHT_BG = "#F2F2F2"  # 浅灰 — 表头背景
COLOR_BORDER = "#BFBFBF"  # 边框线

# ============================================================
# 默认字体
# ============================================================
FONT_CN_BODY = "宋体"  # 正文中文字体
FONT_CN_HEADING = "黑体"  # 标题中文字体
FONT_EN = "Times New Roman"  # 英文/数字字体

# ============================================================
# 页面尺寸
# ============================================================
PAGE_A4 = (Cm(21.0), Cm(29.7))

# ============================================================
# XML 工具函数
# ============================================================


def _hex_to_rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_run_font_xml(
    r, text, font_cn, font_en, size_pt, bold=False, color_hex=COLOR_BLACK
):
    """通过 XML 设置 run 的字体属性，避免 python-docx API 副作用。"""
    rPr = OxmlElement("w:rPr")
    # 字体
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_en)
    rFonts.set(qn("w:hAnsi"), font_en)
    rFonts.set(qn("w:eastAsia"), font_cn)
    rPr.append(rFonts)
    # 字号（单位：half-point）
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(szCs)
    # 颜色
    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), color_hex.lstrip("#"))
    rPr.append(clr)
    # 加粗
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    r.insert(0, rPr)
    # 文字内容
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)


def _set_cell_shading(cell, color_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:shd"))
    if old is not None:
        tcPr.remove(old)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def _set_cell_vcenter(cell):
    """通过 XML 设置单元格上下居中，不受 python-docx API 缓存影响。"""
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:vAlign"))
    if old is not None:
        tcPr.remove(old)
    el = OxmlElement("w:vAlign")
    el.set(qn("w:val"), "center")
    tcPr.append(el)


def _set_cell_center(cell):
    """通过 XML 设置单元格文字左右居中。"""
    p = cell.paragraphs[0]
    pPr = p._p.get_or_add_pPr()
    old = pPr.find(qn("w:jc"))
    if old is not None:
        pPr.remove(old)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)


def _set_para_border(para, edge: str, sz: int = 4, color: str = COLOR_BORDER):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), "1")
    el.set(qn("w:color"), color)
    pBdr.append(el)
    pPr.append(pBdr)


def _add_field(para, field_code: str):
    """插入 Word 域代码（PAGE / NUMPAGES / TOC 等）"""
    for typ in ("begin", "separate", "end"):
        r = para.add_run()
        f = OxmlElement("w:fldChar")
        f.set(qn("w:fldCharType"), typ)
        r._r.append(f)
        if typ == "separate":
            r2 = para.add_run("1")
            r2.font.size = Pt(9)
            r2.font.color.rgb = _hex_to_rgb(COLOR_BLACK)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    # insert instrText after begin
    para.runs[0]._r.addnext(instr)


def _set_run_font(
    run_or_font,
    cn_font: str = FONT_CN_BODY,
    en_font: str = FONT_EN,
    size: int = 12,
    bold: bool = False,
    color_hex: Optional[str] = None,
):
    """Set font on a Run or a style Font object.

    同时支持 python-docx 的 Run（有 .font 属性）和 style Font（无 .font 但自身就是 Font）。
    中文字体通过 rFonts.eastAsia 写入。
    """
    if hasattr(run_or_font, "font"):
        run = run_or_font
        run.font.name = en_font
        run.font.size = Pt(size)
        run.font.bold = bold
        if color_hex:
            run.font.color.rgb = _hex_to_rgb(color_hex)
        rPr = run._element.get_or_add_rPr()
    else:
        f = run_or_font
        f.name = en_font
        f.size = Pt(size)
        f.bold = bold
        if color_hex:
            f.color.rgb = _hex_to_rgb(color_hex)
        rPr = f._element  # style Font._element 已经是 rPr
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)


def _remove_table_borders(table):
    tblPr = table._tbl.tblPr
    xml = (
        f"<w:tblBorders {nsdecls('w')}>"
        + "".join(
            f'<w:{e} w:val="none" w:sz="0" w:space="0"/>'
            for e in ("top", "left", "bottom", "right", "insideH", "insideV")
        )
        + "</w:tblBorders>"
    )
    tblPr.append(parse_xml(xml))


def _set_table_borders_black(table, sz: int = 4):
    """将表格所有边框设为黑色实线。

    Args:
        table: python-docx Table 对象
        sz: 线宽（1/8 磅单位，4 = 0.5pt）
    """
    tblPr = table._tbl.tblPr
    # 移除旧的 tblBorders（如果有）
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    xml = (
        f"<w:tblBorders {nsdecls('w')}>"
        + "".join(
            f'<w:{e} w:val="single" w:sz="{sz}" w:space="0" w:color="000000"/>'
            for e in ("top", "left", "bottom", "right", "insideH", "insideV")
        )
        + "</w:tblBorders>"
    )
    tblPr.append(parse_xml(xml))


def _disable_table_style_look(table):
    """关闭 w:tblLook 的条件格式应用，防止样式覆盖显式单元格设置。"""
    tblPr = table._tbl.tblPr
    look = tblPr.find(qn("w:tblLook"))
    if look is not None:
        look.set(qn("w:firstRow"), "0")
        look.set(qn("w:firstColumn"), "0")
        look.set(qn("w:lastRow"), "0")
        look.set(qn("w:lastColumn"), "0")
        look.set(qn("w:noHBand"), "1")
        look.set(qn("w:noVBand"), "1")


def _set_table_width(table, width_cm: float):
    """设置表格整体宽度（dxa 单位）。

    Args:
        table: python-docx Table 对象
        width_cm: 宽度（厘米）
    """
    tblPr = table._tbl.tblPr
    # 移除旧的 tblW
    old = tblPr.find(qn("w:tblW"))
    if old is not None:
        tblPr.remove(old)
    dxa = int(width_cm * 567)  # 1cm ≈ 567 DXA
    xml = f'<w:tblW {nsdecls("w")} w:w="{dxa}" w:type="dxa"/>'
    tblPr.append(parse_xml(xml))


# ============================================================
# DocBuilder
# ============================================================


class DocBuilder:
    """专业文档排版构建器。"""

    def __init__(self):
        self.doc = Document()
        self._style_body = "_db_body"
        self._style_h1 = "_db_h1"
        self._style_h2 = "_db_h2"
        self._style_h3 = "_db_h3"

    # ── 页面设置 ────────────────────────────────────────────────

    def setup_page(self, margins: Optional[Dict[str, float]] = None):
        """设置 A4 页面和边距。"""
        s = self.doc.sections[0]
        s.page_width, s.page_height = PAGE_A4
        m = margins or {"left": 3.0, "right": 2.5, "top": 2.54, "bottom": 2.54}
        for k in ("left", "right", "top", "bottom"):
            setattr(s, f"{k}_margin", Cm(m[k]))
        return self

    # ── 样式体系 ────────────────────────────────────────────────

    def setup_styles(self, body_size: int = 12, line_spacing: float = 1.5):
        """建立文档样式体系。

        Args:
            body_size: 正文字号（磅）
            line_spacing: 行距倍数
        """
        styles = self.doc.styles

        # Normal
        n = styles["Normal"]
        n.font.name = FONT_EN
        n.font.size = Pt(body_size)
        n.font.color.rgb = _hex_to_rgb(COLOR_BLACK)
        rPr = n.element.get_or_add_rPr()
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:eastAsia"), FONT_CN_BODY)
        rPr.insert(0, rf)
        pf = n.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)

        # 标题样式
        for name, sz, bef, aft in [
            (self._style_h1, 16, 24, 12),
            (self._style_h2, 14, 18, 8),
            (self._style_h3, 12, 12, 6),
        ]:
            st = styles.add_style(name, 1)
            st.font.name = FONT_EN
            st.font.size = Pt(sz)
            st.font.bold = True
            st.font.color.rgb = _hex_to_rgb(COLOR_BLACK)
            rp = st.element.get_or_add_rPr()
            rf = OxmlElement("w:rFonts")
            rf.set(qn("w:eastAsia"), FONT_CN_HEADING)
            rp.insert(0, rf)
            ppf = st.paragraph_format
            ppf.space_before = Pt(bef)
            ppf.space_after = Pt(aft)
            ppf.keep_with_next = True
            ppf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            ppf.line_spacing = line_spacing

        # 正文样式（首行缩进）
        bst = styles.add_style(self._style_body, 1)
        _set_run_font(bst.font, FONT_CN_BODY, FONT_EN, body_size, color_hex=COLOR_BLACK)
        bpf = bst.paragraph_format
        bpf.first_line_indent = Pt(body_size * 2)
        bpf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        bpf.line_spacing = line_spacing
        bpf.space_after = Pt(2)

        # 列表样式
        for ln in ("List Bullet", "List Number"):
            try:
                ls = styles[ln]
                _set_run_font(ls.font, FONT_CN_BODY, FONT_EN, body_size)
                ls.paragraph_format.line_spacing = line_spacing
            except KeyError:
                pass

        return self

    # ── 封面页 ──────────────────────────────────────────────────

    def add_cover(
        self,
        title: str,
        subtitle: str = "",
        date_text: str = "",
        confidential: str = "",
    ):
        """生成封面页。

        Args:
            title: 文档主标题
            subtitle: 副标题/公司名称
            date_text: 日期
            confidential: 密级（如 "内部资料"）
        """
        if confidential:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_run_font(
                p.add_run(confidential),
                FONT_CN_BODY,
                FONT_EN,
                11,
                color_hex=COLOR_BLACK,
            )
            p.paragraph_format.space_after = Pt(80)
        else:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(100)

        if subtitle:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run_font(
                p.add_run(subtitle),
                FONT_CN_HEADING,
                FONT_EN,
                18,
                color_hex=COLOR_BLACK,
            )
            p.paragraph_format.space_after = Pt(60)

        # 主标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(
            p.add_run(title),
            FONT_CN_HEADING,
            FONT_EN,
            22,
            bold=True,
            color_hex=COLOR_BLACK,
        )
        p.paragraph_format.space_after = Pt(60)

        if date_text:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run_font(
                p.add_run(date_text), FONT_CN_BODY, FONT_EN, 14, color_hex=COLOR_BLACK
            )

        self.doc.add_page_break()
        return self

    # ── 目录 ────────────────────────────────────────────────────

    def add_toc(self, title: str = "目  录", levels: int = 3):
        """插入目录域（在 Word 中按 Ctrl+A → F9 更新）。"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(
            p.add_run(title),
            FONT_CN_HEADING,
            FONT_EN,
            16,
            bold=True,
            color_hex=COLOR_BLACK,
        )
        p.paragraph_format.space_after = Pt(20)
        _set_para_border(p, "bottom", 4, COLOR_BLACK)

        pt = self.doc.add_paragraph()
        _add_field(pt, f'TOC \\o "1-{levels}" \\h \\z \\u')

        ph = self.doc.add_paragraph()
        _set_run_font(
            ph.add_run("（在 Word 中按 Ctrl+A → F9 更新目录）"),
            FONT_CN_BODY,
            FONT_EN,
            9,
            color_hex=COLOR_BLACK,
        )
        ph.paragraph_format.space_before = Pt(6)

        self.doc.add_page_break()
        return self

    # ── 页眉页脚 ────────────────────────────────────────────────

    def setup_header(self, text: str = ""):
        """设置页眉（首页无页眉）"""
        sec = self.doc.sections[0]
        sec.different_first_page_header_footer = True
        header = sec.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p._p.getparent().remove(p._p)
        p = header.add_paragraph()
        if text:
            _set_run_font(
                p.add_run(text), FONT_CN_BODY, FONT_EN, 9, color_hex=COLOR_BLACK
            )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_para_border(p, "bottom", 4, COLOR_BORDER)
        return self

    def setup_footer(self, left: str = "", right: str = "page_number"):
        """设置页脚。

        Args:
            left: 左侧文本
            right: "page_number"=页码 | "page_x_of_y"=第X页共Y页 | 自定义文本
        """
        sec = self.doc.sections[0]
        sec.different_first_page_header_footer = True
        footer = sec.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p._p.getparent().remove(p._p)

        tw = sec.page_width - sec.left_margin - sec.right_margin
        p = footer.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(tw, WD_TAB_ALIGNMENT.RIGHT)

        if left:
            _set_run_font(
                p.add_run(left), FONT_CN_BODY, FONT_EN, 9, color_hex=COLOR_BLACK
            )
        r = p.add_run()
        r.add_tab()

        if right == "page_number":
            _add_field(p, "PAGE")
        elif right == "page_x_of_y":
            _add_field(p, "PAGE")
            rr = p.add_run(" / ")
            _set_run_font(rr, FONT_EN, FONT_EN, 9, color_hex=COLOR_BLACK)
            _add_field(p, "NUMPAGES")
        else:
            _set_run_font(
                p.add_run(right), FONT_CN_BODY, FONT_EN, 9, color_hex=COLOR_BLACK
            )

        _set_para_border(p, "top", 4, COLOR_BORDER)
        return self

    # ── 内容元素 ────────────────────────────────────────────────

    def add_heading_1(self, text: str, bold: bool = True):
        """一级标题"""
        p = self.doc.add_paragraph(style=self._style_h1)
        _set_run_font(
            p.add_run(text),
            FONT_CN_HEADING,
            FONT_EN,
            16,
            bold=bold,
            color_hex=COLOR_BLACK,
        )
        return self

    def add_heading_2(self, text: str, bold: bool = True):
        """二级标题"""
        p = self.doc.add_paragraph(style=self._style_h2)
        _set_run_font(
            p.add_run(text),
            FONT_CN_HEADING,
            FONT_EN,
            14,
            bold=bold,
            color_hex=COLOR_BLACK,
        )
        return self

    def add_heading_3(self, text: str, bold: bool = True):
        """三级标题"""
        p = self.doc.add_paragraph(style=self._style_h3)
        _set_run_font(
            p.add_run(text),
            FONT_CN_HEADING,
            FONT_EN,
            12,
            bold=bold,
            color_hex=COLOR_BLACK,
        )
        return self

    def add_body(self, text: str):
        """正文段落（首行缩进）"""
        if not text:
            return self
        p = self.doc.add_paragraph(style=self._style_body)
        _set_run_font(p.add_run(text), FONT_CN_BODY, FONT_EN, 12, color_hex=COLOR_BLACK)
        return self

    def add_body_no_indent(self, text: str):
        """正文段落（无缩进）"""
        if not text:
            return self
        p = self.doc.add_paragraph()
        _set_run_font(p.add_run(text), FONT_CN_BODY, FONT_EN, 12, color_hex=COLOR_BLACK)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        return self

    def add_list_item(self, text: str, level: int = 0):
        """列表项"""
        style = "List Bullet" if level == 0 else "List Bullet 2"
        p = self.doc.add_paragraph(style=style)
        _set_run_font(p.add_run(text), FONT_CN_BODY, FONT_EN, 12, color_hex=COLOR_BLACK)
        return self

    def add_space(self, pt: float = 12):
        """空行"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(pt)
        _set_run_font(p.add_run(""), FONT_CN_BODY, FONT_EN, pt)
        return self

    def add_page_break(self):
        """分页"""
        self.doc.add_page_break()
        return self

    def add_note_box(self, text: str):
        """带灰色边框的提示段落。"""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        _set_para_border(p, "top", 4, COLOR_BORDER)
        _set_para_border(p, "bottom", 4, COLOR_BORDER)
        _set_run_font(p.add_run(text), FONT_CN_BODY, FONT_EN, 12, color_hex=COLOR_BLACK)
        return self

    def add_signature(self, items: List[Tuple[str, str]]):
        """签名块。items = [(标签, 值), ...]"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(36)
        for i, (label, value) in enumerate(items):
            _set_run_font(
                p.add_run(f"{label}：{value}"),
                FONT_CN_BODY,
                FONT_EN,
                12,
                color_hex=COLOR_BLACK,
            )
            if i < len(items) - 1:
                p.add_run("\n")
        p.paragraph_format.line_spacing = 2.0
        return self

    # ── 表格 ────────────────────────────────────────────────────

    def add_table(
        self,
        headers: List[str],
        rows: List[List],
        col_widths: Optional[List[float]] = None,
        caption: str = "",
    ):
        """创建标准表格（浅灰表头，全黑边框，内容居中）。

        Args:
            headers: 表头
            rows: 数据行
            col_widths: 列宽比例，如 [2,3,2]
            caption: 表题
        """
        if caption:
            p = self.doc.add_paragraph()
            _set_run_font(
                p.add_run(caption),
                FONT_CN_HEADING,
                FONT_EN,
                11,
                bold=False,
                color_hex=COLOR_BLACK,
            )
            p.paragraph_format.space_after = Pt(4)

        ncols = len(headers)
        table = self.doc.add_table(rows=len(rows) + 1, cols=ncols)
        table.style = "Table Grid"
        _disable_table_style_look(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_table_borders_black(table)

        # 表头（黑色文字，上下左右居中，浅灰底色 #F2F2F2）
        for j, h in enumerate(headers):
            c = table.rows[0].cells[j]
            _set_cell_shading(c, COLOR_LIGHT_BG)
            # 清空 cell 内容后直接 XML 操作，绕过 python-docx 缓存问题
            tc = c._tc
            for child in list(tc):
                tag = child.tag.split("}")[-1]
                if tag != "w:tcPr":
                    tc.remove(child)
            p = OxmlElement("w:p")
            tc.append(p)
            r = OxmlElement("w:r")
            p.append(r)
            # 设置文字
            _set_run_font_xml(r, h, FONT_CN_BODY, FONT_EN, 10.5, color_hex=COLOR_BLACK)
            _set_cell_vcenter(c)
            _set_cell_center(c)

        # 数据行（居中，上下居中，纯黑色，无底纹）
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                c = table.rows[i + 1].cells[j]
                tc = c._tc
                for child in list(tc):
                    tag = child.tag.split("}")[-1]
                    if tag != "w:tcPr":
                        tc.remove(child)
                p = OxmlElement("w:p")
                tc.append(p)
                r = OxmlElement("w:r")
                p.append(r)
                _set_run_font_xml(
                    r, str(val), FONT_CN_BODY, FONT_EN, 10.5, color_hex=COLOR_BLACK
                )
                _set_cell_vcenter(c)
                _set_cell_center(c)

        # 列宽 + 固定总宽
        if col_widths and len(col_widths) == ncols:
            total = sum(col_widths)
            for j, w in enumerate(col_widths):
                for row in table.rows:
                    row.cells[j].width = Cm(15.0 * w / total)
        _set_table_width(table, 15.0)

        self.add_space(6)
        return self

    def add_kv_table(
        self,
        items: List[Tuple[str, str]],
        label_width: float = 3.0,
        total_width: float = 15.0,
    ):
        """键值对表格（全黑边框，标签列浅灰背景，内容居中）。

        Args:
            items: 键值对列表
            label_width: 标签列宽（cm）
            total_width: 表格总宽（cm），默认 15.0 以匹配 add_table 的总宽
        """
        table = self.doc.add_table(rows=len(items), cols=2)
        table.style = "Table Grid"
        _disable_table_style_look(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_table_borders_black(table)
        val_width = max(total_width - label_width, 1.0)
        for i, (k, v) in enumerate(items):
            ck = table.rows[i].cells[0]
            tc = ck._tc
            for child in list(tc):
                tag = child.tag.split("}")[-1]
                if tag != "w:tcPr":
                    tc.remove(child)
            p = OxmlElement("w:p")
            tc.append(p)
            r = OxmlElement("w:r")
            p.append(r)
            _set_run_font_xml(r, k, FONT_CN_BODY, FONT_EN, 10.5, color_hex=COLOR_BLACK)
            _set_cell_vcenter(ck)
            _set_cell_center(ck)
            ck.width = Cm(label_width)
            cv = table.rows[i].cells[1]
            tc2 = cv._tc
            for child in list(tc2):
                tag = child.tag.split("}")[-1]
                if tag != "w:tcPr":
                    tc2.remove(child)
            p2 = OxmlElement("w:p")
            tc2.append(p2)
            r2 = OxmlElement("w:r")
            p2.append(r2)
            _set_run_font_xml(
                r2, str(v), FONT_CN_BODY, FONT_EN, 10.5, color_hex=COLOR_BLACK
            )
            _set_cell_vcenter(cv)
            _set_cell_center(cv)
            cv.width = Cm(val_width)
        _set_table_width(table, total_width)
        self.add_space(6)
        return self

    # ── 保存 ────────────────────────────────────────────────────

    def save(self, path: str):
        """保存文档。"""
        self.doc.save(str(path))
        return self
