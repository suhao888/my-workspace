"""
审计文档自动填充引擎 — 核心模块
Phase 1: 模板分析 + 实体管理 + 填充引擎框架
"""
import sys, os, json, re, shutil
from pathlib import Path
from collections import defaultdict
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple
import pandas as pd

sys.stdout.reconfigure(encoding='utf-8')


# ============================================================
# 数据模型
# ============================================================

@dataclass
class Entity:
    """审计实体（母公司/子公司/孙公司）"""
    name: str                          # 单位全称
    short_name: str = ''               # 简称
    uscc: str = ''                     # 统一社会信用代码
    level: str = '子公司'              # 母公司/子公司/孙公司
    is_consolidated: bool = False      # 是否合并主体
    parent: str = ''                   # 母公司名称
    registered_capital: str = ''       # 注册资本
    established_year: str = ''         # 成立年份
    address: str = ''                  # 注册地
    industry: str = ''                 # 行业
    business_scope: str = ''           # 经营范围
    controlling_shareholder: str = ''  # 控股股东
    ultimate_controller: str = ''      # 实际控制人
    org_structure: str = ''            # 组织结构简述
    legal_rep: str = ''                # 法定代表人
    report_approver: str = ''          # 财务报告批准报出者

    # 财务数据（从决算套表提取）
    financial_data: Dict[str, float] = field(default_factory=dict)

    @property
    def safe_name(self):
        return self.name.replace('/', '_').replace('\\', '_').replace('\n', '')


@dataclass
class Placeholder:
    """模板中的占位符"""
    text: str           # 原始占位符文本
    category: str       # 分类：公司信息/财务数据/日期/审计意见/其他
    data_source: str    # 数据来源字段名
    default: str = ''   # 无数据时的默认值


# ============================================================
# 模板分析器
# ============================================================

class TemplateAnalyzer:
    """分析Word模板，提取所有占位符"""

    PLACEHOLDER_PATTERNS = [
        # XX + 单位/公司/行业 等 → 公司信息
        (r'XX(?:公司|单位|集团|有限|行业|部|分公司)', '公司信息'),
        # XXXX年/月/日 → 日期
        (r'XXXX年(?:X|XX)?月?(?:X|XX)?日?', '日期'),
        # 注册资本XX万元 → 公司信息
        (r'注册资本(?:为)?XX(?:万元)?', '公司信息'),
        # XX出资/XX持股 → 公司信息
        (r'XX(?:出资|持股)', '公司信息'),
        # 统一社会信用代码XX → 公司信息
        (r'统一社会信用代码XX', '公司信息'),
        # XXX（按实际情况）→ 公司信息
        (r'XXX(?:（按.*?）)?', '公司信息'),
        # 纯XX → 需要人工判断
        (r'(?<![a-zA-Z一-龥])XX(?![a-zA-Z一-龥])', '需人工判断'),
    ]

    def analyze_docx(self, filepath: str) -> Dict:
        """分析一个docx模板，返回占位符映射"""
        from docx import Document
        doc = Document(filepath)

        result = {
            'file': filepath,
            'paragraph_placeholders': [],
            'table_placeholders': [],
            'total_tables': len(doc.tables),
            'total_paragraphs': len(doc.paragraphs),
        }

        # 分析段落
        for i, para in enumerate(doc.paragraphs):
            text = para.text
            for pattern, category in self.PLACEHOLDER_PATTERNS:
                matches = re.finditer(pattern, text)
                for m in matches:
                    result['paragraph_placeholders'].append({
                        'para_index': i,
                        'text': m.group(),
                        'category': category,
                        'context': text[max(0, m.start()-20):m.end()+20]
                    })

        # 分析表格
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    text = cell.text
                    for pattern, category in self.PLACEHOLDER_PATTERNS:
                        matches = re.finditer(pattern, text)
                        for m in matches:
                            result['table_placeholders'].append({
                                'table_index': ti,
                                'row': ri,
                                'col': ci,
                                'text': m.group(),
                                'category': category,
                                'context': text[max(0, m.start()-20):m.end()+20]
                            })

        return result

    def analyze_all_templates(self, template_dir: str) -> Dict:
        """分析所有模板文件"""
        base = Path(template_dir)
        all_results = {}

        for f in base.rglob('*.docx'):
            try:
                result = self.analyze_docx(str(f))
                rel_path = str(f.relative_to(base))
                all_results[rel_path] = result
                print(f'  分析: {rel_path} ({result["total_paragraphs"]}段, {result["total_tables"]}表)')
            except Exception as e:
                print(f'  错误: {f.name} - {e}')

        return all_results


# ============================================================
# 财务数据提取器
# ============================================================

class FinanceDataExtractor:
    """从决算套表提取科目余额和财务数据"""

    # 科目编码→标准科目名映射（企业会计准则）
    ACCOUNT_MAP = {
        # 资产类
        '1001': '现金',
        '1002': '银行存款',
        '1012': '其他货币资金',
        '1101': '交易性金融资产',
        '1121': '应收票据',
        '1122': '应收账款',
        '1123': '预付账款',
        '1131': '应收股利',
        '1132': '应收利息',
        '1221': '其他应收款',
        '1231': '坏账准备',
        '1401': '材料采购',
        '1402': '在途物资',
        '1403': '原材料',
        '1405': '库存商品',
        '1471': '存货跌价准备',
        '1501': '持有至到期投资',
        '1511': '长期股权投资',
        '1521': '投资性房地产',
        '1601': '固定资产',
        '1602': '累计折旧',
        '1603': '固定资产减值准备',
        '1604': '在建工程',
        '1605': '工程物资',
        '1606': '固定资产清理',
        '1701': '无形资产',
        '1702': '累计摊销',
        '1703': '无形资产减值准备',
        '1711': '商誉',
        '1801': '长期待摊费用',
        '1811': '递延所得税资产',
        # 负债类
        '2001': '短期借款',
        '2201': '应付票据',
        '2202': '应付账款',
        '2203': '预收账款',
        '2211': '应付职工薪酬',
        '2221': '应交税费',
        '2231': '应付利息',
        '2232': '应付股利',
        '2241': '其他应付款',
        '2501': '长期借款',
        '2502': '应付债券',
        '2701': '长期应付款',
        '2901': '递延所得税负债',
        # 权益类
        '4001': '实收资本（股本）',
        '4002': '资本公积',
        '4101': '盈余公积',
        '4103': '本年利润',
        '4104': '利润分配',
        # 损益类
        '6001': '营业收入',
        '6401': '营业成本',
        '6403': '营业税金及附加',
        '6601': '销售费用',
        '6602': '管理费用',
        '6603': '财务费用',
    }

    def extract_from_taozhang(self, taozhang_path: str, entity_name: str) -> Dict:
        """从决算套表提取指定实体的财务数据"""
        # 决算套表通常是多Sheet Excel，每个Sheet对应一个报表
        xl = pd.ExcelFile(taozhang_path)

        result = {
            'entity_name': entity_name,
            'balance_sheet': {},     # 资产负债表
            'income_statement': {},  # 利润表
            'cash_flow': {},         # 现金流量表
            'account_balances': {},  # 科目余额
        }

        for sheet in xl.sheet_names:
            df = pd.read_excel(taozhang_path, sheet_name=sheet, header=None)
            # 根据Sheet名称判断报表类型
            if '负债' in sheet or '资产' in sheet:
                result['balance_sheet'][sheet] = df
            elif '利润' in sheet or '损益' in sheet:
                result['income_statement'][sheet] = df
            elif '现金' in sheet:
                result['cash_flow'][sheet] = df

        return result

    def get_account_balance(self, taozhang_data: Dict, account_codes: List[str]) -> float:
        """获取科目余额（多个科目汇总）"""
        total = 0.0
        for code in account_codes:
            total += taozhang_data.get('account_balances', {}).get(code, 0.0)
        return total


# ============================================================
# 文档填充引擎
# ============================================================

class DocFillEngine:
    """Word文档填充引擎 — 保持原格式，替换占位符"""

    def __init__(self):
        self.replacements = {}  # {占位符: 替换值}
        self.warnings = []      # 格式问题警告

    def register_entity(self, entity: Entity):
        """注册实体信息，生成替换字典"""
        r = self.replacements
        s = entity  # shorthand

        # === 公司名称（多种模式） ===
        r['XX公司'] = s.name
        r['XX有限公司'] = s.name
        r['XX集团'] = s.name
        r['XX单位'] = s.name
        r['国家电网XXXX有限公司'] = s.name  # 4个X匹配
        # 模板中的占位符公司名
        r['【ABC公司】'] = s.name
        r['ABC公司'] = s.name
        r['【XX公司】'] = s.name
        r['[ABC公司]'] = s.name

        # === 统一社会信用代码 ===
        uscc = s.uscc or '（待补充）'
        # 精确匹配：只替换明确标注为信用代码的占位符
        r['统一社会信用代码XXXXXXXXXXXXXXXXXX'] = f'统一社会信用代码{uscc}'
        r['统一社会信用代码XX'] = uscc
        # 不要用纯18位X做全局替换（会误匹配其他18字符内容）

        # === 注册资本 ===
        if s.registered_capital:
            r['注册资本为XX万元'] = f'注册资本为{entity.registered_capital}万元'
            r['注册资本XX万元'] = f'注册资本{s.registered_capital}万元'
        else:
            r['注册资本为XX万元'] = '注册资本为（待补充）万元'
            r['注册资本XX万元'] = '注册资本（待补充）万元'

        # === 成立年份 + 沿革（多种句式） ===
        year = s.established_year or '（待补充）'
        r['成立于XX年'] = f'成立于{year}年'
        r['成立于XX年X月X日'] = f'成立于{year}年（待补充）月（待补充）日'
        # 更通用的日期
        r['XX年X月X日'] = f'{year}年（待补充）月（待补充）日'

        r['原为XX'] = '原为（待补充，请提供历史沿革）'
        affiliated = s.controlling_shareholder or '（待补充）'
        r['现隶属于XX'] = f'现隶属于{affiliated}'

        # 组织形式的默认值
        r['为国有独资公司（或其他组织形式）'] = '为国有独资公司（或其他组织形式）'
        # 出资人—需要从决算套表或企业信息中获取
        r['由XX出资'] = '由（待补充出资人信息）出资'
        r['（简述各出资人及出资比例）'] = '（待补充各出资人及出资比例）'

        # === 法定代表人/注册地址 ===
        legal = getattr(s, 'legal_rep', '（待补充）')
        r['法定代表人：XX'] = f'法定代表人：{legal}'
        r['注册地址：XX'] = f'注册地址：{s.address}' if s.address else '注册地址：（待补充）'
        r['企业住所：XX'] = f'企业住所：{s.address}' if s.address else '企业住所：（待补充）'

        # === 行业/经营范围 ===
        r['XX行业'] = s.industry or '（待补充）'
        if s.business_scope:
            r['经营范围主要包括：XX'] = f'经营范围主要包括：{s.business_scope}'
            # 多个XX（经营范围可能有多项用逗号分隔）
            items = [it.strip() for it in s.business_scope.split('、') if it.strip()]
            for idx, item in enumerate(items[:5]):
                r[f'XX，XX，XX，XX，XX'] = s.business_scope  # 整体替换

        # === 股东/控制人 ===
        r['控股股东为XX'] = f'控股股东为{s.controlling_shareholder}' if s.controlling_shareholder else '控股股东为（待补充）'
        r['最终控制人为XX'] = f'最终控制人为{s.ultimate_controller}' if s.ultimate_controller else '最终控制人为（待补充）'

        # === 组织结构 ===
        org = getattr(s, 'org_structure', '')
        if org:
            r['XX部、XX部'] = org

        # === 批准报出 ===
        approver = getattr(s, 'report_approver', '')
        r['批准报出者为XXX'] = f'批准报出者为{approver}' if approver else '批准报出者为（待补充）'
        # 不要注册全局 "XXX" → 太容易误匹配

        # === 日期 ===
        r['XXXX年'] = '2025年'
        r['XXXX年度'] = '2025年度'
        r['XXXX年XX月XX日'] = '2025年12月31日'
        r['XXXX年X月X日'] = '2025年12月31日'
        r['XX月'] = '12月'
        r['XX日'] = '31日'

        # === 简称（最后设置，避免干扰其他替换） ===
        short = s.short_name or s.name
        r['本公司'] = short
        r['本集团'] = short if s.is_consolidated else s.name
        r['以下简称本公司'] = f'以下简称{short}'
        r['以下简称本集团'] = f'以下简称{short}'

    def fill_paragraph(self, para) -> int:
        """填充段落中的占位符（支持跨run替换），返回替换次数"""
        count = 0
        full_text = para.text
        sorted_repl = sorted(self.replacements.items(), key=lambda x: -len(x[0]))

        for placeholder, value in sorted_repl:
            if placeholder not in full_text:
                continue

            # Try simple single-run replacement first
            replaced = False
            for run in para.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
                    count += 1
                    replaced = True
                    full_text = para.text  # refresh
                    break

            if not replaced:
                # Cross-run replacement: merge runs containing parts of placeholder
                count += self._cross_run_replace(para, placeholder, value)
                full_text = para.text  # refresh

        return count

    def _cross_run_replace(self, para, placeholder: str, value: str) -> int:
        """替换跨多个run的占位符"""
        runs = para.runs
        n = len(runs)

        # Build cumulative text positions
        cum_lengths = [0]
        for run in runs:
            cum_lengths.append(cum_lengths[-1] + len(run.text))

        full_text = ''.join(r.text for r in runs)
        pos = full_text.find(placeholder)
        if pos < 0:
            return 0

        end_pos = pos + len(placeholder)

        # Find which runs overlap with [pos, end_pos)
        start_run_idx = 0
        while start_run_idx < n and cum_lengths[start_run_idx + 1] <= pos:
            start_run_idx += 1
        end_run_idx = start_run_idx
        while end_run_idx < n and cum_lengths[end_run_idx] < end_pos:
            end_run_idx += 1

        if start_run_idx >= n:
            return 0

        # Build replacement text for each affected run
        new_texts = []
        for ri in range(start_run_idx, min(end_run_idx + 1, n)):
            run_start = cum_lengths[ri]
            run_end = cum_lengths[ri + 1]
            overlap_start = max(pos, run_start) - run_start
            overlap_end = min(end_pos, run_end) - run_start
            old_text = runs[ri].text
            # The part of placeholder in this run
            ph_part = old_text[overlap_start:overlap_end]
            if ri == start_run_idx:
                # First affected run: replace from overlap_start to overlap_end
                new_text = old_text[:overlap_start] + value + old_text[overlap_end:]
            else:
                # Subsequent runs: their portion of placeholder is removed
                new_text = old_text[:overlap_start] + old_text[overlap_end:]
            new_texts.append(new_text)

        # Apply changes
        runs[start_run_idx].text = new_texts[0]
        for ri in range(start_run_idx + 1, min(end_run_idx + 1, n)):
            if ri - start_run_idx < len(new_texts):
                runs[ri].text = new_texts[ri - start_run_idx]
            else:
                runs[ri].text = ''

        return 1

    def fill_table_cell(self, cell, replacements: Dict[str, str] = None) -> int:
        """填充表格单元格"""
        repl = replacements or self.replacements
        count = 0
        for para in cell.paragraphs:
            full_text = para.text
            for ph, val in repl.items():
                if ph in full_text:
                    for run in para.runs:
                        if ph in run.text:
                            run.text = run.text.replace(ph, val)
                            count += 1
        return count

    def fill_docx(self, template_path: str, output_path: str) -> List[str]:
        """填充整个docx文档"""
        from docx import Document
        doc = Document(template_path)

        total_replacements = 0
        unfilled = []

        # 填充段落
        for i, para in enumerate(doc.paragraphs):
            n = self.fill_paragraph(para)
            total_replacements += n

            # 检查是否有剩余XX占位符
            remaining = re.findall(r'X{2,}', para.text)
            for r_text in remaining:
                if r_text not in self.replacements:
                    unfilled.append(f'段落[{i}]: {r_text} (上下文: {para.text[:80]})')

        # 填充表格
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    n = self.fill_table_cell(cell)
                    total_replacements += n

        # 删除红色指导文字（使用者指引）
        # 红色字体段落：模板说明、样本格式等
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    rgb = str(run.font.color.rgb)
                    if rgb == 'FF0000':  # 红色
                        run.text = ''  # 删除红色指引文字

        doc.save(output_path)
        return unfilled

    def format_amount(self, value: float, unit: str = '元') -> str:
        """格式化金额"""
        if value == 0:
            return '-'
        if unit == '万元':
            return f'{value/10000:,.2f}'
        elif unit == '亿元':
            return f'{value/100000000:,.2f}'
        else:
            return f'{value:,.2f}'


# ============================================================
# 财务报表附注填充器
# ============================================================

class NotesFiller:
    """财务报表附注专用填充器 — 处理272个表 + 段落"""

    def __init__(self, engine: DocFillEngine, entity: Entity):
        self.engine = engine
        self.entity = entity
        self.doc = None
        self.warnings = []

    def load_template(self, template_path: str):
        from docx import Document
        self.doc = Document(template_path)

    def save(self, output_path: str):
        self.doc.save(output_path)

    def fill_section_one_company_info(self):
        """填充一、公司的基本情况"""
        # Para 20-32: 公司基本信息段落
        # 这些由DocFillEngine的段落填充自动处理
        for para in self.doc.paragraphs:
            self.engine.fill_paragraph(para)

    def fill_section_four_policies(self):
        """四、重要会计政策和会计估计 — 裁剪不适用的政策"""
        # 找到Heading 1 "重要会计政策和会计估计" (para 38)
        # 遍历后续段落直到下一个Heading 1 (para 376)
        heading_found = False
        policies_to_remove = []

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()

            if '重要会计政策和会计估计' in text and para.style.name == 'Heading 1':
                heading_found = True
                continue
            if heading_found and para.style.name == 'Heading 1':
                break  # 到达下一章节

            if heading_found and text:
                # 检查段落是否引用不适用科目的政策
                for keyword, accounts in self.POLICY_ACCOUNT_MAP.items():
                    if keyword in text:
                        # 检查科目余额
                        total_balance = sum(
                            self.entity.financial_data.get(acct, 0)
                            for acct in accounts
                        )
                        if total_balance == 0:
                            policies_to_remove.append(i)

        # 删除不适用政策段落
        for i in reversed(policies_to_remove):
            para = self.doc.paragraphs[i]
            para.text = ''  # 清空段落内容
            self.warnings.append(f'已删除不适用的会计政策: {self.doc.paragraphs[i].text[:50]}')

    POLICY_ACCOUNT_MAP = {
        '交易性金融资产': ['1101'],
        '长期股权投资': ['1511'],
        '投资性房地产': ['1521'],
        '商誉': ['1711'],
        '长期待摊费用': ['1801'],
        '短期借款': ['2001'],
        '长期借款': ['2501'],
        '应付债券': ['2502'],
        '长期应付款': ['2701'],
    }

    def fill_section_seven_tables(self):
        """七、合并财务报表重要项目的说明 — 填充272个表"""
        if not self.doc:
            return

        section_started = False
        for i, para in enumerate(self.doc.paragraphs):
            if '合并财务报表重要项目的说明' in para.text and para.style.name == 'Heading 1':
                section_started = True
                continue
            if section_started and para.style.name == 'Heading 1':
                break  # 到达下一章节

        # 七节内的表格按顺序填充
        # Table 18: 货币资金 → 需要实际数据源
        # 当前Phase 1: 标记哪些表需要填充，建立映射框架

        table_index_map = self._build_table_index_map()
        return table_index_map

    def _build_table_index_map(self):
        """建立 段落标题→表格索引 的映射"""
        mapping = {}
        current_title = ''
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if text and not para.style.name.startswith('Heading'):
                # 可能是子标题（如"货币资金"、"应收账款"等）
                if len(text) < 30 and para.runs and para.runs[0].font.bold:
                    current_title = text

            # 找到下一个表格
            # （简化处理：实际需要解析XML结构）
        return mapping


# ============================================================
# 主控制器
# ============================================================

class AuditDocGenerator:
    """审计文档生成主控制器"""

    def __init__(self, template_dir: str, output_base: str):
        self.template_dir = Path(template_dir)
        self.output_base = Path(output_base)
        self.entities: List[Entity] = []
        self.engine = DocFillEngine()
        self.analyzer = TemplateAnalyzer()

    def load_entities_from_excel(self, excel_path: str) -> List[Entity]:
        """从Excel加载实体列表"""
        df = pd.read_excel(excel_path)
        entities = []
        for _, row in df.iterrows():
            e = Entity(
                name=str(row.get('单位名称', '')),
                short_name=str(row.get('简称', '')),
                uscc=str(row.get('统一社会信用代码', '')),
                level=str(row.get('层级', '子公司')),
                is_consolidated=str(row.get('是否合并', '否')) == '是',
                parent=str(row.get('母公司', '')),
                registered_capital=str(row.get('注册资本', '')),
                established_year=str(row.get('成立年份', '')),
                address=str(row.get('注册地', '')),
                industry=str(row.get('行业', '')),
                business_scope=str(row.get('经营范围', '')),
                controlling_shareholder=str(row.get('控股股东', '')),
                ultimate_controller=str(row.get('实际控制人', '')),
            )
            entities.append(e)

        self.entities = entities
        print(f'加载 {len(entities)} 个实体')
        return entities

    def precheck_templates(self) -> Dict:
        """模板预检"""
        results = {}
        for f in self.template_dir.rglob('*'):
            if f.suffix.lower() in ('.doc', '.docx') and not f.name.startswith('~$'):
                rel = str(f.relative_to(self.template_dir))
                if f.suffix == '.doc':
                    results[rel] = {'status': '需要转换为.docx', 'format': '.doc'}
                else:
                    try:
                        from docx import Document
                        doc = Document(str(f))
                        xx_count = len(re.findall(r'X{2,}', '\n'.join(p.text for p in doc.paragraphs)))
                        results[rel] = {
                            'status': 'OK',
                            'format': '.docx',
                            'xx_placeholders': xx_count,
                            'tables': len(doc.tables)
                        }
                    except Exception as e:
                        results[rel] = {'status': f'错误: {e}', 'format': '.docx'}
        return results

    def generate_for_entity(self, entity: Entity, report_type: str = '单体') -> str:
        """为单个实体生成审计文档"""
        self.engine = DocFillEngine()
        self.engine.register_entity(entity)

        output_dir = self.output_base / entity.safe_name / f'{report_type}审计报告'
        output_dir.mkdir(parents=True, exist_ok=True)

        # 选择模板
        if report_type == '合并':
            template_subdir = '1.审计报告/1-1合并审计报告模板'
        else:
            template_subdir = '1.审计报告/1-2单体审计报告模板'

        template_base = self.template_dir / template_subdir

        generated_files = []
        unfilled_all = []

        for template_file in template_base.glob('*'):
            if template_file.suffix == '.docx' and not template_file.name.startswith('~$'):
                output_file = output_dir / template_file.name
                unfilled = self.engine.fill_docx(str(template_file), str(output_file))
                generated_files.append(str(output_file))
                unfilled_all.extend(unfilled)
                print(f'  生成: {template_file.name} → {output_file.name}')

        # 汇总未填充项
        if unfilled_all:
            print(f'  ⚠️ {len(unfilled_all)} 个占位符未填充')
            for u in unfilled_all[:10]:
                print(f'    - {u}')

        return str(output_dir)


# ============================================================
# CLI入口
# ============================================================

if __name__ == '__main__':
    TEMPLATE_DIR = 'D:/Users/12844/Desktop/已下发的文档模版/2.审计文档'
    OUTPUT_BASE = 'D:/Users/12844/Desktop/审计文档输出'
    ENTITY_EXCEL = 'D:/Users/12844/Desktop/企业信息表.xlsx'

    gen = AuditDocGenerator(TEMPLATE_DIR, OUTPUT_BASE)

    # Step 1: 模板预检
    print('=' * 60)
    print('Step 1: 模板预检')
    print('=' * 60)
    results = gen.precheck_templates()
    ok = sum(1 for v in results.values() if v['status'] == 'OK')
    need_convert = sum(1 for v in results.values() if '转换' in v['status'])
    print(f'  正常: {ok}, 需转换: {need_convert}')

    # Step 2: 加载实体（如果Excel存在）
    if Path(ENTITY_EXCEL).exists():
        print(f'\nStep 2: 加载实体')
        gen.load_entities_from_excel(ENTITY_EXCEL)
        print(f'  已加载 {len(gen.entities)} 个实体')

        # Step 3: 为第一个实体生成文档（测试）
        if gen.entities:
            print(f'\nStep 3: 测试生成（{gen.entities[0].name}）')
            gen.generate_for_entity(gen.entities[0], '单体')
    else:
        print(f'\n⚠️ 企业信息表不存在: {ENTITY_EXCEL}')
        print('  请准备包含以下列的Excel:')
        print('  单位名称, 简称, 统一社会信用代码, 层级, 是否合并, 母公司,')
        print('  注册资本, 成立年份, 注册地, 行业, 经营范围, 控股股东, 实际控制人')
