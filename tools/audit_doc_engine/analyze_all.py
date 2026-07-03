"""
全量模板分析器 — 扫描所有49个.docx，建立完整占位符目录
"""
import sys, re, json
from pathlib import Path
from collections import defaultdict, Counter
from docx import Document
sys.stdout.reconfigure(encoding='utf-8')

DOCX_DIR = Path('D:/Users/12844/Desktop/已下发的文档模版/2.审计文档_docx')
OUT_JSON = Path('D:/Users/12844/WorkBuddy/audit_doc_engine/placeholder_catalog.json')


def analyze_all():
    results = {}
    total_xx = 0
    all_patterns = Counter()

    for f in sorted(DOCX_DIR.rglob('*.docx')):
        if f.name.startswith('~$'):
            continue

        rel = str(f.relative_to(DOCX_DIR))
        try:
            doc = Document(str(f))

            paras_with_xx = []
            tables_with_xx = []

            # Scan paragraphs
            for pi, para in enumerate(doc.paragraphs):
                matches = re.findall(r'X{2,}', para.text)
                for m in matches:
                    if len(set(m)) == 1:  # pure X only
                        ctx = para.text.strip()[:200]
                        paras_with_xx.append({
                            'para': pi,
                            'placeholder': m,
                            'context': ctx
                        })
                        all_patterns[m] += 1

            # Scan tables
            for ti, table in enumerate(doc.tables):
                for ri, row in enumerate(table.rows):
                    for ci, cell in enumerate(row.cells):
                        matches = re.findall(r'X{2,}', cell.text)
                        for m in matches:
                            if len(set(m)) == 1:
                                ctx = cell.text.strip()[:150]
                                tables_with_xx.append({
                                    'table': ti,
                                    'row': ri,
                                    'col': ci,
                                    'placeholder': m,
                                    'context': ctx
                                })
                                all_patterns[m] += 1

            count = len(paras_with_xx) + len(tables_with_xx)
            total_xx += count

            results[rel] = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'xx_in_paras': len(paras_with_xx),
                'xx_in_tables': len(tables_with_xx),
                'total_xx': count,
                'paras_detail': paras_with_xx[:20],  # first 20 per file
                'tables_detail': tables_with_xx[:20],
            }

            if count > 0:
                print(f'  [{count:>3} XX] {rel}')

        except Exception as e:
            print(f'  [ERROR] {rel}: {e}')

    # Build pattern→category mapping
    pattern_catalog = build_catalog(all_patterns)

    # Count categories
    cat_counts = Counter()
    for p, info in pattern_catalog.items():
        cat_counts[info['category']] += all_patterns[p]

    print(f'\n总文件: {len(results)} | 含占位符文件: {sum(1 for v in results.values() if v["total_xx"]>0)}')
    print(f'总占位符: {total_xx} 处')
    print(f'唯一模式: {len(pattern_catalog)} 种')
    print()

    print('=== 占位符分类 ===')
    for cat, count in cat_counts.most_common():
        print(f'  {cat}: {count} 处')

    # Save full catalog
    output = {
        'summary': {
            'total_files': len(results),
            'files_with_placeholders': sum(1 for v in results.values() if v['total_xx'] > 0),
            'total_placeholders': total_xx,
            'unique_patterns': len(pattern_catalog),
            'categories': dict(cat_counts),
        },
        'pattern_catalog': pattern_catalog,
        'files': {k: {
            'paras': v['paragraphs'],
            'tables': v['tables'],
            'xx_paras': v['xx_in_paras'],
            'xx_tables': v['xx_in_tables'],
            'total_xx': v['total_xx'],
        } for k, v in results.items()},
    }

    with open(OUT_JSON, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)

    print(f'\n完整目录: {OUT_JSON}')
    return results, pattern_catalog


def build_catalog(all_patterns):
    """将占位符模式映射到分类和数据源"""
    catalog = {}

    for pattern, count in all_patterns.items():
        info = classify_pattern(pattern)
        catalog[pattern] = {
            'count': count,
            'category': info['category'],
            'data_source': info['source'],
            'template_text': info['template_text'],
        }

    return catalog


def classify_pattern(pattern):
    """判断占位符的分类和数据来源"""
    plen = len(pattern)

    # === 公司信息类 ===
    if plen >= 18 and plen <= 20:
        return {'category': '公司信息', 'source': '统一社会信用代码',
                'template_text': 'XXXXXXXXXXXXXXXXXX'}
    if 'XXXX有限公司' in pattern or 'XXXX公司' in pattern:
        return {'category': '公司信息', 'source': '单位名称', 'template_text': pattern}

    # === 日期类 ===
    if 'XXXX年' in pattern and 'XX月' in pattern and 'XX日' in pattern:
        return {'category': '日期', 'source': '报告日期', 'template_text': 'XXXX年XX月XX日'}
    if 'XXXX年' in pattern and 'XX月' in pattern:
        return {'category': '日期', 'source': '报告日期', 'template_text': 'XXXX年XX月'}
    if 'XXXX年' in pattern:
        return {'category': '日期', 'source': '报告年份', 'template_text': 'XXXX年'}
    if 'XXXX' in pattern and plen == 4:
        if pattern.count('X') == 4:
            return {'category': '日期', 'source': '年份', 'template_text': 'XXXX'}

    # === 金额数字类 ===
    if 'XX万元' in pattern:
        return {'category': '财务数据', 'source': '注册资本/金额', 'template_text': 'XX万元'}
    if 'XX亿元' in pattern:
        return {'category': '财务数据', 'source': '金额(亿元)', 'template_text': 'XX亿元'}
    if 'XX元' in pattern:
        return {'category': '财务数据', 'source': '金额(元)', 'template_text': 'XX元'}

    # === 其他公司信息 ===
    if 'XX行业' in pattern:
        return {'category': '公司信息', 'source': '行业', 'template_text': 'XX行业'}
    if 'XX部' in pattern:
        return {'category': '公司信息', 'source': '组织结构', 'template_text': 'XX部'}

    # === 通用占位符 ===
    if plen == 2:
        # XX: 需要根据上下文判断
        return {'category': '需上下文判断', 'source': '根据所在段落确定', 'template_text': 'XX'}
    if plen == 3:
        return {'category': '需上下文判断', 'source': '根据所在段落确定', 'template_text': 'XXX'}

    return {'category': '其他', 'source': '待定', 'template_text': pattern}


if __name__ == '__main__':
    print('=' * 60)
    print('全量模板占位符分析')
    print('=' * 60)
    analyze_all()
