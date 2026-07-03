"""
审计报告扩展模块 — AuditDocExtensions
======================================
为 DocBuilder 提供审计报告专属内容元素。

用法:
    builder = DocBuilder()
    builder.setup_page().setup_styles()
    add_report_header(builder, "华油审字[2026]第001号", "华北油田公司")
    add_audit_opinion(builder, "无保留意见", "我们认为...")
    add_finding(builder, "1", "报销附件不合规", ...)
    builder.save("审计报告.docx")

也可子类化:
    class AuditDocBuilder(DocBuilder, AuditDocExtensions):
        pass
"""

from __future__ import annotations

from typing import Optional, List, Tuple

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .doc_builder import (
    DocBuilder,
    _set_run_font,
    _set_cell_shading,
    _set_para_border,
    COLOR_BLACK,
    COLOR_LIGHT_BG,
    COLOR_BORDER,
    FONT_CN_BODY,
    FONT_CN_HEADING,
    FONT_EN,
)


# ── 报告头 ──────────────────────────────────────────────────


def add_report_header(
    builder: DocBuilder,
    report_no: str = "",
    addressee: str = "",
    subject: str = "",
    date_text: str = "",
):
    """审计报告头部信息。

    包含报告文号、致送单位、报告主题、日期。
    """
    if report_no:
        p = builder.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_run_font(
            p.add_run(report_no), FONT_CN_BODY, FONT_EN, 11, color_hex=COLOR_BLACK
        )
        p.paragraph_format.space_after = Pt(4)

    if addressee:
        p = builder.doc.add_paragraph()
        _set_run_font(
            p.add_run(addressee + "："),
            FONT_CN_BODY,
            FONT_EN,
            12,
            color_hex=COLOR_BLACK,
        )
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    if subject:
        p = builder.doc.add_paragraph()
        _set_run_font(
            p.add_run(f"关于{subject}的审计报告"),
            FONT_CN_HEADING,
            FONT_EN,
            16,
            bold=True,
            color_hex=COLOR_BLACK,
        )
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if date_text:
        p = builder.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(
            p.add_run(date_text), FONT_CN_BODY, FONT_EN, 12, color_hex=COLOR_BLACK
        )
        p.paragraph_format.space_after = Pt(12)

    return builder


# ── 审计意见段 ──────────────────────────────────────────────


def add_audit_opinion(
    builder: DocBuilder,
    opinion_type: str = "无保留意见",
    opinion_text: str = "",
    basis_text: str = "",
):
    """审计意见段。

    Args:
        opinion_type: 无保留意见 / 保留意见 / 否定意见 / 无法表示意见
        opinion_text: 意见结论正文
        basis_text: 意见依据（可选）
    """
    builder.add_heading_1("一、审计意见")

    p = builder.doc.add_paragraph()
    _set_run_font(
        p.add_run(f"审计意见类型：{opinion_type}"),
        FONT_CN_HEADING,
        FONT_EN,
        12,
        bold=True,
        color_hex=COLOR_BLACK,
    )
    p.paragraph_format.space_after = Pt(8)

    if opinion_text:
        builder.add_body(opinion_text)
    if basis_text:
        builder.add_heading_2("（一）审计意见基础")
        builder.add_body(basis_text)

    return builder


# ── 审计范围段 ──────────────────────────────────────────────


def add_scope_section(
    builder: DocBuilder,
    scope_text: str = "",
    scope_items: Optional[List[str]] = None,
):
    """审计范围段。"""
    builder.add_heading_1("二、审计范围与方法")
    if scope_text:
        builder.add_body(scope_text)
    if scope_items:
        for item in scope_items:
            builder.add_list_item(item)
    return builder


# ── 审计发现（单条） ────────────────────────────────────────


def add_finding(
    builder: DocBuilder,
    finding_id: str,
    title: str,
    phenomenon: str,
    basis: str = "",
    impact: str = "",
    recommendation: str = "",
    amount: str = "",
):
    """单条审计发现（四要素完整）。

    Args:
        finding_id: 问题编号，如 "1"、"2-1"
        title: 问题标题
        phenomenon: 现象描述（事实）
        basis: 依据（违反的法规/制度条款）
        impact: 影响/风险
        recommendation: 建议
        amount: 涉及金额
    """
    # 问题标题行
    p = builder.doc.add_paragraph()
    _set_run_font(
        p.add_run(f"【问题{finding_id}】{title}"),
        FONT_CN_HEADING,
        FONT_EN,
        12,
        bold=True,
        color_hex=COLOR_BLACK,
    )
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    # 金额行（如有）
    if amount:
        pa = builder.doc.add_paragraph()
        _set_run_font(
            pa.add_run(f"涉及金额：{amount}"),
            FONT_CN_BODY,
            FONT_EN,
            11,
            bold=False,
            color_hex=COLOR_BLACK,
        )
        pa.paragraph_format.space_after = Pt(4)

    # 现象
    if phenomenon:
        builder.add_body(phenomenon)
    # 依据
    if basis:
        pb = builder.doc.add_paragraph()
        _set_run_font(
            pb.add_run(f"制度依据：{basis}"),
            FONT_CN_BODY,
            FONT_EN,
            11,
            color_hex=COLOR_BLACK,
        )
        pb.paragraph_format.space_after = Pt(2)
    # 影响
    if impact:
        pc = builder.doc.add_paragraph()
        _set_run_font(
            pc.add_run(f"风险影响：{impact}"),
            FONT_CN_BODY,
            FONT_EN,
            11,
            color_hex=COLOR_BLACK,
        )
        pc.paragraph_format.space_after = Pt(2)
    # 建议
    if recommendation:
        pd = builder.doc.add_paragraph()
        _set_run_font(
            pd.add_run(f"审计建议：{recommendation}"),
            FONT_CN_BODY,
            FONT_EN,
            11,
            color_hex=COLOR_BLACK,
        )
        pd.paragraph_format.space_after = Pt(4)

    # 分隔线
    psep = builder.doc.add_paragraph()
    _set_para_border(psep, "bottom", 4, COLOR_BORDER)
    psep.paragraph_format.space_after = Pt(4)

    return builder


# ── 审计发现汇总表 ──────────────────────────────────────────


def add_finding_summary(
    builder: DocBuilder,
    headers: Optional[List[str]] = None,
    rows: Optional[List[List]] = None,
    col_widths: Optional[List[float]] = None,
    caption: str = "",
):
    """审计发现汇总表。

    默认列: 序号 | 问题类型 | 涉及金额 | 违规依据（可自定义）
    """
    h = headers or ["序号", "问题类型", "涉及金额", "违规依据"]
    r = rows or []
    cw = col_widths or [1, 3, 2, 4]
    if not caption:
        builder.add_heading_2("审计发现汇总表")
    builder.add_table(headers=h, rows=r, col_widths=cw, caption=caption)
    return builder


# ── 管理建议 ────────────────────────────────────────────────


def add_recommendation(
    builder: DocBuilder,
    rec_id: str,
    title: str,
    current_status: str,
    impact: str,
    suggestion: str,
):
    """单条管理建议（现状→影响→建议）。"""
    p = builder.doc.add_paragraph()
    _set_run_font(
        p.add_run(f"{rec_id}. {title}"),
        FONT_CN_HEADING,
        FONT_EN,
        12,
        bold=True,
        color_hex=COLOR_BLACK,
    )
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

    builder.add_body(f"现状：{current_status}")
    builder.add_body(f"影响：{impact}")
    builder.add_body(f"建议：{suggestion}")
    return builder


# ── 责任段 ──────────────────────────────────────────────────


def add_responsibility_section(
    builder: DocBuilder,
    mgmt_resp: str = "",
    auditor_resp: str = "",
):
    """责任段（管理层责任 + 注册会计师责任）。"""
    builder.add_heading_1("三、责任")
    if not mgmt_resp:
        mgmt_resp = (
            "被审计单位管理层负责按照《企业内部控制基本规范》及相关规定建立健全内部控制，"
            "保证其有效运行，确保财务报表的真实性和完整性。"
        )
    if not auditor_resp:
        auditor_resp = (
            "我们的责任是在执行审计工作的基础上对财务报表发表审计意见。"
            "我们按照中国注册会计师审计准则的规定执行了审计工作。"
            "审计准则要求我们遵守职业道德规范，计划和执行审计工作以对财务报表是否不存在重大错报获取合理保证。"
        )
    builder.add_heading_2("（一）管理层责任")
    builder.add_body(mgmt_resp)
    builder.add_heading_2("（二）注册会计师责任")
    builder.add_body(auditor_resp)
    return builder


# ── 审计签字 ────────────────────────────────────────────────


def add_audit_signature(
    builder: DocBuilder,
    firm: str = "",
    cpa1: str = "",
    cpa2: str = "",
    date_text: str = "",
):
    """审计报告签字页。"""
    builder.add_space(30)

    # 会计师事务所
    if firm:
        p = builder.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_run_font(
            p.add_run(f"会计师事务所（盖章）：{firm}"),
            FONT_CN_BODY,
            FONT_EN,
            12,
            color_hex=COLOR_BLACK,
        )
        p.paragraph_format.space_after = Pt(18)

    if cpa1 or cpa2:
        p = builder.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_run_font(
            p.add_run(f"中国注册会计师：{cpa1}"),
            FONT_CN_BODY,
            FONT_EN,
            12,
            color_hex=COLOR_BLACK,
        )
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(18)

        p = builder.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_run_font(
            p.add_run(f"中国注册会计师：{cpa2}"),
            FONT_CN_BODY,
            FONT_EN,
            12,
            color_hex=COLOR_BLACK,
        )
        p.paragraph_format.space_after = Pt(18)

    if date_text:
        ps = builder.doc.add_paragraph()
        ps.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_run_font(
            ps.add_run(date_text),
            FONT_CN_BODY,
            FONT_EN,
            12,
            color_hex=COLOR_BLACK,
        )
        ps.paragraph_format.space_before = Pt(24)

    return builder


# ============================================================
# 子类方式（二选一：函数式 或 继承式）
# ============================================================


class AuditDocBuilder(DocBuilder):
    """带审计扩展的DocBuilder（继承式）。"""

    def add_report_header(
        self,
        report_no: str = "",
        addressee: str = "",
        subject: str = "",
        date_text: str = "",
    ):
        return add_report_header(self, report_no, addressee, subject, date_text)

    def add_audit_opinion(
        self,
        opinion_type: str = "无保留意见",
        opinion_text: str = "",
        basis_text: str = "",
    ):
        return add_audit_opinion(self, opinion_type, opinion_text, basis_text)

    def add_scope_section(
        self,
        scope_text: str = "",
        scope_items: Optional[List[str]] = None,
    ):
        return add_scope_section(self, scope_text, scope_items)

    def add_finding(
        self,
        finding_id: str,
        title: str,
        phenomenon: str,
        basis: str = "",
        impact: str = "",
        recommendation: str = "",
        amount: str = "",
    ):
        return add_finding(
            self, finding_id, title, phenomenon, basis, impact, recommendation, amount
        )

    def add_finding_summary(
        self,
        headers: Optional[List[str]] = None,
        rows: Optional[List[List]] = None,
        col_widths: Optional[List[float]] = None,
        caption: str = "",
    ):
        return add_finding_summary(self, headers, rows, col_widths, caption)

    def add_recommendation(
        self,
        rec_id: str,
        title: str,
        current_status: str,
        impact: str,
        suggestion: str,
    ):
        return add_recommendation(
            self, rec_id, title, current_status, impact, suggestion
        )

    def add_responsibility_section(
        self,
        mgmt_resp: str = "",
        auditor_resp: str = "",
    ):
        return add_responsibility_section(self, mgmt_resp, auditor_resp)

    def add_audit_signature(
        self,
        firm: str = "",
        cpa1: str = "",
        cpa2: str = "",
        date_text: str = "",
    ):
        return add_audit_signature(self, firm, cpa1, cpa2, date_text)
